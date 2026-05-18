#!/usr/bin/env python3
"""LTSA 기반 유상 서비스 연계 방향 — 워드 제안서 생성기.

회사 주제(2030 매출 증대 중장기 계획 + 업무 프로세스 개선)에 답하는 후속 제안.
'24/7 무상 기술지원 웹사이트 구축 제안'(선행)을 전제로,
무상 기술지원 → AI 케이스 인계 → HI-CMS 알림 → LTSA 자연 전환의 4축을 설계.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE_DIR, "LTSA 유상 서비스 연계 방향 제안.docx")

PRIMARY   = '1B3A5C'
SECONDARY = '2B579A'
LIGHT_BG  = 'E8F0FE'
LIGHT     = 'F0F4F8'
ACCENT    = 'D35400'   # subtle accent for numbers
RED       = 'E74C3C'
ORANGE    = 'F39C12'
GREEN     = '27AE60'

# ── style helpers (mirror generate_free_support_web.py for visual continuity) ──

def init_doc(doc):
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    style.paragraph_format.line_spacing = 1.2
    style.paragraph_format.space_after = Pt(4)

    for level, (size, color) in enumerate([
        (20, PRIMARY), (15, PRIMARY), (12, SECONDARY),
    ], 1):
        h = doc.styles[f'Heading {level}']
        h.font.name = '맑은 고딕'
        h.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        h.font.size = Pt(size)
        h.font.bold = True
        r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
        h.font.color.rgb = RGBColor(r, g, b)
        h.paragraph_format.space_before = Pt(16 if level == 1 else 10)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)


def shade(cell, color):
    tcPr = cell._element.get_or_add_tcPr()
    s = tcPr.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    tcPr.append(s)


def set_cell(cell, text, *, bold=False, color=None, size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(str(text))
    r.font.size = size
    r.font.bold = bold
    r.font.name = '맑은 고딕'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color is not None:
        r.font.color.rgb = color


def keep_together(table):
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trPr.append(trPr.makeelement(qn('w:cantSplit'), {}))
    for row in table.rows[:-1]:
        for cell in row.cells:
            for p in cell.paragraphs:
                pPr = p._p.get_or_add_pPr()
                pPr.append(pPr.makeelement(qn('w:keepNext'), {}))
                pPr.append(pPr.makeelement(qn('w:keepLines'), {}))


def spacer(doc, pt=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(pt)
    run = p.add_run('')
    run.font.size = Pt(1)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, PRIMARY)
        set_cell(c, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                 align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                shade(c, LIGHT)
            if isinstance(val, str) and val.startswith('**') and val.endswith('**'):
                set_cell(c, val[2:-2], bold=True, size=Pt(9))
            else:
                set_cell(c, val, size=Pt(9))
    keep_together(t)
    spacer(doc, 10)
    return t


def add_ba_table(doc, rows):
    """Before/After comparison table."""
    t = doc.add_table(rows=1 + len(rows), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['항목', 'Before (현재)', 'After (제안 후)']):
        c = t.rows[0].cells[i]
        shade(c, PRIMARY)
        set_cell(c, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                 align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10))
    for ri, (item, before, after) in enumerate(rows):
        set_cell(t.rows[ri + 1].cells[0], item, bold=True, size=Pt(9))
        cb = t.rows[ri + 1].cells[1]
        shade(cb, 'F5F5F5')
        set_cell(cb, before, size=Pt(9), color=RGBColor(0x77, 0x77, 0x77))
        ca = t.rows[ri + 1].cells[2]
        shade(ca, 'EAFAF1')
        set_cell(ca, after, size=Pt(9), bold=True, color=RGBColor(0x1B, 0x3A, 0x5C))
    keep_together(t)
    spacer(doc, 10)
    return t


def title_block(doc):
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.add_run('━' * 41).font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('LTSA 기반 유상 서비스 연계 방향\n2030년 안정 매출 구조와 업무 프로세스 개선 통합 제안')
    r.font.name = '맑은 고딕'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    r.font.size = Pt(24); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run('무상 기술지원 → AI 케이스 인계 → HI-CMS 알림 → LTSA 자연 전환')
    rs.font.size = Pt(12); rs.font.bold = True
    rs.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = sub2.add_run('— 24/7 무상 기술지원 웹사이트 구축 제안의 후속, '
                       '회사 주제(중장기 매출 증대 + 업무 프로세스 개선)에 대한 통합 답변 —')
    rs2.font.size = Pt(10); rs2.font.italic = True
    rs2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    rule2 = doc.add_paragraph()
    rule2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule2.add_run('━' * 41).font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run('2026. 04\nHD현대마린솔루션테크')
    rm.font.size = Pt(11); rm.font.bold = True
    rm.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)


# ── content sections -----------------------------------------------------

def add_executive_summary(doc):
    doc.add_heading('Executive Summary — 한 장으로 보는 제안', level=1)
    doc.add_paragraph(
        '본 제안은 회사가 내려준 두 가지 주제 — "2030년까지 매출 증대를 위한 중장기 계획" '
        '과 "업무 프로세스 개선" — 을 단일 메커니즘으로 통합한다. '
        '핵심 명제는 "무상 보증기간 동안 축적되는 기술지원 이력을 자산화해, '
        'LTSA(Long Term Service Agreement, 장기 서비스 계약) 형태의 안정 매출로 '
        '자연 전환시킨다" 는 것. 매출 증대(중장기 계획)는 결과이고, '
        '업무 프로세스 개선(AI 기반 PM 인계 자동화)은 그 과정에서 부수적으로 달성된다.')

    doc.add_paragraph(
        '본 제안은 선행 제안인 "24/7 무상 기술지원 웹사이트 구축 제안"의 후속이며, '
        '무상 웹 포털을 "데이터 입구", LTSA를 "유상 출구", 그 사이를 잇는 4가지 '
        '메커니즘 — 무상 기술지원 / AI 케이스 인계 / HI-CMS 통합 알림 / LTSA 전환 — '
        '을 설계한다.')

    add_table(doc, ['축', '메커니즘', '현재 상태', '2030년 목표'], [
        ['입구',  '무상 기술지원 포털',  '메일·전화에 산재',          '전 선주 단일 창구'],
        ['통로 1', 'AI 케이스 인계',     'PM이 메일 수기 정리(수일)',   '자동 요약 푸시(분 단위)'],
        ['통로 2', 'HI-CMS 통합 알림',   '사내 분석용·선주 미전달',    '이상 감지→웹 알림→케이스 자동 생성'],
        ['출구',  'LTSA 자연 전환',     '단발 유상 (예측 불가)',       '구독+사건 하이브리드 (예측 가능)'],
    ])

    doc.add_paragraph(
        '※ 본 제안은 신규 인력·신규 시스템 일괄 투입을 요구하지 않음. '
        '24/7 무상 기술지원 웹사이트(선행 제안)가 구축된 시점부터 "위에 얹는" 형태로 '
        '단계적 적용 가능. 2026년 PoC, 2027년 정식 운영, 2030년 LTSA 비중 50%↑ 도달이 목표.'
    ).runs[0].font.italic = True


def add_intro(doc):
    doc.add_heading('서론 — 왜 LTSA 연계인가', level=1)

    doc.add_heading('회사 주제 두 가지를 어떻게 하나의 답변으로 묶는가', level=2)
    doc.add_paragraph(
        '회사는 "2030년까지 매출 증대를 위한 중장기 계획" 또는 "업무 프로세스 개선" '
        '두 가지 주제를 제시했다. 본 제안은 "또는"이 아니라 "그리고"로 답한다. '
        '중장기 매출 증대를 핵심 골자로 두되, 그것을 달성하기 위한 수단으로 '
        '업무 프로세스 개선(AI 기반 PM 인계, HI-CMS 통합 알림)이 자연스럽게 포함된다.')
    doc.add_paragraph('▸  중장기 계획의 끝점 — 2030년 LTSA 매출 비중 50% 이상')
    doc.add_paragraph('▸  업무 프로세스 개선 — 그 끝점에 도달하기 위한 PM·기술지원·HI-CMS 운영 효율화')
    doc.add_paragraph('▸  통합 메커니즘 — 두 가지가 하나의 데이터 흐름(무상→AI인계→LTSA)에서 동시 달성')

    doc.add_heading('선행 제안과의 관계', level=2)
    doc.add_paragraph(
        '본 제안은 "24/7 무상 기술지원 웹사이트 구축 제안" (2026.04, 사내자료) 을 전제로 한다. '
        '선행 제안이 "선주를 어떻게 우리 웹으로 자연스럽게 끌어오는가"에 답했다면, '
        '본 제안은 그 다음 질문 — "끌어온 다음, 어떻게 안정 매출 구조로 연결하는가" — 에 답한다.')
    doc.add_paragraph(
        '즉 선행 제안 = "무상 웹의 진입로 설계", 본 제안 = "무상 웹 → 유상 LTSA 전환로 설계".'
    ).runs[0].font.bold = True

    doc.add_heading('본 제안이 답하는 5개 질문', level=2)
    add_table(doc, ['#', '질문', '답변 위치'], [
        ['Q1', '왜 LTSA인가? OEM 3사가 30년 걸려 도달한 구조의 본질',          '본론 1'],
        ['Q2', '무상 → 유상 전환 단절은 어떻게 해소하는가',                      '본론 2'],
        ['Q3', 'AI 케이스 인계는 어떻게 작동하며 PM 부담을 얼마나 줄이는가',     '본론 3'],
        ['Q4', 'HI-CMS는 어떻게 선주를 끌어오는 알림 채널이 되는가',             '본론 4'],
        ['Q5', '2030년까지 어떻게 도달하는가 (단계별 매출 시뮬레이션)',           '본론 5'],
    ])


def add_section1(doc):
    doc.add_heading('본론 1 — 현황: 무상-유상 단절과 매출 불안정', level=1)

    doc.add_heading('1-1. HMS의 현재 구조 — 무상에 머무르는 데이터, 단절된 유상', level=2)
    doc.add_paragraph(
        'HD현대마린솔루션테크는 HD현대 조선소가 건조한 약 9,890척 규모의 선박에 대해 '
        '무상 보증기간(통상 1~2년) 동안 기술지원·CLAIM 처리·선박 모니터링을 제공한다. '
        '문제는 보증기간이 끝나는 시점에 "유상 서비스로의 자연스러운 전환로"가 없다는 점.')
    doc.add_paragraph('▸  무상 기간 동안 축적된 기술지원·CLAIM·모니터링 이력은 메일·전화·개별 PC에 산재')
    doc.add_paragraph('▸  보증 종료 시점에 선주에게 "다음 단계 옵션(유상 LCA·구독·사건 기반)"이 제안되지 않음')
    doc.add_paragraph('▸  결과적으로 유상 서비스 매출은 단발성 사고·고장 대응에 의존 — 예측·계획 불가')

    doc.add_heading('1-2. PM 조직의 어려움 — 케이스 이력 재구성 부담', level=2)
    doc.add_paragraph(
        '유상 서비스 요청이 들어왔을 때, PM(Project Manager, 유상 서비스 조직)은 '
        '해당 선박의 무상 기술지원 이력에 직접 접근할 수 없다. 결과적으로 한 척의 선박에 '
        '대해 PM이 다음과 같은 절차를 매번 반복한다.')
    add_table(doc, ['단계', 'PM이 하는 일', '소요 시간 (현재)'], [
        ['1', '담당 기술지원 엔지니어 식별',         '0.5~1일'],
        ['2', '메일·전화 이력 수집 (개인 PC·메일함 산재)', '1~3일'],
        ['3', '주요 케이스 정리·진단 내용 재구성',     '0.5~2일'],
        ['4', '선박 데이터(Hi-4S·SmartCare·HI-CMS) 별도 조회', '0.5~1일'],
        ['5', '유상 서비스 견적·범위 산정',            '0.5~1일'],
        ['', '**합계**',                                '**3~8일 / 척**'],
    ])
    doc.add_paragraph(
        '결과적으로 PM 1인이 동시에 처리 가능한 선박 수가 제한되며, '
        '응답 지연 → 선주 만족도 ↓ → 유상 전환율 ↓ → 매출 손실로 이어진다. '
        '이는 "무상-유상 사이의 데이터 단절" 이라는 구조적 문제이며, '
        '인력 추가만으로는 근본 해결이 불가능하다.'
    ).runs[0].font.bold = True

    doc.add_heading('1-3. 매출 구조의 불안정 — 단발성 vs 안정 LTSA', level=2)
    doc.add_paragraph(
        '바르질라·콩스버그·ABB 등 OEM 3사는 매출의 40~60%를 LCA(Lifecycle Agreement, '
        'OEM에서 LTSA를 부르는 명칭)에서 확보한다. 연 단위·다년 계약 기반으로 '
        '예측 가능하고 안정적이다. 반면 HMS의 현재 유상 매출은 사고·고장 시 일회성 청구가 '
        '대부분 — 분기·연도별 변동성이 크다.')
    add_table(doc, ['지표', 'OEM 3사 평균', 'HMS 현재 (추정)'], [
        ['LTSA/LCA 매출 비중',         '40~60%',     '5~10%'],
        ['연 단위 매출 예측 정확도',    '±5%',         '±25%'],
        ['선박당 연 평균 LCA 수익',     '확보',        '미확보 (사건 기반)'],
        ['장기 고객 이탈률',           '<5%',         '측정 불가 (데이터 산재)'],
    ])
    doc.add_paragraph(
        '본 제안의 목적은 LTSA 비중을 2030년까지 OEM 3사 수준(40~50%)으로 끌어올려 '
        '매출 구조의 안정성과 예측 가능성을 동시에 확보하는 것이다.'
    ).runs[0].font.bold = True


def add_section2(doc):
    doc.add_heading('본론 2 — LTSA 연계 모델: 무상에서 유상으로의 흐름 설계', level=1)

    doc.add_heading('2-1. 단계별 흐름 — "무상 입구 → 데이터 자산 → 유상 출구"', level=2)
    doc.add_paragraph(
        '본 제안의 핵심은 "선주가 무상 보증기간 동안 자연스럽게 사용한 기술지원 포털 자체가 '
        '유상 LTSA의 영업 도구가 된다" 는 점이다. 별도의 영업 캠페인이나 신규 접점 발굴이 '
        '필요 없다. 4단계로 설계.')
    add_table(doc, ['단계', '기간', '선주 입장', 'HMS 입장'], [
        ['Stage 1 · 도입',     '인도~보증 1년차',    '무상 기술지원 포털 사용 시작', '데이터 자산 축적 시작'],
        ['Stage 2 · 신뢰',     '보증 1~2년차',      '응답 속도·이력 일원화 체감',    'AI 케이스 인계 가동'],
        ['Stage 3 · 전환 제안', '보증 종료 6개월 전', 'LTSA 옵션 자동 안내 (포털)',   '예측 매출로 사전 편입'],
        ['Stage 4 · LTSA 운영', '보증 종료 후 N년',   '구독·사건 LTSA 사용 중',        '안정 매출 확보·갱신 유도'],
    ])

    doc.add_heading('2-2. LTSA 옵션 비교 — 구독 vs 사건 기반 vs 하이브리드', level=2)
    doc.add_paragraph(
        '유상 전환 시점에서 어떤 LTSA 형태가 효율적인지 — 본 제안의 분석 결과는 '
        '"하이브리드(기본 구독 + 초과 사건 청구)"가 선주·HMS 양쪽에 가장 유리하다는 것.')
    add_table(doc, ['모델', '구조', '선주 부담', 'HMS 매출 안정성', '권고'], [
        ['A · 구독형',     '연 정액 (선박당)', '예측 가능 / 사용량 무관 부담', '매우 안정 (예측 ±5%)', '소형 선주 적합'],
        ['B · 사건 기반',   '발생 시 청구',     '저사용 시 유리 / 변동',         '불안정 (현재 구조)',     '대형·자가 정비 선주 적합'],
        ['C · 하이브리드',  '기본료+초과 사용',  '예측 가능 + 사용량 반영',       '안정 (기본료 확보)',     '**권고 (전 선주 기본)**'],
    ])
    doc.add_paragraph(
        '하이브리드 모델은 "기본 구독료로 일정 매출을 확보"하면서도 '
        '"실제 발생한 기술지원·원격진단·부품·엔지니어 파견"은 별도 정산하는 구조. '
        '선주는 갑작스런 대형 청구에 노출되지 않고, HMS는 매출 예측 가능성을 확보한다.')

    doc.add_heading('2-3. 전환 트리거 — "보증 종료" 만이 아니다', level=2)
    doc.add_paragraph(
        'LTSA 전환은 보증기간 종료 시점에 한 번만 일어나는 이벤트가 아니다. '
        '아래 5가지 트리거가 활성화될 때마다 포털에서 자동으로 LTSA 제안이 노출된다.')
    add_table(doc, ['트리거', '발생 시점', '전환 권유 메시지 예시'], [
        ['T1 · 보증 종료',          '인도+1~2년',           '보증 종료 6개월 전 알림 + LTSA 옵션'],
        ['T2 · 고빈도 기술 문의',    '월 3건 이상 누적',      '"이 정도면 정기 LTSA가 더 유리합니다"'],
        ['T3 · 이상 데이터 감지',    'HI-CMS 임계 초과',     '"예측정비 LTSA로 사전 대응 가능"'],
        ['T4 · 사고·정비 대형 케이스', '1건당 비용 임계 초과',  '"향후 동일 위험 LTSA로 흡수 가능"'],
        ['T5 · 선주 신규 선박 인도',  '동일 선주 추가 인도 시', '"플릿 단위 LTSA 패키지 제공"'],
    ])


def add_section3(doc):
    doc.add_heading('본론 3 — AI 케이스 인계 자동화 (업무 프로세스 개선의 핵심)', level=1)

    doc.add_heading('3-1. 문제 재정의 — "PM이 메일을 뒤지는 시간"이 매출 손실', level=2)
    doc.add_paragraph(
        '본론 1-2에서 정량화된 "PM 1척당 3~8일 케이스 재구성 시간"은 단순한 효율 문제가 아니다. '
        '그 시간만큼 LTSA 제안이 늦어지고, 그 사이 선주는 "이번에도 단발성 견적으로 처리"하는 '
        '의사결정으로 기울어진다. 즉 PM 인계 지연 = LTSA 전환 기회 손실.')

    doc.add_heading('3-2. 제안 구조 — 케이스 종료 시점 AI 자동 요약', level=2)
    doc.add_paragraph(
        '무상 기술지원 포털에서 케이스가 종료되거나 일정 기간 누적되는 시점에, '
        'LLM 기반 요약 에이전트가 다음 항목을 자동 생성하여 PM 대시보드로 푸시한다.')
    doc.add_paragraph('▸  장비·부위 — 어떤 장비의 어느 부위인지 (계통 분류)')
    doc.add_paragraph('▸  증상 요약 — 선주 호소 내용 + 관측된 데이터')
    doc.add_paragraph('▸  진단 결과 — Tier 1~4 단계별 분석 종합')
    doc.add_paragraph('▸  임시 조치 — 무상 단계에서 적용된 응급·단기 조치')
    doc.add_paragraph('▸  잔여 리스크 — 재발 가능성·연관 시스템 영향 평가')
    doc.add_paragraph('▸  LTSA 권고 액션 — 정기 점검 / 부품 교체 / 모니터링 강화 등')
    doc.add_paragraph('▸  매출 기회 추정 — 권고 액션 수행 시 예상 LTSA 매출')

    doc.add_heading('3-3. 흐름도 — 케이스 → AI → PM → LTSA', level=2)
    doc.add_paragraph(
        '아래는 무상 기술지원 케이스가 닫히는 시점부터 LTSA 제안이 발송되기까지의 흐름.')
    add_table(doc, ['단계', '주체', '소요 시간 (목표)', '산출물'], [
        ['1 · 케이스 종료',     '기술지원 엔지니어',  '0',          '케이스 자동 클로즈'],
        ['2 · AI 요약 생성',    'AI 에이전트',       '5~10초',      '7개 항목 정리본'],
        ['3 · PM 대시보드 푸시', '시스템',           '실시간',       '담당 PM 알림'],
        ['4 · PM 검토',        'PM',               '5~15분',      '권고 액션 검토·승인'],
        ['5 · LTSA 제안 발송',  'PM (자동 템플릿)',  '실시간',       '선주 포털 + 메일'],
        ['', '**총 소요**',                          '**15~30분**', '**(현재 3~8일 대비 99% 감축)**'],
    ])

    doc.add_heading('3-4. 정량 효과 — PM 처리 능력과 매출 기회', level=2)
    add_ba_table(doc, [
        ('PM 1인당 동시 처리 선박',     '20~30척 (수기 한계)',           '150~200척 (AI 인계)'),
        ('케이스 인계 소요 시간',        '3~8일 / 척',                   '15~30분 / 척'),
        ('LTSA 제안 발송 지연',          '평균 2주 (수기 정리 후)',         '실시간 (케이스 종료 직후)'),
        ('PM 1인당 연 LTSA 제안 건수',   '40~60건',                       '300~400건'),
        ('AI 요약 신뢰도 (검토 소요 시)', '해당 없음',                      '95%+ (PM 검토 5분 이내)'),
        ('영업 기회 손실률',             '추정 30~50% (지연 기인)',        '5% 이하'),
    ])
    doc.add_paragraph(
        '※ 위 수치는 보수적 가정. 실제 PoC를 통해 검증할 예정. 다만 "PM이 메일을 뒤지는 시간을 '
        'AI가 대신한다"는 구조 자체는 일반화된 RAG·LLM 워크플로로 충분히 구현 가능한 수준.'
    ).runs[0].font.italic = True


def add_section4(doc):
    doc.add_heading('본론 4 — HI-CMS 통합 알림: 선주를 끌어오는 자석', level=1)

    doc.add_heading('4-1. HI-CMS는 무엇인가', level=2)
    doc.add_paragraph(
        'HI-CMS(HD현대마린솔루션테크 통합 모니터링 시스템)는 두 가지 사내 시스템을 통합한 결과물.')
    add_table(doc, ['구성', '약어 풀이', '담당 영역'], [
        ['HI-ODS', 'Operational Data Service',     '기관·연료·운항 등 운영 데이터 통합 수집·저장'],
        ['ISS',    'Integrated Smart Ship',         '자동화·항해·제어 시스템 통합'],
        ['HI-CMS', 'HI-ODS + ISS 통합',              '선박 운영 전 영역 통합 모니터링·분석 플랫폼'],
    ])
    doc.add_paragraph(
        '문제는 HI-CMS의 풍부한 데이터가 현재 사내 분석·R&D 용도에 머물러 있고, '
        '선주에게는 가공된 알림으로 전달되지 않는다는 점이다. 즉 "데이터는 있는데 선주가 모름".')

    doc.add_heading('4-2. 통합 흐름 — HI-CMS → 웹 알림 → 케이스 → LTSA', level=2)
    doc.add_paragraph('HI-CMS의 이상 감지 신호가 선주의 행동을 어떻게 유도하는지의 4단계 흐름.')
    add_table(doc, ['단계', '시스템 동작', '선주 경험'], [
        ['1 · 이상 감지', 'HI-CMS 임계치 초과·이상 패턴 학습 모델 알림',  '(선주는 모름)'],
        ['2 · 웹 알림',    '선주 포털 + 모바일 푸시 + 이메일',              '"엔진 진동 평소보다 +18%" 알림 클릭'],
        ['3 · 케이스 자동 생성', '관련 데이터·이력 자동 첨부된 케이스 초안',       '클릭 한 번으로 기술지원 시작'],
        ['4 · LTSA 트리거',   'AI 인계 → PM이 예측정비 LTSA 제안 발송',      '"이 위험을 LTSA로 흡수 가능" 안내 수신'],
    ])

    doc.add_heading('4-3. 알림 트리거 예시 (실제 적용 가능 시나리오)', level=2)
    add_table(doc, ['트리거', '감지 데이터', '선주 알림 메시지', 'LTSA 연결 가능성'], [
        ['연료 효율 저하',  '주행거리/연료 소비 비율 5%↓ 지속',  '"연료 효율 5% 저하 감지 — 점검 권고"',  '높음 (정기 정비 LTSA)'],
        ['엔진 진동 증가',  '진동 센서 +15% 이상',                '"엔진 진동 평소 대비 +18%"',           '매우 높음 (예측정비 LTSA)'],
        ['윤활유 이상',    '오일 분석 + 압력 센서 패턴',           '"윤활유 점도 이상 — 샘플 권고"',         '높음 (소모품 LTSA)'],
        ['DPS 결함 신호',  '제어 시스템 자가진단 코드',            '"DPS 모듈 자가진단 경고 발생"',         '매우 높음 (긴급 LTSA + 부품)'],
        ['배기·NOx 임계',  '배기 가스 센서 + 운항 패턴',          '"NOx 규제 임계 접근 — 사전 점검 권고"',  '중간 (규제 대응 LTSA)'],
    ])
    doc.add_paragraph(
        '핵심: HI-CMS 알림은 "선주가 안 와도 됐던 영역"을 "선주가 즉시 확인하게 만드는 영역"으로 '
        '바꾼다. 알림 클릭 → 케이스 생성 → 무상 기술지원 → 데이터 자산화 → LTSA 제안의 흐름이 '
        '선주 입장에서는 "한 번의 클릭"으로 시작된다.'
    ).runs[0].font.bold = True


def add_section5(doc):
    doc.add_heading('본론 5 — 2030 중장기 로드맵과 매출 시뮬레이션', level=1)

    doc.add_heading('5-1. 5-Phase 로드맵 — 2026~2030', level=2)
    add_table(doc, ['Phase', '시기', '핵심 과업', 'KPI'], [
        ['Phase 1 · MVP',     '2026 · 1~2분기',
         '무상 웹 + AI 인계 PoC + HI-CMS 알림 1종',
         '10~20개사 Closed Beta'],
        ['Phase 2 · 정착',    '2026 · 3~4분기',
         'AI 요약 신뢰도 95%↑ / HI-CMS 알림 5종 / LTSA 제안 자동화',
         '월 LTSA 제안 100건↑'],
        ['Phase 3 · 확장',    '2027~2028',
         '전 무상 보증 선주 적용 / 구독+사건 하이브리드 정식 출시',
         'LTSA 비중 20~30% 도달'],
        ['Phase 4 · 통합',    '2028~2029',
         'HI-CMS 알림 전 선박·전 계통 / 예측정비 LTSA 본격 운영',
         'LTSA 비중 35~40%'],
        ['Phase 5 · 안정',    '2030',
         '플릿 단위 LTSA 패키지 / OEM 3사 수준 매출 안정성',
         'LTSA 비중 50%↑'],
    ])

    doc.add_heading('5-2. 매출 시뮬레이션 (보수 기준 가정)', level=2)
    doc.add_paragraph(
        '아래 시뮬레이션은 다음 가정에 근거. (a) 적용 대상 선박 9,890척 중 '
        'LTSA 적용 가능 영역은 약 60% (5,900척), (b) 선박당 연 LTSA 단가 가정치는 '
        'OEM 3사 LCA 평균(외부 자료 기준)의 70% 수준, (c) 비중 % 는 HMS 전체 '
        '서비스 매출 대비 LTSA 매출 비율.')
    add_table(doc, ['연도', 'LTSA 가입 척수 (누적)', 'LTSA 매출 비중', '매출 안정성 (예측 정확도)'], [
        ['2026', '50~100척 (Pilot)',           '5% 이하 (현재 수준 유지)', '±25%'],
        ['2027', '500~800척',                  '10~15%',                 '±20%'],
        ['2028', '1,500~2,000척',              '20~30%',                 '±15%'],
        ['2029', '2,500~3,500척',              '35~40%',                 '±10%'],
        ['2030', '4,000~5,000척',              '**50%↑**',               '**±5% (OEM 3사 수준)**'],
    ])
    doc.add_paragraph(
        '※ 위 시뮬레이션은 본 제안 단계의 가정 모델이며, Phase 1 PoC 종료 시점에 '
        '실측 데이터 기반으로 재보정 예정. 다만 "LTSA 비중을 50%로 끌어올리는 데 '
        '필요한 단일 메커니즘이 무상 웹 + AI 인계 + HI-CMS 알림" 이라는 구조는 변하지 않는다.'
    ).runs[0].font.italic = True

    doc.add_heading('5-3. 단계별 의사결정 게이트', level=2)
    add_table(doc, ['게이트', '시점', '의사결정 항목', 'Go/No-Go 기준'], [
        ['G1', 'Phase 1 종료',  'AI 요약 신뢰도, PM 수용도',     '신뢰도 90%↑ + PM 만족도 4.0/5.0↑'],
        ['G2', 'Phase 2 종료',  '하이브리드 LTSA 정식 출시 여부',  'Closed Beta 갱신율 80%↑'],
        ['G3', 'Phase 3 종료',  '전 선박 확대 / 인력 보강 규모',  'LTSA 비중 20%↑ + 매출 예측 ±15%↓'],
        ['G4', 'Phase 4 종료',  '플릿 단위 LTSA 패키지 도입 여부',  'LTSA 비중 35%↑'],
    ])


def add_section6(doc):
    doc.add_heading('본론 6 — 업무 프로세스 개선 효과 (부수적 가치)', level=1)
    doc.add_paragraph(
        '본 제안의 1차 목표는 매출 증대(LTSA 비중 50%↑)이지만, 그 과정에서 발생하는 '
        '업무 프로세스 개선 효과 또한 회사 주제의 직접 답이 된다.')

    doc.add_heading('6-1. R&R 명확화', level=2)
    add_ba_table(doc, [
        ('기술지원 엔지니어 ↔ PM 인계',       '메일 포워딩·구두 인수',                'AI 자동 요약 + 대시보드 알림'),
        ('PM ↔ 영업 (LTSA 제안)',             '월간 회의에서 케이스 공유',              '케이스 종료 즉시 LTSA 제안 자동 생성'),
        ('영업 ↔ 선주 (전환 시점)',            '단발성 견적 발송',                     '포털에서 LTSA 옵션 자동 노출'),
        ('R&D ↔ 현장 (HI-CMS 데이터)',         '사내 분석에 머무름',                   '이상 감지 → 선주 알림 자동 발송'),
    ])

    doc.add_heading('6-2. 데이터 기반 KPI 측정', level=2)
    doc.add_paragraph(
        '현재는 "어떤 기술지원이 LTSA로 전환됐는가"를 사후 분석조차 하기 어렵다. '
        '본 제안 후에는 케이스 단위로 다음 KPI가 자동 측정된다.')
    doc.add_paragraph('▸  케이스 → LTSA 전환율 (per case, per 선주, per 선박)')
    doc.add_paragraph('▸  AI 요약 정확도 (PM 수정률 기준)')
    doc.add_paragraph('▸  HI-CMS 알림 → 케이스 생성 전환율')
    doc.add_paragraph('▸  LTSA 갱신율·이탈율 (선주 단위)')
    doc.add_paragraph('▸  Tier별 응답 SLA 준수율 (선행 제안의 핵심 KPI 연계)')

    doc.add_heading('6-3. Before / After 종합', level=2)
    add_ba_table(doc, [
        ('PM 1척당 케이스 재구성 시간',  '3~8일 (메일 수기 정리)',     '15~30분 (AI 자동 인계)'),
        ('LTSA 제안 발송 지연',           '평균 2주~1개월',              '실시간 (케이스 종료 직후)'),
        ('HI-CMS 활용률 (선주 노출)',      '0% (사내 한정)',              '100% (이상 감지 즉시 알림)'),
        ('LTSA 매출 비중',                 '5~10%',                     '50%↑ (2030)'),
        ('매출 예측 정확도',                '±25% (단발성 의존)',          '±5% (OEM 3사 수준)'),
        ('선주 이탈 리스크',                '측정 불가',                   '갱신율·이탈율 실측 가능'),
    ])


def add_conclusion(doc):
    doc.add_heading('결론 — 4가지 메커니즘, 하나의 흐름', level=1)
    doc.add_paragraph(
        '본 제안은 회사의 두 가지 주제 — 2030년 매출 증대 중장기 계획, 업무 프로세스 개선 — '
        '을 단일 메커니즘으로 통합한다. 그 메커니즘은 다음 4가지 축의 연쇄 흐름이다.')
    add_table(doc, ['축', '역할', '한 줄 설명'], [
        ['① 무상 기술지원 포털',  '입구',    '선주가 자연스럽게 들어오는 단일 창구'],
        ['② AI 케이스 인계',      '통로 1',  'PM이 메일을 뒤지지 않고 즉시 LTSA 제안 가능'],
        ['③ HI-CMS 통합 알림',    '통로 2',  '선주가 안 와도 됐던 영역을 "와야 하는 영역"으로'],
        ['④ LTSA 자연 전환',     '출구',    '단발성 매출 → 안정적 구독+사건 하이브리드'],
    ])
    doc.add_paragraph(
        '이 4가지가 동시에 작동할 때 "매출 증대(중장기 계획)"와 "업무 프로세스 개선" 이 '
        '하나의 흐름에서 동시에 달성된다. 별도 영업 캠페인·신규 시장 진출·신규 인력 일괄 '
        '투입은 필요 없다. 우리가 이미 가지고 있는 데이터와 시스템(HI-CMS, 무상 보증 선박 9,890척, '
        'AS 운영 조직)을 "연결"만 하면 되는 구조다.')

    doc.add_heading('단 한 줄의 약속', level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('"메일을 열던 그 손이, 자연스럽게 LTSA 제안을 클릭하는 순간을 만든다."')
    r.font.size = Pt(13); r.font.bold = True; r.font.italic = True
    r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)


def add_appendix(doc):
    doc.add_heading('부록 A — 발표 시간 배분 (총 10분)', level=1)
    add_table(doc, ['시간', '구간', '핵심 메시지'], [
        ['0:00~1:00',   '서론',                 '회사 주제 두 가지를 하나의 메커니즘으로 통합'],
        ['1:00~3:00',   '본론 1 — 현황·단절',     'PM 1척당 3~8일 / LTSA 비중 5~10%'],
        ['3:00~4:30',   '본론 2 — LTSA 모델',     '하이브리드 (구독+사건) 권고'],
        ['4:30~6:30',   '본론 3 — AI 인계',       '99% 시간 단축 / 영업 기회 손실 5%↓'],
        ['6:30~7:30',   '본론 4 — HI-CMS 알림',   '선주를 "와야 하는 사람"으로 전환'],
        ['7:30~9:00',   '본론 5 — 2030 로드맵',   'LTSA 비중 5% → 50% / 5-Phase'],
        ['9:00~10:00',  '결론 + Q&A',           '4축 메커니즘 / 단 한 줄의 약속'],
    ])

    doc.add_heading('부록 B — 핵심 질문 1줄 답변', level=1)
    add_table(doc, ['#', '질문', '1줄 답변'], [
        ['Q1', '왜 LTSA인가',
         'OEM 3사가 매출 안정성을 확보한 유일한 구조. HMS도 동일하게 도달 가능.'],
        ['Q2', '무상-유상 단절은 어떻게',
         '무상 포털 자체가 유상 영업 도구. AI 인계 + HI-CMS 알림이 그 사이를 잇는다.'],
        ['Q3', 'AI 인계의 효과',
         'PM 1척당 3~8일 → 15~30분 (99% 시간 단축, 처리 능력 5~7배).'],
        ['Q4', 'HI-CMS는 어떻게',
         '이상 감지 → 웹 알림 → 케이스 자동 생성 → LTSA 제안 (선주 클릭 1회).'],
        ['Q5', '2030년 도달 방법',
         '5-Phase 단계별 게이트. PoC(2026) → 정착(2027) → 확장(2028) → 통합(2029) → 안정(2030).'],
    ])

    doc.add_heading('부록 C — 예상 Q&A', level=1)
    add_table(doc, ['Q', '예상 질문', '준비된 답변'], [
        ['1', 'AI 요약 신뢰도가 부족하면?',
         'PM 검토 단계 5~15분 유지. 신뢰도 90%↑ 도달 시까지 PoC 단계에서 검증.'],
        ['2', 'HI-CMS 데이터 보안·선주 동의는?',
         '선주별 옵트인 정책. 공개 범위는 LCA 표준 약관 기반(OEM 3사 동등 수준).'],
        ['3', '구독형 LTSA가 선주 부담만 키우지 않을까?',
         '하이브리드(기본료+사용량) 권고. 저사용 선주에게도 손해 없는 구조 설계.'],
        ['4', '무상 → 유상 전환 시 선주 저항은?',
         '무상은 그대로 유지. 유상은 "추가" 가치(예측정비·우선 응답)에 한정. 강요 없음.'],
        ['5', '필요 인력·예산 규모?',
         'Phase 1 기준 AI/플랫폼 PoC 팀 5~7명. 신규 영업 인력 불필요. (선행 제안 인력 활용)'],
        ['6', 'OEM 3사가 동일하게 막아오면?',
         'OEM 3사는 자기 장비만 커버. HMS는 HD현대 조선소 건조선 전체 시스템 통합 가능 — 차별점.'],
    ])


def add_references(doc):
    doc.add_heading('참고문헌 (References)', level=1)
    doc.add_paragraph(
        '본 제안은 선행 제안의 참고문헌을 그대로 승계하며, LTSA·LCA·AI 인계·예측정비 영역에 '
        '한해 추가 자료를 보강.')

    sections = [
        ('1. 사내 자료 (선행·연계 제안)', [
            ('[1] HD현대마린솔루션테크 — 24/7 무상 기술지원 웹사이트 구축 제안 (2026.04)',
             '본 제안의 직접 전제 · 무상 웹 진입로 설계'),
            ('[2] HD현대마린솔루션테크 — 원격 기술 지원 24/7 전문가 대기 체계 제안 (2026.04)',
             '4-Tier · Follow-the-Sun 운영 체계'),
            ('[3] HD현대마린솔루션테크 — 선도기업 AM 전략 벤치마킹 보고서 (2026.04)',
             'OEM 3사 LCA 매출 구조 분석'),
            ('[4] HD현대마린솔루션테크 — 배전반 서비스 리포트 고도화 제안 (2026.04)',
             '선주 접점 재설계·HI-CMS 데이터 활용'),
        ]),
        ('2. LCA · LTSA 운영 사례 (OEM 선도기업)', [
            ('[1] Wärtsilä — Lifecycle Agreement 표준 약관 및 사례',
             'https://www.wartsila.com/marine/services/lifecycle-agreements'),
            ('[2] Wärtsilä — Expert Insight Service (예측정비 LCA 연계)',
             'https://www.wartsila.com/marine/services/lifecycle-agreements/expert-insight'),
            ('[3] ABB — Service Agreements (Marine)',
             'https://new.abb.com/marine/services/service-agreements'),
            ('[4] Kongsberg Maritime — Lifecycle Services',
             'https://www.kongsberg.com/maritime/services/'),
        ]),
        ('3. AI 기반 운영 인계 · 케이스 자동 요약', [
            ('[1] McKinsey — Generative AI in Customer Service Operations (2024)',
             '케이스 요약·인계 자동화 효과 분석'),
            ('[2] BCG — AI for Field Service Operations (2024)',
             '현장 엔지니어 → PM 인계 효율화 사례'),
            ('[3] Gartner — Hyperautomation in B2B Service (2024)',
             'RAG·LLM 기반 인계 워크플로 트렌드'),
        ]),
        ('4. 매출 안정화 · 구독 모델 이론', [
            ('[1] HBR — The Economics of Subscription Businesses (재게시 분석)',
             '구독+사건 하이브리드 효율 비교'),
            ('[2] BCG — Service Models in Capital-Intensive Industries',
             'LCA 구조의 매출 안정화 메커니즘'),
            ('[3] Deloitte — Servitization in Industrial OEMs',
             '제품 매출 → 서비스 매출 전환 로드맵'),
        ]),
        ('5. 시장 · 산업 동향', [
            ('[1] DNV — Maritime Forecast to 2050 (원격·디지털화 전망)',
             'https://www.dnv.com/maritime/publications/maritime-forecast-2023/'),
            ('[2] Clarksons Research — World Fleet Register',
             '대상 선박 모집단(9,890척) 산정 근거'),
            ('[3] Fortune Business Insights — Maritime Satellite Communication Market',
             '선박 ↔ 본사 실시간 연결성 전망'),
        ]),
    ]

    for title, items in sections:
        h = doc.add_paragraph()
        rh = h.add_run(title)
        rh.font.bold = True; rh.font.size = Pt(11)
        rh.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        for item, sub in items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            r1 = p.add_run(item + '\n')
            r1.font.size = Pt(10); r1.font.bold = True
            r2 = p.add_run('    ' + sub)
            r2.font.size = Pt(9); r2.font.italic = True
            r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        spacer(doc, 6)

    doc.add_paragraph(
        '※ 모든 외부 자료 URL은 발표 시점 기준 최신 공개본. 사내 자료는 발표 자료집 또는 '
        '팀 공유 드라이브 참조.'
    ).runs[0].font.italic = True


def main():
    doc = Document()
    init_doc(doc)
    title_block(doc)
    add_executive_summary(doc)
    add_intro(doc)
    add_section1(doc)
    add_section2(doc)
    add_section3(doc)
    add_section4(doc)
    add_section5(doc)
    add_section6(doc)
    add_conclusion(doc)
    add_appendix(doc)
    add_references(doc)
    doc.save(OUTPUT)
    print(f"✓ saved {OUTPUT}")


if __name__ == "__main__":
    main()
