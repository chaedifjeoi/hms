#!/usr/bin/env python3
"""
원격 기술 지원 — 24/7 전문가 대기 체계 벤치마킹 발표 보고서
서론-본론-결론 구조 / 10분 발표용
"""

import os, re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

# ── 경로 ──
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, "report_images")
os.makedirs(IMG_DIR, exist_ok=True)
OUTPUT = os.path.join(BASE_DIR, "원격 기술 지원 — 24_7 전문가 대기 체계 제안.docx")

# ── 컬러 (네이비/블루 기조) ──
PRIMARY = '1B3A5C'
SECONDARY = '2B579A'
LIGHT = 'F0F4F8'
LIGHT_BG = 'E8F0FE'
RED = 'E74C3C'
YELLOW = 'F39C12'
GREEN = '27AE60'

plt.rcParams['font.family'] = 'AppleGothic'
plt.rcParams['axes.unicode_minus'] = False
C_PRIMARY = '#1B3A5C'
C_SECONDARY = '#2B579A'
C_GRAY = '#BDC3C7'


# ══════════════════════════════════════
#  헬퍼 함수
# ══════════════════════════════════════

def init_doc(doc):
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)

    for level, (size, color) in enumerate([
        (20, PRIMARY), (16, SECONDARY), (13, '333333'),
    ], 1):
        h = doc.styles[f'Heading {level}']
        h.font.name = '맑은 고딕'
        h.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        h.font.size = Pt(size)
        h.font.bold = True
        r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
        h.font.color.rgb = RGBColor(r, g, b)
        h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def set_shading(cell, color):
    tcPr = cell._element.get_or_add_tcPr()
    s = tcPr.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    tcPr.append(s)


def set_text(cell, text, bold=False, color_rgb=None, size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(str(text))
    run.font.size = size
    run.font.bold = bold
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color_rgb:
        run.font.color.rgb = color_rgb


def keep_together(table):
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trPr.append(trPr.makeelement(qn('w:cantSplit'), {}))
    for row in table.rows[:-1]:
        for cell in row.cells:
            for p in cell.paragraphs:
                pPr = p._p.get_or_add_pPr()
                pPr.append(pPr.makeelement(qn('w:keepNext'), {}))
                pPr.append(pPr.makeelement(qn('w:keepLines'), {}))
    for cell in table.rows[-1].cells:
        for p in cell.paragraphs:
            pPr = p._p.get_or_add_pPr()
            pPr.append(pPr.makeelement(qn('w:keepLines'), {}))


def spacer(doc, pt=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(pt)
    run = p.add_run('')
    run.font.size = Pt(1)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_shading(cell, PRIMARY)
        set_text(cell, h, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                set_shading(cell, LIGHT)
            if isinstance(val, str) and val.startswith('**') and val.endswith('**'):
                set_text(cell, val[2:-2], bold=True)
            else:
                set_text(cell, val)
    keep_together(table)
    spacer(doc, 12)
    return table


def add_gap_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_shading(cell, PRIMARY)
        set_text(cell, h, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                set_shading(cell, LIGHT)
            if ci == len(headers) - 1:
                if '큼' in str(val):
                    set_shading(cell, 'FDEDEC')
                    set_text(cell, val, bold=True, color_rgb=RGBColor(0xE7, 0x4C, 0x3C))
                elif '중간' in str(val):
                    set_shading(cell, 'FEF9E7')
                    set_text(cell, val, bold=True, color_rgb=RGBColor(0xF3, 0x9C, 0x12))
                elif '양호' in str(val) or '우위' in str(val):
                    set_shading(cell, 'EAFAF1')
                    set_text(cell, val, bold=True, color_rgb=RGBColor(0x27, 0xAE, 0x60))
                else:
                    set_text(cell, val)
            else:
                set_text(cell, val)
    keep_together(table)
    spacer(doc, 12)
    return table


def add_ba_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['항목', 'Before (현재)', 'After (24/7 전문가 대기)']):
        cell = table.rows[0].cells[i]
        set_shading(cell, PRIMARY)
        set_text(cell, h, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, (item, before, after) in enumerate(rows):
        set_text(table.rows[ri + 1].cells[0], item, bold=True)
        c_b = table.rows[ri + 1].cells[1]
        set_shading(c_b, 'F5F5F5')
        set_text(c_b, before, color_rgb=RGBColor(0x99, 0x99, 0x99))
        c_a = table.rows[ri + 1].cells[2]
        set_shading(c_a, LIGHT_BG)
        set_text(c_a, after, bold=True, color_rgb=RGBColor(0x1B, 0x3A, 0x5C))
    keep_together(table)
    spacer(doc, 12)
    return table


def insight(doc, text, bg='E8F0FE', border='1B3A5C'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_shading(cell, bg)
    tcPr = cell._element.get_or_add_tcPr()
    borders = tcPr.makeelement(qn('w:tcBorders'), {})
    for side, w in [('left', '24'), ('top', '4'), ('bottom', '4'), ('right', '4')]:
        b = borders.makeelement(qn(f'w:{side}'), {
            qn('w:val'): 'single', qn('w:sz'): w,
            qn('w:color'): border if side == 'left' else 'D5D5D5',
            qn('w:space'): '0',
        })
        borders.append(b)
    tcPr.append(borders)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    r, g, b = int(border[:2], 16), int(border[2:4], 16), int(border[4:], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    keep_together(table)
    spacer(doc, 12)


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        else:
            run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    p.clear()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        else:
            run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def chart(doc, path, caption, width=Inches(5.8)):
    spacer(doc, 18)
    doc.add_picture(path, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(f'[그림] {caption}')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    spacer(doc, 18)


# ══════════════════════════════════════
#  차트 생성
# ══════════════════════════════════════

def make_247_map():
    """글로벌 24/7 전문가 대기 네트워크 비교"""
    fig, ax = plt.subplots(figsize=(8, 4))

    companies = ['바르질라\n(Smart Support Centre)', 'ABB\n(Integrated Ops Center)',
                 '콩스버그\n(Global Support)', 'HD현대마린솔루션테크\n(SmartCare 부산)']
    centers = [4, 5, 6, 1]
    colors = [C_SECONDARY, C_SECONDARY, C_SECONDARY, C_PRIMARY]

    bars = ax.barh(companies, centers, color=colors, zorder=3, height=0.55)
    for i, (bar, val) in enumerate(zip(bars, centers)):
        ax.text(val + 0.15, bar.get_y() + bar.get_height()/2,
                f'{val}개소', va='center', fontsize=11, fontweight='bold',
                color=C_PRIMARY)

    ax.set_xlim(0, 7.5)
    ax.set_xlabel('글로벌 24/7 지원센터 수', fontsize=10, color='#666666')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='x', alpha=0.3, linestyle='--')
    ax.set_axisbelow(True)
    ax.invert_yaxis()
    plt.title('글로벌 24/7 전문가 대기 네트워크 비교',
              fontsize=13, fontweight='bold', color=C_PRIMARY, pad=15)
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "support_247_network.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def make_followsun_chart():
    """Follow-the-Sun 24시간 교대 개념도"""
    fig, ax = plt.subplots(figsize=(8, 3.5))

    hours = list(range(24))
    asia = [1 if 0 <= h < 8 else 0 for h in hours]
    europe = [1 if 8 <= h < 16 else 0 for h in hours]
    america = [1 if 16 <= h < 24 else 0 for h in hours]

    ax.fill_between(hours, 0, asia, step='post', color='#3498DB', alpha=0.8, label='아시아 센터 (부산/싱가포르)')
    ax.fill_between(hours, 0, europe, step='post', color='#2B579A', alpha=0.8, label='유럽 센터 (Turku/Drammen)')
    ax.fill_between(hours, 0, america, step='post', color='#1B3A5C', alpha=0.8, label='미주 센터 (Houston)')

    ax.set_xticks(range(0, 25, 4))
    ax.set_xticklabels([f'{h:02d}:00' for h in range(0, 25, 4)], fontsize=9)
    ax.set_yticks([])
    ax.set_xlim(0, 24)
    ax.set_ylim(0, 1.2)
    ax.set_xlabel('UTC 시간대', fontsize=10, color='#666666')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.legend(loc='upper center', bbox_to_anchor=(0.5, -0.2), ncol=3, fontsize=9, frameon=False)
    plt.title('Follow-the-Sun 방식 — 지역 센터 간 24시간 교대 대응',
              fontsize=12, fontweight='bold', color=C_PRIMARY, pad=12)
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "followsun.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def make_response_time_chart():
    """대응 속도 비교"""
    fig, ax = plt.subplots(figsize=(7, 3.5))
    categories = ['최초 응답\n(First Response)', '원인 분석\n(Diagnosis)',
                  '해결 완료\n(Resolution)', '엔지니어 파견 시\n(현장)']
    before = [6, 24, 72, 120]
    after = [0.25, 1, 4, 24]

    x = np.arange(len(categories))
    w = 0.35
    ax.bar(x - w/2, before, w, label='Before (영업시간 내 대응)', color=C_GRAY, zorder=3)
    ax.bar(x + w/2, after, w, label='After (24/7 전문가 대기)', color=C_PRIMARY, zorder=3)

    for i, (b, a) in enumerate(zip(before, after)):
        ax.text(x[i] - w/2, b + 2, f'{b}h', ha='center', fontsize=8, color='#666666')
        ax.text(x[i] + w/2, a + 2, f'{a}h', ha='center', fontsize=8,
                color=C_PRIMARY, fontweight='bold')

    ax.set_ylabel('소요 시간 (시간)', fontsize=9, color='#666666')
    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontsize=9, fontweight='bold')
    ax.set_ylim(0, 140)
    ax.legend(fontsize=9, loc='upper left')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_axisbelow(True)
    plt.title('24/7 전문가 대기 도입 시 대응 속도 변화',
              fontsize=13, fontweight='bold', color=C_PRIMARY, pad=15)
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "response_time_247.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


print("차트 생성 중...")
img_network = make_247_map()
img_followsun = make_followsun_chart()
img_response = make_response_time_chart()
print("차트 생성 완료")

# ══════════════════════════════════════
#  문서 생성
# ══════════════════════════════════════

print("Word 문서 생성 중...")
doc = Document()
init_doc(doc)

# ────── 표지 ──────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(160)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
run = p.add_run('━' * 40)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run('선도기업 벤치마킹을 통한\n24/7 전문가 대기 체계 도입 제안')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
run = p.add_run('24/7 Remote Expert Standby — Benchmarking Report')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 40)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.font.size = Pt(10)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(140)

for t in ['2026. 04', 'HD현대마린솔루션테크']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(t)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_page_break()

# ════════════════════════════════════════════
#  서론 — 왜 24/7 전문가 대기인가 (약 2분)
# ════════════════════════════════════════════

doc.add_heading('서론 — 왜 24/7 전문가 대기인가', level=1)

doc.add_heading('선박은 24시간 멈추지 않는다', level=2)

body(doc,
    '선박은 **24시간 365일 운항**. '
    '문제는 한국 영업시간(09~18시)에만 생기지 않음. '
    '야간·주말·공휴일, 그리고 **정반대 시간대의 대양 한복판**에서도 발생')

bullet(doc, '**시간대 불일치** — 선박이 대서양에 있을 때, 한국은 새벽. 전화해도 받을 사람이 없음')
bullet(doc, '**대응 지연의 진짜 비용** — 1시간 지연이 아니라 \"다음 업무일까지 기다리는\" 12시간+')
bullet(doc, '**긴급도 판단 불가** — 문의가 접수돼도 담당자 판단 전까지는 단순 민원과 구분 안 됨')
bullet(doc, '**선주 입장** — \"밤새 엔진 알람이 울렸는데 아무도 받지 않았다\" = 신뢰 붕괴')

spacer(doc, 6)

doc.add_heading('핵심 질문', level=2)

insight(doc,
    '선박이 언제 고장나든, 즉시 전문가와 연결된다면?\n'
    '고객이 전화를 걸었을 때, 그 자리에서 해결이 시작된다면?')

spacer(doc, 6)

doc.add_heading('시장 동향', level=2)

body(doc,
    '글로벌 해양 서비스 시장은 **\"상시 대응(Always-On)\"** 으로 빠르게 전환 중. '
    '해양 위성통신 시장 **45억$(2025) → 118억$(2034), CAGR 11.4%** — '
    '위성통신 인프라 확장이 24/7 원격 지원의 기술적 전제를 마련')

body(doc,
    '선도 OEM(바르질라, 콩스버그, ABB)은 이미 **전 세계 다중 거점**으로 '
    '24/7 전문가 대기 체계를 운영. 이는 단순 \"야간 당직\"이 아닌 '
    '**Follow-the-Sun(해를 따라가는) 글로벌 교대 체계**')

insight(doc,
    '24/7 전문가 대기 = 고객이 언제 연락해도 "살아있는 전문가"가 답하는 것.\n'
    '이것이 고객을 우리 서비스에 묶어두는 가장 기본적인 약속')

doc.add_page_break()

# ════════════════════════════════════════════
#  본론 1 — 선도기업 24/7 체계 벤치마킹 (약 3.5분)
# ════════════════════════════════════════════

doc.add_heading('본론 1 — 선도기업 24/7 전문가 대기 체계', level=1)

# ── 바르질라 ──
doc.add_heading('바르질라 (Wärtsilä) — Smart Support Centres', level=2)

body(doc,
    '바르질라는 **전 세계 4개 거점**에서 Smart Support Centre를 운영. '
    '각 센터는 **서로 다른 시간대**에 위치하여, 한 센터가 퇴근할 때 다음 센터가 이어받는 '
    '**Follow-the-Sun 방식**의 24/7 대응 체계를 구현')

add_table(doc,
    ['센터', '위치', '시간대', '역할'],
    [
        ['**Turku Centre**', '핀란드 Turku', 'UTC+2', '본사 — 제품 엔지니어링 직결'],
        ['**Singapore Centre**', '싱가포르', 'UTC+8', '아시아·태평양 선박 대응'],
        ['**Houston Centre**', '미국 Houston', 'UTC-6', '미주 선박 대응'],
        ['**Drammen Centre**', '노르웨이 Drammen', 'UTC+1', '북유럽·대서양 대응'],
    ])

body(doc,
    '**핵심 특징:** 어느 시간에 고객이 전화해도 \"지금 근무 중인\" 엔지니어가 받음. '
    '야간 당직이 아니라 그 시간대가 정규 근무인 엔지니어가 응답하므로 '
    '**응답 품질과 속도가 균일**')

bullet(doc, '**Expert Insight** 와 연동 — 데이터 이상 징후 감지 즉시 해당 시간대 센터로 자동 라우팅')
bullet(doc, '**다국어 지원** — 영어/독일어/스페인어/중국어 등 지역별 언어 커버')
bullet(doc, '**케이스 인수인계** — 교대 시 진행 중 케이스를 다음 센터로 디지털 인수')

# ── ABB ──
doc.add_heading('ABB Marine — Integrated Operations Centers (IOCs)', level=2)

body(doc,
    'ABB는 **IOC(통합 운영 센터)** 를 중심으로 24/7 원격 지원을 운영. '
    '**전 세계 600척 이상의 선박을 24시간 모니터링**하며, '
    '이상 징후 발생 즉시 전문가가 개입')

add_table(doc,
    ['항목', '상세'],
    [
        ['**운영 규모**', '전 세계 5개 IOC — 노르웨이·미국·싱가포르·이탈리아·그리스'],
        ['**모니터링 대상**', '600척+ 선박 실시간 모니터링 (24/7)'],
        ['**최초 응답 시간**', '평균 15분 이내 (시간대 무관)'],
        ['**전문가 풀**', '각 IOC에 도메인별 엔지니어 상주 (전기·추진·제어·자동화)'],
        ['**에스컬레이션**', '복잡 건은 R&D 본사로 24시간 내 전달 — 다중 전문가 협업'],
    ])

body(doc,
    '**핵심 효과:** 엔지니어 방문 **70% 감소**, 정비 비용 **50% 절감** — '
    '이는 원격 기술 자체의 효과가 아니라 \"**상시 대기하는 전문가가 즉시 개입**\" 해서 '
    '**작은 문제를 초기에 해결**하기 때문')

# ── 콩스버그 ──
doc.add_heading('콩스버그 (Kongsberg) — Global 24/7 Support Network', level=2)

body(doc,
    '콩스버그는 **세계 최대 규모의 원격 지원 네트워크** 보유. '
    '전 세계 **6개 지원 허브**와 **수백 명의 엔지니어 풀**로 '
    '어떤 선박, 어떤 시간대도 커버')

add_table(doc,
    ['기능', '상세'],
    [
        ['**24/7 Helpdesk**', '전화/이메일/포털 다채널, 최초 응답 15분 이내 SLA'],
        ['**Follow-the-Sun 교대**', '아시아 → 유럽 → 미주 순 시간대 교대, 케이스 연속성 보장'],
        ['**도메인 전문가 풀**', '자동화·DP·에너지관리 등 도메인별 엔지니어 분류'],
        ['**위성통신 허브**', '대양 항해 중에도 즉시 연결 — 통신 두절 없음'],
        ['**원격 시스템 접속**', '필요 시 선박 제어시스템에 직접 원격 로그인'],
    ])

chart(doc, img_followsun, 'Follow-the-Sun 방식 — 한 센터가 퇴근할 때 다음 센터가 이어받는 24시간 교대 체계')

doc.add_page_break()

# ── 벤치마킹 핵심 요약 ──
doc.add_heading('벤치마킹 핵심 요약', level=2)

body(doc,
    '3사 공통점 = **\"어느 시간에 전화해도 전문가가 받는다\"** 는 단순한 약속. '
    '그 약속을 지키기 위한 장치가 **복수 시간대 거점 + 도메인 전문가 풀 + 케이스 인수인계 체계**')

chart(doc, img_network, '글로벌 24/7 지원센터 수 비교 — 선도기업 vs 우리')

insight(doc,
    '"원격 지원"의 본질은 기술이 아니라 "사람"\n'
    '언제든 전화 받을 "근무 중인" 전문가가 있어야 24/7이 성립')

doc.add_page_break()

# ════════════════════════════════════════════
#  본론 2 — 우리 현황과 Gap (약 1.5분)
# ════════════════════════════════════════════

doc.add_heading('본론 2 — HD현대마린솔루션테크 현황과 Gap', level=1)

doc.add_heading('보유 역량', level=2)

body(doc,
    '24/7 전문가 대기 체계의 **기반 요소는 이미 존재**')

bullet(doc, '**Hi-4S** — 24/7 스마트케어 시스템, 선박 데이터 실시간 모니터링')
bullet(doc, '**SmartCare 부산 센터** — 2024.05 출범, 24시간 운영, 동남권 선박 관리')
bullet(doc, '**디지털 관제센터** — 350척+ 모니터링, 150명+ 전문가 협업')
bullet(doc, '**9,890척** AS 네트워크 — HD현대 건조 선박의 유일한 AS 사업자')

spacer(doc, 6)

doc.add_heading('Gap 분석 — 24/7 대응 체계 관점', level=2)

add_gap_table(doc,
    ['영역', '선도기업', 'HD현대마린솔루션테크', 'Gap'],
    [
        ['글로벌 거점 수', '4~6개소 (다중 시간대)', '부산 1개소', '큼'],
        ['Follow-the-Sun 교대', '시간대별 정규 근무로 커버', '야간 당직 체제', '큼'],
        ['최초 응답 SLA', '15분 이내 보장', 'SLA 문서화 미흡', '중간'],
        ['도메인 전문가 풀', '도메인별 상주 엔지니어', '담당자 개인 의존', '큼'],
        ['케이스 인수인계', '디지털 시스템으로 교대 시 이관', '구두/메일 의존', '큼'],
        ['다국어 대응', '지역별 현지어 커버', '한국어·영어 중심', '중간'],
        ['관제 인프라', '전용 IOC + 통합 플랫폼', 'SmartCare 운영 중', '양호'],
    ])

insight(doc,
    '관제 "센터"는 있지만, 글로벌 24/7을 지탱할 "시간대 거점 + 전문가 풀" 체계 부재\n'
    '부산 센터 하나로 전 세계 선박의 24시간을 커버하는 건 구조적으로 불가능',
    bg='FDEDEC', border=RED)

doc.add_page_break()

# ════════════════════════════════════════════
#  본론 3 — 도입 제안 (약 2분)
# ════════════════════════════════════════════

doc.add_heading('본론 3 — 24/7 전문가 대기 체계 도입 제안', level=1)

doc.add_heading('3단계 도입 로드맵', level=2)

body(doc,
    '기존 SmartCare 부산을 **아시아·태평양 허브**로 확장하고, '
    '이후 유럽·미주로 거점을 확대하는 단계적 접근')

add_table(doc,
    ['단계', '시점', '내용'],
    [
        ['**1단계: 내부 체계화**', '6개월', 'SmartCare 부산 내 24/7 교대·SLA·케이스 관리 시스템 정립'],
        ['**2단계: 아시아 허브 확장**', '1년', '싱가포르·중동 거점 추가 (파트너사 협력 또는 소규모 지사)'],
        ['**3단계: 글로벌 Follow-the-Sun**', '2~3년', '유럽·미주 거점 추가로 완전한 24시간 교대 체계 완성'],
    ])

doc.add_heading('1단계 상세 — 내부 체계화 (Quick Win)', level=2)

body(doc,
    '해외 거점 확장 전에도 **내부 체계화만으로** 24/7 대응 수준을 대폭 개선 가능')

bullet(doc, '**3교대 근무제** — SmartCare 부산에 주간/야간/심야 교대 엔지니어 배치')
bullet(doc, '**SLA 문서화** — 최초 응답 15분, 1차 진단 2시간, 해결 완료 24시간 등 등급별 SLA 명시')
bullet(doc, '**전문가 풀 명단** — 도메인(전기·추진·제어·자동화)별 전문가 풀 구성, 호출 프로토콜 수립')
bullet(doc, '**케이스 관리 시스템** — 교대 시 진행 중 케이스를 다음 근무자에게 디지털 인수')
bullet(doc, '**고객 포털** — 선주가 직접 케이스 등록·진행 조회 가능한 웹 포털')

doc.add_heading('배전반 AS 적용 시나리오', level=2)

body(doc, '**시나리오: 새벽 03:00, 대서양 항해 중인 컨테이너선에서 배전반 경보 발생**')

spacer(doc, 4)

add_table(doc,
    ['단계', '기존 방식', '24/7 전문가 대기 도입 후'],
    [
        ['1. 경보 접수', '메일 접수 → 익일 아침 확인', '24/7 포털 자동 접수 → 심야 근무자 15분 내 응답'],
        ['2. 1차 진단', '담당자 출근 후 전화 통화', 'Hi-4S 데이터 조회 → 즉시 상황 파악'],
        ['3. 전문가 연결', '담당 엔지니어 섭외', '도메인 풀에서 상주 배전반 전문가 즉시 호출'],
        ['4. 해결 진행', '현장 출장 일정 협의', '원격 조치 + 필요 시 익일 현장 파견 동시 진행'],
        ['5. 소요 시간', '12~24시간 대기 후 시작', '15분 내 시작, 2~4시간 내 1차 해결'],
    ])

chart(doc, img_response, '24/7 전문가 대기 도입 시 대응 속도 변화')

doc.add_heading('기대 효과', level=2)

add_ba_table(doc, [
    ('최초 응답 시간', '영업시간까지 대기 (최대 12h+)', '15분 이내 (시간대 무관)'),
    ('심야/주말 대응', '불가 또는 당직자 개인 판단', '근무 중인 정규 엔지니어가 응답'),
    ('전문가 배정', '개별 담당자 의존 → 속도 편차', '도메인 풀에서 즉시 배정 → 균일한 품질'),
    ('케이스 누락', '인수인계 누락 발생', '디지털 시스템으로 추적 — 누락 제로'),
    ('고객 경험', '"언제 회신 올까" 불안', '"지금 전문가가 보고 있다" 확신'),
    ('엔지니어 파견', '매 건 현장 방문', '원격 해결 70%, 현장 필요 건만 파견'),
])

doc.add_page_break()

# ════════════════════════════════════════════
#  결론 (약 1분)
# ════════════════════════════════════════════

doc.add_heading('결론', level=1)

spacer(doc, 24)

insight(doc,
    '바르질라는 4개 센터로 24시간을 나눠 지킵니다.\n'
    'ABB는 5개 IOC에서 600척을 24/7 모니터링합니다.\n'
    '콩스버그는 6개 허브로 어느 시간·어느 대양도 커버합니다.\n'
    '\n'
    '우리는 부산 한 곳에서 전 세계를 책임지려 하고 있습니다.')

spacer(doc, 12)

body(doc,
    '그러나 **체계화로 먼저 달려갈 수 있음**')

bullet(doc, 'SmartCare 부산 = **Follow-the-Sun의 아시아 허브**가 될 준비 완료')
bullet(doc, 'Hi-4S + 관제센터 = **상시 모니터링 인프라** 이미 구축')
bullet(doc, '9,890척 AS 네트워크 = **전 세계 선박 접근권** 이미 보유')

spacer(doc, 12)

body(doc,
    '필요한 건 기술이 아니라 **사람과 프로세스** — '
    '3교대 근무, SLA, 도메인 전문가 풀, 케이스 인수인계 시스템. '
    '이것부터 정비하면 **6개월 안에 24/7 대응 수준이 단계 변화**')

spacer(doc, 18)

insight(doc,
    '"언제 전화해도 전문가가 받는다"\n'
    '이 단 한 줄의 약속이 고객을 우리에게 묶어두는 가장 강력한 무기입니다.\n'
    '\n'
    '기술보다 먼저, 약속을 지킬 "체계"를 만드는 것이\n'
    '이 제안의 핵심입니다.')

doc.add_page_break()

# ════════════════════════════════════════════
#  참고 자료
# ════════════════════════════════════════════

doc.add_heading('참고 자료', level=1)

refs = [
    ('Wärtsilä Smart Support Centre',
     'https://www.wartsila.com/marine/products/data-service-downloads/smart-support-centre'),
    ('Wärtsilä Expert Insight Service',
     'https://www.wartsila.com/marine/services/lifecycle-agreements/expert-insight-service'),
    ('Wärtsilä Lifecycle Agreement — Japanese Ferry Operator (2025)',
     'https://www.wartsila.com/are/media/news/04-11-2025-wartsila-lifecyle-agreement-selected-by-japanese-ferry-operator-to-support-service-reliability-3678494'),
    ('ABB Ability™ Remote Diagnostics and Predictive Maintenance',
     'https://new.abb.com/abb-ability/transport/marine/remote-diagnostics-and-predictive-maintenance'),
    ('ABB — Integrated Operations Centers (IOC) Overview',
     'https://new.abb.com/marine/systems-and-solutions/digital-solutions'),
    ('ABB — Remote Diagnostic Services Brochure',
     'https://library.e.abb.com/public/97e02350b7e6330bc1257c47004b1622/RDS%20Marine_Brochure%202014.pdf'),
    ('Maritime Executive — ABB Latest Update of Remote Diagnostic Services',
     'https://maritime-executive.com/corporate/abbs-latest-update-of-remote-diagnostic-services'),
    ('Kongsberg Maritime — Remote Support (24/7)',
     'https://www.kongsberg.com/maritime/services/kongsberg-remote-services/remote-support/'),
    ('Kongsberg Maritime — KONGSBERG Remote Services (K-IMS)',
     'https://www.kongsberg.com/maritime/products/digital/k_ims_applications/kongsberg-remote-services/'),
    ('Fortune Business Insights — Maritime Satellite Communication Market',
     'https://www.fortunebusinessinsights.com/maritime-satellite-communication-market-113315'),
    ('HD현대마린솔루션 — 디지털 서비스',
     'https://www.hd-marinesolution.com/kor/CMS/Contents/Contents.do?mCode=MN094'),
    ('서울경제 — HD현대마린솔루션 부산 스마트케어 센터 설치',
     'https://www.sedaily.com/NewsView/2D9ACT14IL'),
]

for i, (title, url) in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2

    run_num = p.add_run(f'[{i}]  ')
    run_num.font.size = Pt(9)
    run_num.font.bold = True
    run_num.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run_num.font.name = '맑은 고딕'
    run_num._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    run_title = p.add_run(title)
    run_title.font.size = Pt(9)
    run_title.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run_title.font.name = '맑은 고딕'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    run_url = p.add_run(f'\n      {url}')
    run_url.font.size = Pt(8)
    run_url.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run_url.font.name = '맑은 고딕'
    run_url._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ── 저장 ──
doc.save(OUTPUT)
print(f"\n문서 생성 완료: {OUTPUT}")
