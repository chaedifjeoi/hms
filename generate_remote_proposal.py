#!/usr/bin/env python3
"""원격 기술 지원(Remote Expert Support) 도입 제안 — 워드 생성"""
import os, re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE_DIR, "원격 기술 지원(Remote Expert Support) 도입 제안.docx")

PRIMARY = '1B3A5C'
SECONDARY = '2B579A'
LIGHT = 'F0F4F8'
LIGHT_BG = 'E8F0FE'


def init_doc(doc):
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)
    for level, (size, color) in enumerate([(20, PRIMARY), (16, SECONDARY), (13, '333333')], 1):
        h = doc.styles[f'Heading {level}']
        h.font.name = '맑은 고딕'
        h.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        h.font.size = Pt(size)
        h.font.bold = True
        r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
        h.font.color.rgb = RGBColor(r, g, b)
        h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def set_shading(cell, color):
    tcPr = cell._element.get_or_add_tcPr()
    s = tcPr.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    tcPr.append(s)


def set_text(cell, text, bold=False, color_rgb=None, size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(str(text))
    run.font.size = size
    run.font.bold = bold
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color_rgb:
        run.font.color.rgb = color_rgb


def keep_together(table):
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trPr.append(trPr.makeelement(qn('w:cantSplit'), {}))


def spacer(doc, pt=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(pt)
    p.add_run('').font.size = Pt(1)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_shading(cell, PRIMARY)
        set_text(cell, h, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                set_shading(cell, LIGHT)
            if isinstance(val, str) and val.startswith('**') and val.endswith('**'):
                set_text(cell, val[2:-2], bold=True, color_rgb=RGBColor(0x1B, 0x3A, 0x5C))
            else:
                set_text(cell, val)
    keep_together(table)
    spacer(doc, 12)
    return table


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    for part in re.split(r'(\*\*.*?\*\*)', text):
        run = p.add_run(part[2:-2] if part.startswith('**') else part)
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        if part.startswith('**') and part.endswith('**'):
            run.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        else:
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    p.clear()
    for part in re.split(r'(\*\*.*?\*\*)', text):
        run = p.add_run(part[2:-2] if part.startswith('**') else part)
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        if part.startswith('**') and part.endswith('**'):
            run.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        else:
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def insight(doc, text, bg='E8F0FE', border='1B3A5C'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_shading(cell, bg)
    tcPr = cell._element.get_or_add_tcPr()
    borders = tcPr.makeelement(qn('w:tcBorders'), {})
    for side, w in [('left', '24'), ('top', '4'), ('bottom', '4'), ('right', '4')]:
        b = borders.makeelement(qn(f'w:{side}'), {
            qn('w:val'): 'single', qn('w:sz'): w,
            qn('w:color'): border if side == 'left' else 'D5D5D5',
            qn('w:space'): '0',
        })
        borders.append(b)
    tcPr.append(borders)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for line in text.split('\n'):
        if p.runs:
            p.add_run('\n')
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.bold = True
        r, g, b = int(border[:2], 16), int(border[2:4], 16), int(border[4:], 16)
        run.font.color.rgb = RGBColor(r, g, b)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    keep_together(table)
    spacer(doc, 12)


def divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)


# ═══════════════════════════════════════════════════════════════
doc = Document()
init_doc(doc)

# 표지
divider(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('원격 기술 지원(Remote Expert Support) 도입 제안')
run.font.size = Pt(24); run.font.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('HD현대마린솔루션테크의 서비스 대응 속도와 범위를 혁신하는 방안')
run.font.size = Pt(13); run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
spacer(doc, 20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('2026. 04\nHD현대마린솔루션테크')
run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
divider(doc)

doc.add_page_break()

# ───────── 서론 ─────────
doc.add_heading('서론 — 왜 필요한가', level=1)

doc.add_heading('1. 현재 기술 지원 방식의 한계', level=2)
body(doc, '우리 회사의 현재 기술 지원 프로세스는 **"전화/메일 → 엔지니어 직접 방문"** 구조입니다.')
insight(doc, '선주 문의 전화 → PM 상황 파악 → 엔지니어 출장 스케줄 → 이동 → 현장 도착 → 진단 → 조치',
        bg='F5F5F5', border='888888')
body(doc, '이 구조의 문제점은 다음과 같습니다.')
add_table(doc, ['문제', '상세'], [
    ['**대응 시간 느림**', '전화 → 엔지니어 배정 → 출장 → 도착까지 최소 수일. 해외라면 수 주'],
    ['**출장 비용 부담**', '항공료, 숙박, 일당. 해외 출장 1건 수백만 원. 무상 헬스체크면 전액 회사 부담'],
    ['**전화 한계**', '"이 부분에서 소리가 나요" → 전화로 정확한 상황 파악 불가. 사진 몇 장으로도 부족'],
    ['**엔지니어 가용성**', '동시에 여러 건 발생 시 인원 부족 → 대기 발생'],
    ['**글로벌 대응 한계**', '선박은 전 세계. 중동·유럽·남미 선박 대응 시 시간+비용 엄청남'],
    ['**간단한 문제도 출장**', '현장 가보니 "스위치 하나 리셋하면 되는 거였다" → 출장비 낭비'],
])
insight(doc, '"엔지니어가 현장에 가지 않고도, 현장의 눈과 귀가 될 수 있다면?"')

doc.add_heading('2. 선도기업은 이미 원격 지원을 하고 있다', level=2)

doc.add_heading('바르질라 — Smart Support Centre + AR 원격 가이드', level=3)
body(doc, '바르질라는 2018년부터 AR(증강현실) 기반 원격 지원을 도입했고, 2020년에 **Smart Support Centre**를 정식 런칭했습니다.')
body(doc, '**작동 방식**')
bullet(doc, '선상에서 문제 발생 → 선원이 스마트 글래스/태블릿/스마트폰 착용')
bullet(doc, '카메라 + 마이크로 현장 상황 실시간 전송')
bullet(doc, '바르질라 전문가가 원격에서 실시간으로 확인')
bullet(doc, '음성 지시 + 화면에 AR 마커로 "이 부분을 확인하세요" 표시')
bullet(doc, '선원이 가이드에 따라 조치 → 엔지니어 출장 없이 해결')

body(doc, '**바르질라 실제 성과**')
add_table(doc, ['지표', '수치', '출처'], [
    ['**원격 해결률**', '95% — 거의 모든 오류를 원격으로 해결', 'Wärtsilä Smart Support Centre'],
    ['**대응 시간**', '즉시 (24/7 전문가 대기)', '—'],
    ['**AR 테스트 결과**', 'WiFi 안테나 + 스마트 글래스로 선상 어디서든 실시간 통신 검증', '2019년 테스트'],
])
insight(doc, '"95%의 오류를 원격으로 해결할 수 있습니다.\n다음 항구에서 서비스 엔지니어를 부를 필요가 없습니다." — 바르질라 공식 발표')

doc.add_heading('콩스버그 — Remote Services (원격 진단/원격 가상 서비스)', level=3)
body(doc, '콩스버그는 여러 단계의 원격 서비스를 제공합니다.')
add_table(doc, ['서비스', '내용'], [
    ['**Remote Diagnostics (1st Line)**', '실시간 선상 시스템 상태 원격 확인 → 즉석 진단 + 트러블슈팅'],
    ['**Remote Virtual Services**', '모바일 기기로 선상 인력과 콩스버그 전문가 실시간 협업 → 트러블슈팅·운영 가이드·원격 검사'],
    ['**Remote Configuration**', '원격 시스템 설정 변경, 진단, 소프트웨어 업데이트'],
])
insight(doc, '"첫 시도에 해결한다(resolve on first attempt).\n그래야 비용과 다운타임을 최소화할 수 있다." — 콩스버그')
body(doc, '모든 통신은 **PKI 기반 암호화**로 보안 처리됩니다.')

doc.add_heading('미 해군 / 캐나다 해군 — AR 정비 시스템 실전 배치', level=3)
body(doc, '군용 사례지만 시사점이 큽니다.')
add_table(doc, ['사례', '내용'], [
    ['**미 해군 ARMS**', 'AR 헤드셋으로 전문가 원격 가이드. 기술 매뉴얼·3D 도면·정비 카드를 AR 표시. 현장 방문 필요성 대폭 감소'],
    ['**캐나다 해군 MIRRAS**', 'HoloLens + AI로 선상 정비/수리. 숙련도 격차 해소 + 원격 지식 전달'],
    ['**정비 속도**', 'AR 원격 지원 시 고장 수리 속도 47% 향상 (기술 매뉴얼 대비)'],
    ['**오류율**', 'AR 원격 지원 시 작업 오류 13% vs 종이 매뉴얼 53%'],
])

doc.add_page_break()

# ───────── 본론 ─────────
doc.add_heading('본론 — 도입 방안과 기대 이득', level=1)

doc.add_heading('1. 우리 회사에 맞는 원격 지원 시스템 설계', level=2)

doc.add_heading('서비스 단계 — 3단계 대응 체계', level=3)

body(doc, '**Level 1. 전화/채팅 원격 지원 (즉시 도입 가능)**')
body(doc, '지금 당장 할 수 있는 것. 추가 투자 거의 없음.')
add_table(doc, ['항목', '내용'], [
    ['**방식**', '전화 + 카카오톡/WhatsApp/Teams 영상통화'],
    ['**절차**', '선원이 스마트폰으로 배전반 촬영 → 우리 엔지니어가 보면서 음성 가이드'],
    ['**도구**', '기존 스마트폰 + 기존 메신저 앱 (별도 투자 없음)'],
    ['**한계**', '화면 공유, AR 마킹, 녹화/기록 기능 없음'],
    ['**적합 상황**', '간단한 문의, 스위치 조작 가이드, 알람 확인, 기본 트러블슈팅'],
])

body(doc, '**Level 2. 화상/화면 공유 원격 지원 (3~6개월 내 도입)**')
body(doc, '전문 원격 지원 플랫폼 도입. 기록과 품질이 올라감.')
add_table(doc, ['항목', '내용'], [
    ['**방식**', '전용 원격 지원 플랫폼 (Microsoft Teams / Zoom + 원격 지원 앱)'],
    ['**절차**', '선원이 앱 실행 → 우리 전문가에게 연결 → 실시간 영상 + 화면 공유'],
    ['**추가 기능**', '화면 캡처/녹화 (기록용), 포인터로 위치 지시, 파일(도면/매뉴얼) 실시간 전송'],
    ['**기록**', '세션 녹화 → 서비스 이력에 자동 첨부 → 웹 포털 연동'],
    ['**적합 상황**', '중간 난이도 트러블슈팅, 측정 가이드, 점검 절차 안내'],
])

body(doc, '**Level 3. AR 원격 가이드 (12개월~ 도입)**')
body(doc, '바르질라 Smart Support Centre 수준. 가장 강력한 원격 지원.')
add_table(doc, ['항목', '내용'], [
    ['**방식**', 'AR 글래스(스마트 글래스) 또는 태블릿 AR 앱'],
    ['**절차**', '선원이 AR 기기 착용 → 우리 전문가와 연결 → 전문가가 선원 시야를 실시간으로 봄'],
    ['**AR 기능**', '전문가가 선원 시야에 화살표, 원, 텍스트를 AR로 표시'],
    ['**핸즈프리**', '선원은 두 손으로 작업하면서 글래스로 가이드 (폰 들 필요 없음)'],
    ['**추가 기능**', '매뉴얼/도면 AR 오버레이, 3D 모델 표시, 세션 전체 녹화'],
    ['**적합 상황**', '복잡한 수리 작업, 부품 교체 가이드, 신입 선원 지원, 긴급 대응'],
])

doc.add_heading('단계별 비교', level=3)
add_table(doc, ['항목', 'Level 1 전화/채팅', 'Level 2 화상/공유', 'Level 3 AR 가이드'], [
    ['**투자**', '0원', '낮음 (플랫폼 라이선스)', '중간 (AR 기기 + 플랫폼)'],
    ['**도입 시점**', '즉시', '3~6개월', '12개월~'],
    ['**해결 범위**', '간단한 문의/가이드', '중간 난이도 트러블슈팅', '복잡한 수리까지'],
    ['**핸즈프리**', '불가 (폰 들어야 함)', '불가 (폰/태블릿)', '가능 (글래스)'],
    ['**AR 마킹**', '불가', '불가', '가능'],
    ['**기록/녹화**', '비공식', '가능', '가능'],
    ['**선원 부담**', '낮음', '낮음', '중간 (기기 착용)'],
])

doc.add_heading('2. 구체적 운영 방식', level=2)
doc.add_heading('원격 지원 요청 프로세스', level=3)
bullet(doc, '① 웹 포털 또는 앱에서 "원격 지원 요청" 클릭 (긴급도 선택: 일반 / 긴급)')
bullet(doc, '② 우리 엔지니어에게 자동 배정 + 알림 (긴급이면 당직 엔지니어 즉시 알림)')
bullet(doc, '③ 엔지니어가 원격 세션 시작 (Level 1: 전화/영상통화 · Level 2: 화상 공유 · Level 3: AR 세션)')
bullet(doc, '④ 실시간 진단 + 가이드 (선원 카메라로 현장 확인, 엔지니어가 음성/화면 지시, 도면/매뉴얼 전송)')
bullet(doc, '⑤ 결과 기록 (세션 녹화/스크린샷, 진단·조치 기록, 웹 포털 서비스 이력 자동 등록)')
bullet(doc, '⑥ 후속 조치 판단 (원격 해결 완료 → 종료 / 현장 방문 필요 → 정확한 상황 파악된 상태로 출장)')

doc.add_heading('운영 체계', level=3)
add_table(doc, ['항목', '내용'], [
    ['**운영 시간**', '초기: 업무 시간 (09:00~18:00 KST) · 확대 후: 24/7'],
    ['**대응 인원**', '배전반 전문 엔지니어 2~3명 교대 (기존 인력 활용)'],
    ['**대응 시간 목표**', '일반: 4시간 이내 세션 시작 / 긴급: 1시간 이내'],
    ['**지원 언어**', '한국어, 영어'],
    ['**보안**', '영상 통화 암호화, 녹화 데이터 사내 서버 저장'],
])

doc.add_heading('3. 도입 시 우리 회사에 돌아오는 이득', level=2)

doc.add_heading('이득 1. 출장 비용 대폭 절감', level=3)
add_table(doc, ['항목', '현재 (모든 건 출장)', '원격 지원 도입 후'], [
    ['**간단한 문의**', '엔지니어 출장 (수백만 원)', '원격 10분에 해결 (0원)'],
    ['**중간 트러블슈팅**', '출장 → 현장 진단 → 조치', '원격 진단 → 70% 원격 해결, 30% 출장'],
    ['**복잡한 수리**', '출장 (가봐야 뭔지 앎)', '원격으로 사전 진단 완료 → 정확한 준비 후 출장'],
])
insight(doc, '비용 절감 시뮬레이션 (연간 출장 100건 가정)\n'
             '현재: 100건 × ₩3M = ₩300M\n'
             '도입 후: 원격 해결 50건(₩0) + 사전진단 출장 30건(₩90M) + 기존 출장 20건(₩60M) = ₩150M\n'
             '→ 연간 절감액 ₩150M (50% 절감)')

doc.add_heading('이득 2. 대응 속도 혁신', level=3)
add_table(doc, ['구분', '현재', '원격 지원 도입 후'], [
    ['국내 선박', '1~3일 (출장 준비 + 이동)', '**수 시간 이내 (원격 세션 즉시)**'],
    ['해외 선박', '1~2주 (비자·항공·이동)', '**수 시간 이내 (원격)**'],
    ['긴급 상황', '"최대한 빨리 가겠습니다"', '**1시간 이내 원격 대응 시작**'],
])
body(doc, '**핵심:** 선박은 전 세계에 있지만, 원격 지원은 거리와 무관하게 동일한 속도입니다.')

doc.add_heading('이득 3. 엔지니어 생산성 향상', level=3)
add_table(doc, ['항목', '현재', '원격 지원 도입 후'], [
    ['엔지니어 1인 주간 처리', '출장 1~2건 (이동 시간 포함)', '원격 5~10건 + 필요 시 출장 1건'],
    ['이동 시간', '전체 업무의 30~50%', '0% (원격 건)'],
    ['동시 대응', '불가 (1명이 1곳에)', '순차 대응 가능 (원격이라 바로 다음 건)'],
])
body(doc, '**같은 인원으로 3~5배 더 많은 건을 처리할 수 있습니다.**')

doc.add_heading('이득 4. 헛걸음 방지 + 정확한 출장 준비', level=3)
add_table(doc, ['항목', '현재', '원격 지원 도입 후'], [
    ['출장 전 상황 파악', '전화로 대충 파악 → 가봐야 앎', '영상으로 정확히 파악 완료'],
    ['부품/공구 준비', '추정으로 준비 → 현장에서 부족', '원격 진단 후 정확한 준비 목록'],
    ['헛걸음', '"스위치 리셋이면 됐는데..."', '원격 선확인 → 갈 필요 없는 건 안 감'],
])

doc.add_heading('이득 5. 서비스 기록 체계화', level=3)
add_table(doc, ['항목', '현재', '원격 지원 도입 후'], [
    ['기술 지원 기록', '전화 통화 → 기록 없음', '세션 녹화 + 스크린샷 + 조치 기록 자동 저장'],
    ['인수인계', '"이 선박 뭐가 문제였더라..."', '과거 원격 세션 녹화본 열람 → 즉시 파악'],
    ['노하우 축적', '엔지니어 머릿속', '원격 세션이 교육 자료가 됨 (신입 학습용)'],
])

doc.add_heading('이득 6. 고객 경험 혁신 — 경쟁사 차별화', level=3)
add_table(doc, ['항목', '현재 (경쟁사도 동일)', '원격 지원 도입 후'], [
    ['고객 경험', '"전화했더니 출장 잡는다고 기다리래"', '"바로 연결되서 화면 보면서 해결해줬다"'],
    ['차별화', '없음 (다 비슷)', '"이 회사는 원격으로 바로 봐준다"'],
    ['경쟁 우위', '가격 경쟁', '서비스 속도/품질 경쟁'],
])

doc.add_heading('4. 구체적 도입 프로세스', level=2)

doc.add_heading('Step 1. Level 1 즉시 도입 (1~2개월)', level=3)
add_table(doc, ['활동', '상세'], [
    ['**내부 준비**', '원격 지원 전담 엔지니어 2~3명 지정 (기존 인력)'],
    ['**도구**', '기존 스마트폰 + 영상통화 앱 (Teams/WhatsApp)'],
    ['**프로세스 수립**', '원격 지원 요청 → 대응 → 기록 프로세스 표준화'],
    ['**고객 안내**', '주요 고객에게 "원격 기술 지원 서비스 시작" 안내'],
    ['**기록**', '원격 세션 후 간단 보고서 작성 → 서비스 이력 등록'],
    ['**비용**', '₩0 (추가 투자 없음)'],
])

doc.add_heading('Step 2. Level 2 화상 플랫폼 도입 (3~6개월)', level=3)
add_table(doc, ['활동', '상세'], [
    ['**플랫폼 선정**', '원격 지원 전용 플랫폼 도입 (VSight, TeamViewer, Microsoft Remote Assist 등)'],
    ['**기능**', '화면 공유, 포인터, 캡처/녹화, 파일 전송'],
    ['**웹 포털 연동**', '원격 세션 기록이 웹 포털 서비스 이력에 자동 등록'],
    ['**교육**', '엔지니어 + 주요 고객 선원 대상 사용법 교육'],
    ['**파일럿**', '주요 고객 5~10곳 대상 파일럿 운영 + 피드백'],
    ['**비용**', '낮음 (플랫폼 라이선스 월 수십만 원)'],
])

doc.add_heading('Step 3. Level 3 AR 도입 (12개월~)', level=3)
add_table(doc, ['활동', '상세'], [
    ['**AR 기기 선정**', 'RealWear Navigator, Microsoft HoloLens, 또는 태블릿 AR 앱'],
    ['**파일럿**', '프리미엄 고객 대상 AR 기기 배치'],
    ['**콘텐츠 제작**', '배전반 주요 작업 AR 가이드 콘텐츠 (차단기 교체 절차, 측정 방법 등)'],
    ['**확대**', '성과 확인 후 대상 확대'],
    ['**비용**', '중간 (AR 기기 대당 ₩1~3M + 플랫폼 비용)'],
])

doc.add_heading('도입 타임라인', level=3)
bullet(doc, '**[즉시~2개월]** Level 1 — 전화/영상통화 원격 지원 시작 (비용 ₩0)')
bullet(doc, '**[3~6개월]** Level 2 — 화상 플랫폼 도입 + 웹 포털 연동')
bullet(doc, '**[6~12개월]** Level 2 확대 + 성과 분석')
bullet(doc, '**[12개월~]** Level 3 — AR 원격 가이드 파일럿')

doc.add_page_break()

# ───────── 결론 ─────────
doc.add_heading('결론 — 왜 지금 해야 하는가', level=1)

doc.add_heading('1. 선도기업은 이미 운영 중이다', level=2)
bullet(doc, '**바르질라**: Smart Support Centre로 24/7 원격 지원 운영 중. 95%를 원격 해결.')
bullet(doc, '**콩스버그**: Remote Diagnostics + Remote Virtual Services 운영 중. "첫 시도에 해결" 원칙.')
bullet(doc, '**미/캐나다 해군**: AR 기반 원격 정비를 실전 배치 완료. 수리 속도 47% 향상.')
insight(doc, '이건 미래 기술이 아니라, 이미 작동하고 있는 현재 기술입니다.')

doc.add_heading('2. Level 1은 내일 시작할 수 있다', level=2)
bullet(doc, '스마트폰 + 영상통화 = **추가 투자 0원**')
bullet(doc, '기존 엔지니어가 전화로 하던 걸 **영상통화로 바꾸기만** 하면 됨')
bullet(doc, '이것만 해도 간단한 건의 **50% 이상을 출장 없이 해결** 가능')

doc.add_heading('3. 비용 절감이 즉각적이다', level=2)
bullet(doc, '출장 100건 중 50건만 원격 해결해도 → **연간 ₩150M 절감** (가정)')
bullet(doc, '같은 인원으로 **3~5배 더 많은 건 처리** 가능')
bullet(doc, '해외 선박 대응이 **수 주 → 수 시간**으로 단축')

doc.add_heading('4. 웹 포털과의 시너지', level=2)
bullet(doc, '원격 지원 요청을 **웹 포털에서 접수** → 통합 고객 접점')
bullet(doc, '원격 세션 녹화·기록이 **웹 포털 서비스 이력에 자동 축적** → 데이터 자산')
bullet(doc, '고객은 **웹 포털 하나로 요청·추적·원격 상담**까지 처리')

doc.add_heading('5. 한 문장 요약', level=2)
insight(doc,
        '"엔지니어를 보내는 게 아니라, 엔지니어의 눈과 귀를 보낸다.\n'
        '바르질라는 95%를 원격으로 해결하고 있다.\n'
        '우리는 스마트폰 하나로 내일부터 시작할 수 있다.\n'
        '시작하는 데 드는 비용은 0원이다."')

doc.add_page_break()

# ───────── 참고 자료 ─────────
doc.add_heading('참고 자료', level=1)
refs = [
    ('Wärtsilä Smart Support Centre', 'https://www.wartsila.com/marine/products/data-service-downloads/smart-support-centre'),
    ('Wärtsilä AR 원격 가이드 테스트 (2019)', 'https://www.wartsila.com/media/news/15-03-2019-wartsila-successfully-tests-remote-guidance-service-capabilities-2401858'),
    ('Wärtsilä AR 정비 서비스 도입 (2018)', 'https://www.wartsila.com/media/news/31-07-2018-augmented-reality-creates-a-new-dimension-in-marine-maintenance-services'),
    ('Wärtsilä Smart Support Centre 빠른 원격 대응 (2020)', 'https://www.wartsila.com/media/news/16-04-2020-wartsila-s-smart-support-centre-delivers-fast-remote-service-response-2688324'),
    ('Kongsberg Remote Support', 'https://www.kongsberg.com/maritime/support/remote-support/'),
    ('Kongsberg Remote Diagnostics', 'https://www.kongsberg.com/maritime/services/kongsberg-remote-services/remote-diagnostics-1st-line/'),
    ('Kongsberg Remote Virtual Services', 'https://www.kongsberg.com/maritime/services/kongsberg-remote-services/remote-virtual-service/'),
    ('미 해군 ARMS AR 정비 시스템', 'https://seapowermagazine.org/first-augmented-reality-maintenance-systems-operational-on-five-ships/'),
    ('캐나다 해군 MIRRAS AR 정비', 'https://maritime-executive.com/article/canadian-navy-trials-augmented-reality-for-shipboard-maintenance'),
    ('AR 원격 지원 시 수리 속도 47% 향상', 'https://vsight.io/blog/augmented-reality-in-maritime-industry-operations/'),
    ('해양 AR 협업 체계적 문헌 리뷰', 'https://www.tandfonline.com/doi/full/10.1080/10447318.2023.2209838'),
]
for i, (title, url) in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'[{i}]  {title}\n      {url}')
    run.font.size = Pt(9)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.save(OUTPUT)
print(f'✅ 저장 완료: {OUTPUT}')
