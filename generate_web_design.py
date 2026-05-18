#!/usr/bin/env python3
"""
24/7 무상 기술 지원 웹사이트 — 상세 설계서
서론 → 본론(아키텍처·IA·기능·유저플로우·데이터모델·무상→유상 로드맵) → 결론
PM 업무 개선 관점의 기대 효과 포함
"""

import os, re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

# ── 경로 ──
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, "report_images")
os.makedirs(IMG_DIR, exist_ok=True)
OUTPUT = os.path.join(BASE_DIR, "24_7 무상 기술지원 웹사이트 상세 설계.docx")

# ── 네이비/블루 기조 ──
PRIMARY   = '1B3A5C'
SECONDARY = '2B579A'
LIGHT_BG  = 'E8F0FE'
LIGHT     = 'F0F4F8'
RED       = 'E74C3C'
ORANGE    = 'F39C12'

plt.rcParams['font.family'] = 'AppleGothic'
plt.rcParams['axes.unicode_minus'] = False
C_PRIMARY   = '#1B3A5C'
C_SECONDARY = '#2B579A'
C_ACCENT    = '#3498DB'
C_DARK      = '#0F2439'
C_GRAY      = '#BDC3C7'
C_LIGHT     = '#E8F0FE'
C_GREEN     = '#27AE60'
C_RED       = '#E74C3C'
C_ORANGE    = '#F39C12'


# ══════════════════════════════════════
#  헬퍼
# ══════════════════════════════════════

def init_doc(doc):
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    style.paragraph_format.line_spacing = 1.2
    style.paragraph_format.space_after = Pt(4)

    for level, (size, color) in enumerate([
        (20, PRIMARY), (15, PRIMARY), (12, SECONDARY),
    ], 1):
        h = doc.styles[f'Heading {level}']
        h.font.name = '맑은 고딕'
        h.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        h.font.size = Pt(size)
        h.font.bold = True
        r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
        h.font.color.rgb = RGBColor(r, g, b)
        h.paragraph_format.space_before = Pt(16 if level == 1 else 10)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)


def set_shading(cell, color):
    tcPr = cell._element.get_or_add_tcPr()
    s = tcPr.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    tcPr.append(s)


def set_text(cell, text, bold=False, color_rgb=None, size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(str(text))
    run.font.size = size
    run.font.bold = bold
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color_rgb:
        run.font.color.rgb = color_rgb


def keep_together(table):
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trPr.append(trPr.makeelement(qn('w:cantSplit'), {}))
    for row in table.rows[:-1]:
        for cell in row.cells:
            for p in cell.paragraphs:
                pPr = p._p.get_or_add_pPr()
                pPr.append(pPr.makeelement(qn('w:keepNext'), {}))
                pPr.append(pPr.makeelement(qn('w:keepLines'), {}))
    for cell in table.rows[-1].cells:
        for p in cell.paragraphs:
            pPr = p._p.get_or_add_pPr()
            pPr.append(pPr.makeelement(qn('w:keepLines'), {}))


def spacer(doc, pt=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(pt)
    run = p.add_run('')
    run.font.size = Pt(1)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_shading(cell, PRIMARY)
        set_text(cell, h, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF),
                 align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                set_shading(cell, LIGHT)
            if isinstance(val, str) and val.startswith('**') and val.endswith('**'):
                set_text(cell, val[2:-2], bold=True, size=Pt(9))
            else:
                set_text(cell, val, size=Pt(9))
    keep_together(table)
    spacer(doc, 10)
    return table


def add_ba_table(doc, rows, before_label='Before', after_label='After'):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['항목', before_label, after_label]):
        cell = table.rows[0].cells[i]
        set_shading(cell, PRIMARY)
        set_text(cell, h, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF),
                 align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10))
    for ri, (item, before, after) in enumerate(rows):
        set_text(table.rows[ri + 1].cells[0], item, bold=True, size=Pt(9))
        c_b = table.rows[ri + 1].cells[1]
        set_shading(c_b, 'F5F5F5')
        set_text(c_b, before, color_rgb=RGBColor(0x99, 0x99, 0x99), size=Pt(9))
        c_a = table.rows[ri + 1].cells[2]
        set_shading(c_a, LIGHT_BG)
        set_text(c_a, after, bold=True, color_rgb=RGBColor(0x1B, 0x3A, 0x5C), size=Pt(9))
    keep_together(table)
    spacer(doc, 10)
    return table


def insight(doc, text, bg=LIGHT_BG, border=PRIMARY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_shading(cell, bg)
    tcPr = cell._element.get_or_add_tcPr()
    borders = tcPr.makeelement(qn('w:tcBorders'), {})
    for side, w in [('left', '24'), ('top', '4'), ('bottom', '4'), ('right', '4')]:
        b = borders.makeelement(qn(f'w:{side}'), {
            qn('w:val'): 'single', qn('w:sz'): w,
            qn('w:color'): border if side == 'left' else 'D5D5D5',
            qn('w:space'): '0',
        })
        borders.append(b)
    tcPr.append(borders)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    r, g, b = int(border[:2], 16), int(border[2:4], 16), int(border[4:], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    keep_together(table)
    spacer(doc, 10)


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        else:
            run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)


def bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.left_indent = Cm(0.5)

    prefix = p.add_run('▸  ')
    prefix.font.size = Pt(10)
    prefix.font.bold = True
    prefix.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    prefix.font.name = '맑은 고딕'
    prefix._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        else:
            run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)


def sub_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Cm(1.2)
    prefix = p.add_run('·  ')
    prefix.font.size = Pt(10)
    prefix.font.bold = True
    prefix.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    prefix.font.name = '맑은 고딕'
    prefix._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        else:
            run = p.add_run(part)
            run.font.size = Pt(9.5)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def chart(doc, path, caption, width=Inches(5.9)):
    spacer(doc, 10)
    doc.add_picture(path, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(f'[그림] {caption}')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    spacer(doc, 10)


# ══════════════════════════════════════
#  차트 / 다이어그램
# ══════════════════════════════════════

def draw_box(ax, x, y, w, h, text, color, text_color='white', fontsize=9.5, fontweight='bold'):
    box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.03",
                          linewidth=1.2, edgecolor=color, facecolor=color)
    ax.add_patch(box)
    ax.text(x + w/2, y + h/2, text, ha='center', va='center',
            fontsize=fontsize, fontweight=fontweight, color=text_color)


def make_system_arch():
    """시스템 아키텍처 다이어그램"""
    fig, ax = plt.subplots(figsize=(9, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 8)
    ax.axis('off')

    # 외부 사용자 (상단)
    draw_box(ax, 0.3, 6.7, 1.8, 0.9, '선주 PM\n(웹/모바일)',          C_SECONDARY, fontsize=9)
    draw_box(ax, 2.3, 6.7, 1.8, 0.9, '선박 사관\n(모바일/위성)',       C_SECONDARY, fontsize=9)
    draw_box(ax, 4.3, 6.7, 1.8, 0.9, 'HMS 엔지니어\n(부산·해외지사)',  C_PRIMARY, fontsize=9)
    draw_box(ax, 6.3, 6.7, 1.8, 0.9, '도메인 전문가\n(Tier 3·4)',      C_PRIMARY, fontsize=9)
    draw_box(ax, 8.3, 6.7, 1.5, 0.9, '관리자\n(Admin)',                C_DARK, fontsize=9)

    # 프레젠테이션 계층
    draw_box(ax, 0.3, 5.1, 9.5, 1.1, 'Presentation Layer\n웹 포털 (PC) · 모바일 앱 (iOS/Android) · 이메일 게이트웨이', C_ACCENT, fontsize=10)

    # 애플리케이션 계층
    draw_box(ax, 0.3, 3.3, 1.75, 1.5, '인증/SSO\n(OAuth 2.0)',       C_PRIMARY, fontsize=9)
    draw_box(ax, 2.25, 3.3, 1.75, 1.5, '케이스 관리\n(접수·진단·종료)', C_PRIMARY, fontsize=9)
    draw_box(ax, 4.2, 3.3, 1.75, 1.5, '실시간 채팅·Tier\n라우팅 엔진',  C_PRIMARY, fontsize=9)
    draw_box(ax, 6.15, 3.3, 1.75, 1.5, '지식베이스·FAQ\n(AI 매칭)',    C_PRIMARY, fontsize=9)
    draw_box(ax, 8.1, 3.3, 1.7, 1.5, '알림·리포트\n엔진',              C_PRIMARY, fontsize=9)

    # 데이터 계층
    draw_box(ax, 0.3, 1.6, 3.0, 1.2, 'Core DB\n(선박·케이스·유저·메시지·SLA)', C_DARK, fontsize=9.5)
    draw_box(ax, 3.4, 1.6, 3.0, 1.2, 'Knowledge Store\n(FAQ·해결사례·매뉴얼·도면)', C_DARK, fontsize=9.5)
    draw_box(ax, 6.5, 1.6, 3.3, 1.2, 'Event Log / Audit\n(타임라인·핸드오프·SLA 기록)', C_DARK, fontsize=9.5)

    # 외부 연계
    draw_box(ax, 0.3, 0.1, 2.3, 1.1, 'Hi-4S\n(선박 데이터·경보)',     '#7F8C8D', fontsize=9)
    draw_box(ax, 2.8, 0.1, 2.3, 1.1, 'SmartCare 관제\n(모니터링 CMS)', '#7F8C8D', fontsize=9)
    draw_box(ax, 5.3, 0.1, 2.3, 1.1, 'ERP·AS 계약 시스템\n(유상 정산)', '#7F8C8D', fontsize=9)
    draw_box(ax, 7.8, 0.1, 2.0, 1.1, 'Mail Gateway\n(SMTP/IMAP)',       '#7F8C8D', fontsize=9)

    # 레이어 라벨
    ax.text(-0.1, 7.15, 'Users',       fontsize=10, fontweight='bold', color=C_PRIMARY, rotation=90, va='center')
    ax.text(-0.1, 5.65, 'Frontend',    fontsize=10, fontweight='bold', color=C_ACCENT,  rotation=90, va='center')
    ax.text(-0.1, 4.05, 'Application', fontsize=10, fontweight='bold', color=C_PRIMARY, rotation=90, va='center')
    ax.text(-0.1, 2.2,  'Data',        fontsize=10, fontweight='bold', color=C_DARK,    rotation=90, va='center')
    ax.text(-0.1, 0.65, 'External',    fontsize=10, fontweight='bold', color='#7F8C8D', rotation=90, va='center')

    plt.title('24/7 무상 기술 지원 웹사이트 — 시스템 아키텍처 5계층 구조',
              fontsize=13, fontweight='bold', color=C_PRIMARY, pad=12)
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "web_system_arch.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def make_sitemap():
    """사이트 맵"""
    fig, ax = plt.subplots(figsize=(9, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 7)
    ax.axis('off')

    # 루트
    draw_box(ax, 3.8, 6.0, 2.4, 0.7, '홈 / 대시보드', C_PRIMARY, fontsize=10)

    # 1차 메뉴
    menus = [
        (0.2,  4.6, '케이스\n(문의·이력)'),
        (2.1,  4.6, '지식베이스\n(FAQ·매뉴얼)'),
        (4.0,  4.6, '선박별 현황\n(내 선대)'),
        (5.9,  4.6, '리포트\n(월간·SLA)'),
        (7.8,  4.6, '내 계정\n(알림·권한)'),
    ]
    for x, y, t in menus:
        draw_box(ax, x, y, 1.8, 0.9, t, C_SECONDARY, fontsize=9)
        # 선
        ax.plot([5.0, x + 0.9], [6.0, y + 0.9], color='#AAAAAA', lw=0.8)

    # 2차 (케이스)
    sub1 = ['새 케이스\n등록', '진행 중 목록', '완료 이력', '긴급 알림']
    for i, t in enumerate(sub1):
        draw_box(ax, 0.02 + i*0.45 - 0.0, 3.0, 0.35, 1.2, t, C_ACCENT, fontsize=7.5)
        ax.plot([1.1, 0.17 + i*0.45 + 0.175], [4.6, 4.2], color='#AAAAAA', lw=0.6)

    # 2차 (지식베이스)
    sub2 = ['전기/전자', '추진', '제어/자동화', '배전반']
    for i, t in enumerate(sub2):
        draw_box(ax, 1.92 + i*0.45, 3.0, 0.35, 1.2, t, C_ACCENT, fontsize=7.5)
        ax.plot([3.0, 2.07 + i*0.45 + 0.175], [4.6, 4.2], color='#AAAAAA', lw=0.6)

    # 2차 (선박별)
    sub3 = ['선박\n프로필', '경보\n이력', '정비 일정', '계약 현황']
    for i, t in enumerate(sub3):
        draw_box(ax, 3.82 + i*0.45, 3.0, 0.35, 1.2, t, C_ACCENT, fontsize=7.5)
        ax.plot([4.9, 3.97 + i*0.45 + 0.175], [4.6, 4.2], color='#AAAAAA', lw=0.6)

    # 2차 (리포트)
    sub4 = ['SLA 리포트', '월간 요약', '재발 이슈', '비용 분석\n(유상)']
    for i, t in enumerate(sub4):
        draw_box(ax, 5.72 + i*0.45, 3.0, 0.35, 1.2, t, C_ACCENT, fontsize=7.5)
        ax.plot([6.8, 5.87 + i*0.45 + 0.175], [4.6, 4.2], color='#AAAAAA', lw=0.6)

    # 2차 (내 계정)
    sub5 = ['알림\n설정', '사용자\n관리', '권한/역할', 'SSO 연결']
    for i, t in enumerate(sub5):
        draw_box(ax, 7.62 + i*0.45, 3.0, 0.35, 1.2, t, C_ACCENT, fontsize=7.5)
        ax.plot([8.7, 7.77 + i*0.45 + 0.175], [4.6, 4.2], color='#AAAAAA', lw=0.6)

    # 유상 서비스 영역 (하단)
    draw_box(ax, 2.0, 0.8, 6.0, 1.2, '▼ 유상 서비스 (Phase 3 ~)   원격 진단 요청 · 부품 주문 · 엔지니어 파견 요청 · LCA 계약',
             C_DARK, fontsize=9)
    ax.text(5.0, 2.3, '무상 사용자는 자연스럽게 진입 — 별도 회원가입/결제 UI 없음',
            ha='center', fontsize=8.5, color='#666666', style='italic')

    plt.title('정보 아키텍처 (IA) — 5개 1차 메뉴 + 유상 서비스 확장 영역',
              fontsize=13, fontweight='bold', color=C_PRIMARY, pad=12)
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "web_sitemap.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def make_user_flow():
    """유저 플로우 — 경보 발생부터 해결까지"""
    fig, ax = plt.subplots(figsize=(9, 4.2))
    ax.set_xlim(0, 11); ax.set_ylim(0, 4.5)
    ax.axis('off')

    steps = [
        (0.2,  2.5, 'Hi-4S\n경보 감지',        '#7F8C8D'),
        (1.8,  2.5, '포털 자동 케이스\n생성 (초안)', C_ACCENT),
        (3.6,  2.5, '선주 PM\n모바일 푸시',    C_SECONDARY),
        (5.4,  2.5, 'PM 웹 진입\n케이스 확인',  C_SECONDARY),
        (7.2,  2.5, 'Tier 1 봇\n초기 분석',     C_PRIMARY),
        (9.0,  2.5, '당번 엔지니어\n15분 내 응답', C_PRIMARY),
    ]
    for x, y, t, c in steps:
        draw_box(ax, x, y, 1.5, 1.0, t, c, fontsize=8.5)

    # 화살표
    for i in range(5):
        arr = FancyArrowPatch((steps[i][0] + 1.55, 3.0), (steps[i+1][0] - 0.05, 3.0),
                               arrowstyle='->', mutation_scale=15, color='#333333', lw=1.3)
        ax.add_patch(arr)

    # 분기 (2번째 루트: 해결 → 종료)
    draw_box(ax, 3.6, 0.8, 1.5, 1.0, '해결 방안\n전달', C_PRIMARY, fontsize=8.5)
    draw_box(ax, 5.4, 0.8, 1.5, 1.0, 'PM 확인·\n종료', C_SECONDARY, fontsize=8.5)
    draw_box(ax, 7.2, 0.8, 1.5, 1.0, '이력 자동\n저장', '#7F8C8D', fontsize=8.5)
    draw_box(ax, 9.0, 0.8, 1.5, 1.0, '지식베이스\n자동 갱신', '#7F8C8D', fontsize=8.5)

    # 아래 화살표
    for i, (x1, x2) in enumerate([(3.6, 5.4), (5.4, 7.2), (7.2, 9.0)]):
        arr = FancyArrowPatch((x1 + 1.55, 1.3), (x2 - 0.05, 1.3),
                               arrowstyle='->', mutation_scale=15, color='#333333', lw=1.3)
        ax.add_patch(arr)
    # 위에서 아래로
    arr = FancyArrowPatch((9.75, 2.45), (4.35, 1.85),
                           arrowstyle='->', mutation_scale=15, color='#999999', lw=1.0, linestyle='dashed')
    ax.add_patch(arr)

    plt.title('유저 플로우 예시 — "새벽 경보" 시나리오 (발생부터 종료까지)',
              fontsize=13, fontweight='bold', color=C_PRIMARY, pad=12)
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "web_user_flow.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def make_data_model():
    """데이터 모델 (ERD)"""
    fig, ax = plt.subplots(figsize=(9, 5.2))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6)
    ax.axis('off')

    entities = [
        (0.3, 4.0, 2.0, 1.6, 'User\n─────\nid · name\nrole · email\ncompany · ship_ids'),
        (3.0, 4.0, 2.0, 1.6, 'Ship\n─────\nimo · name\ntype · flag\ndelivery_date'),
        (5.7, 4.0, 2.0, 1.6, 'Case\n─────\nid · ship_id\ntitle · severity\nstatus · sla_tier'),
        (8.0, 4.0, 1.7, 1.6, 'Message\n─────\ncase_id · author\ntext · attachments\ntimestamp'),
        (0.3, 1.5, 2.2, 1.6, 'KnowledgeArticle\n─────\nid · title · domain\nbody · linked_cases'),
        (2.8, 1.5, 2.2, 1.6, 'AlertEvent\n─────\nship_id · source\nseverity · raw_data\ncase_id (link)'),
        (5.3, 1.5, 2.2, 1.6, 'SLA_Log\n─────\ncase_id · stage\ntarget_ts · actual_ts\nbreach_flag'),
        (7.8, 1.5, 2.0, 1.6, 'Handover\n─────\ncase_id · from_user\nto_user · note\nhandover_ts'),
    ]
    for x, y, w, h, t in entities:
        box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.03",
                              linewidth=1.3, edgecolor=C_PRIMARY, facecolor='white')
        ax.add_patch(box)
        ax.text(x + w/2, y + h - 0.15, t.split('\n')[0],
                ha='center', va='top', fontsize=10, fontweight='bold', color=C_PRIMARY)
        lines = t.split('\n')[1:]
        for i, line in enumerate(lines):
            ax.text(x + 0.1, y + h - 0.45 - i*0.22, line,
                    ha='left', va='top', fontsize=8, color='#444444', family='monospace')

    # 관계선
    def rel(a, b):
        ax.annotate('', xy=b, xytext=a,
                    arrowprops=dict(arrowstyle='-', color='#888888', lw=1.0))
    rel((2.3, 4.8), (3.0, 4.8))   # User - Ship
    rel((5.0, 4.8), (5.7, 4.8))   # Ship - Case
    rel((7.7, 4.8), (8.0, 4.8))   # Case - Message
    rel((6.7, 4.0), (6.4, 3.1))   # Case - SLA_Log
    rel((6.7, 4.0), (8.8, 3.1))   # Case - Handover
    rel((6.7, 4.0), (3.9, 3.1))   # Case - AlertEvent
    rel((6.7, 4.0), (1.4, 3.1))   # Case - Knowledge

    plt.title('데이터 모델 (Core Entities) — 선박·케이스를 중심으로 한 관계 구조',
              fontsize=13, fontweight='bold', color=C_PRIMARY, pad=12)
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "web_data_model.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def make_tier_flow():
    """Tier 에스컬레이션"""
    fig, ax = plt.subplots(figsize=(8, 3.6))
    ax.set_xlim(0, 10); ax.set_ylim(-1.4, 2.0)
    ax.axis('off')

    tiers = [
        (0.3, 'Tier 1\n자동 응답\n(0초)',   '#5D8CA8', 'FAQ·지식베이스\nAI 매칭 · 봇'),
        (2.6, 'Tier 2\n당번 엔지니어\n(15분)', C_SECONDARY, 'Follow-the-Sun\n해외지사 정규근무'),
        (4.9, 'Tier 3\n도메인 전문가\n(2시간)',  C_PRIMARY,  '전기·추진·제어·\n배전반 전문가 풀'),
        (7.2, 'Tier 4\nR&D·본사\n(24시간)',     C_DARK,     '복잡 건 에스컬레이션\n영업시간 내'),
    ]
    for x, t, c, desc in tiers:
        draw_box(ax, x, 0.5, 2.0, 1.2, t, c, fontsize=10)
        ax.text(x + 1.0, -0.3, desc, ha='center', va='top',
                fontsize=8.5, color='#444444')

    for i in range(3):
        arr = FancyArrowPatch((tiers[i][0] + 2.05, 1.1), (tiers[i+1][0] - 0.02, 1.1),
                               arrowstyle='->', mutation_scale=18, color='#333333', lw=1.5)
        ax.add_patch(arr)
        ax.text(tiers[i][0] + 2.15, 1.35, '해결 못함 시\n에스컬레이션',
                fontsize=7, color='#666666', style='italic')

    plt.title('Tier 에스컬레이션 플로우 — 단순은 자동, 복잡은 전문가에게',
              fontsize=13, fontweight='bold', color=C_PRIMARY, pad=12)
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "web_tier_flow.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def make_roadmap_timeline():
    """무상 → 유상 확장 타임라인"""
    fig, ax = plt.subplots(figsize=(9, 3.4))
    ax.set_xlim(0, 12); ax.set_ylim(0, 4)
    ax.axis('off')

    phases = [
        (0.3, 'Phase 1\n무상 MVP',    '0~6개월',     C_DARK,
         'FAQ · 채팅 · 케이스\n기존 선주 Closed Beta'),
        (3.1, 'Phase 2\n24/7 정식 운영', '6~12개월',  C_PRIMARY,
         '4-Tier · 해외지사\nSLA · 모바일 앱'),
        (5.9, 'Phase 3\n유상 서비스 도입', '1~2년',   C_SECONDARY,
         '원격 진단 · 부품\n파견 요청 · 정산 연동'),
        (8.7, 'Phase 4\n통합 포털',     '2년+',        C_ACCENT,
         'LCA 계약 · 대시보드\n예측정비 · 리포트'),
    ]
    for x, t, d, c, desc in phases:
        draw_box(ax, x, 2.0, 2.5, 1.2, t, c, fontsize=10.5)
        ax.text(x + 1.25, 1.65, d, ha='center', fontsize=9,
                fontweight='bold', color=c)
        ax.text(x + 1.25, 0.7, desc, ha='center', va='top',
                fontsize=8.5, color='#444444')

    # 무상/유상 구분선
    ax.axvline(5.7, ymin=0.05, ymax=0.95, color=C_RED, linestyle='--', lw=1.2, alpha=0.6)
    ax.text(2.9, 3.7, '◀  무상 영역  ▶', ha='center', fontsize=10,
            fontweight='bold', color=C_PRIMARY)
    ax.text(9.2, 3.7, '◀  유상 영역 (선사 회계 가능 건 한정)  ▶',
            ha='center', fontsize=10, fontweight='bold', color=C_SECONDARY)

    plt.title('무상 → 유상 확장 로드맵 — 선주가 쓰던 창구 안에서 자연 전환',
              fontsize=13, fontweight='bold', color=C_PRIMARY, pad=12)
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "web_roadmap.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def make_pm_improvement():
    """PM 업무 개선 정량 효과"""
    fig, ax = plt.subplots(figsize=(8, 3.8))
    categories = ['메일/포워딩\n처리 시간', '케이스당\n평균 응답', '이력 조회\n소요 시간', '주간 회의\n준비 시간', '재발 이슈\n대응 시간']
    before = [100, 100, 100, 100, 100]
    after  = [35, 15, 8, 40, 30]

    x = np.arange(len(categories))
    w = 0.36
    ax.bar(x - w/2, before, w, label='Before (메일 기반)', color=C_GRAY, zorder=3)
    ax.bar(x + w/2, after,  w, label='After (웹 포털 도입)', color=C_PRIMARY, zorder=3)

    for i, (b, a) in enumerate(zip(before, after)):
        ax.text(x[i] - w/2, b + 2, f'{b}',  ha='center', fontsize=9, color='#666666')
        ax.text(x[i] + w/2, a + 2, f'{a}%', ha='center', fontsize=9,
                color=C_PRIMARY, fontweight='bold')

    ax.set_ylabel('소요 시간 (Before = 100 기준)', fontsize=9, color='#666666')
    ax.set_xticks(x); ax.set_xticklabels(categories, fontsize=9.5, fontweight='bold')
    ax.set_ylim(0, 115); ax.legend(fontsize=10, loc='upper right')
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.grid(axis='y', alpha=0.3, linestyle='--'); ax.set_axisbelow(True)
    plt.title('PM 업무 개선 정량 효과 — 웹 포털 도입 후 소요시간 변화 (예상치)',
              fontsize=12, fontweight='bold', color=C_PRIMARY, pad=15)
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "web_pm_improvement.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


print("차트 생성 중...")
img_arch      = make_system_arch()
img_sitemap   = make_sitemap()
img_user_flow = make_user_flow()
img_data      = make_data_model()
img_tier_flow = make_tier_flow()
img_roadmap   = make_roadmap_timeline()
img_pm        = make_pm_improvement()
print("차트 생성 완료")


# ══════════════════════════════════════
#  문서 생성
# ══════════════════════════════════════

print("Word 문서 생성 중...")
doc = Document()
init_doc(doc)

# ────── 표지 ──────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(120)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
run = p.add_run('━' * 42)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C); run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
run = p.add_run('24/7 무상 기술 지원 웹사이트\n상세 설계서')
run.font.size = Pt(24); run.font.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C); run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run('아키텍처 · 정보구조 · 기능 · 유저플로우 · 데이터모델\n그리고 PM 업무 개선을 위한 유상 서비스 확장 로드맵')
run.font.size = Pt(12); run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
run = p.add_run('━' * 42)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C); run.font.size = Pt(10)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(130)

for t in ['2026. 04', 'HD현대마린솔루션테크']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    run = p.add_run(t); run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88); run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_page_break()

# ════════════════════════════════════════════
#  목차
# ════════════════════════════════════════════

doc.add_heading('목차', level=1)

add_table(doc,
    ['장', '제목', '주요 내용'],
    [
        ['서론',      '배경·목적·설계 원칙',        '메일 기반 한계 / PM 업무 과부하 / 설계 목표 5가지'],
        ['본론 1',    '시스템 아키텍처',             '5계층 구조 · 외부 연계 · 비기능 요구사항'],
        ['본론 2',    '정보 아키텍처 (IA)',          '사이트맵 · 화면 구성 · 메뉴 계층'],
        ['본론 3',    '핵심 기능 상세',              '채팅·케이스·지식베이스·알림·리포트'],
        ['본론 4',    '유저 플로우',                 '4가지 대표 시나리오 (경보·문의·사관·유상)'],
        ['본론 5',    '데이터 모델 · 보안',           '핵심 엔터티 · 권한 · 인증 · 감사'],
        ['본론 6',    '무상 → 유상 확장 로드맵',      '4-Phase · PM 업무 개선 중심 유상 설계'],
        ['본론 7',    '기대 효과 (정량·정성)',       'PM 업무 / 선주 / HMS 관점별'],
        ['결론',      '요약·성공 조건·다음 단계',     '핵심 메시지 · 리스크 · 실행 권고'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════
#  서론
# ════════════════════════════════════════════

doc.add_heading('서론 — 왜 웹사이트를 설계하는가', level=1)

doc.add_heading('1. 배경', level=2)

body(doc,
    '현재 HD현대마린솔루션테크의 선주·선박 기술 지원은 **이메일·전화 중심**으로 운영. '
    '이 방식은 익숙하지만 구조적 한계를 가짐')

bullet(doc, '**시간대 편차** — 선박은 24시간 운항하나, 지원은 영업시간 중심')
bullet(doc, '**이력 파편화** — 각 PM 메일함에 분산되어 재발 이슈 추적 불가')
bullet(doc, '**전문가 매칭 지연** — 메일 담당자가 직접 포워딩해야 함')
bullet(doc, '**PM 업무 과부하** — 회신·이관·요약·재촉이 모두 수동')
bullet(doc, '**지식 자산화 실패** — 해결 사례가 메일함에 갇혀 조직 자산이 되지 못함')

doc.add_heading('2. 본 설계서의 목적', level=2)

body(doc,
    '**24/7 무상 기술 지원 웹사이트**를 상세 설계. '
    '초기에는 무상으로 접근성을 확보하고, 이후 **PM 업무 개선을 위한 유상 서비스**로 확장하는 단계적 구축안')

insight(doc,
    '설계 원칙 5가지\n'
    '① 메일보다 편해야 한다 (대체 아닌 "상위 호환")\n'
    '② 선주가 "새로 가입"하지 않아도 된다 (SSO · 자동 계정)\n'
    '③ HMS가 먼저 시작하는 구조 (경보 → 케이스 자동 생성)\n'
    '④ 선박 단위 이력이 자산이 된다 (검색·재사용 가능)\n'
    '⑤ 무상 → 유상 전환이 "새 시스템 가입"이 아닌 "버튼 클릭"')

doc.add_heading('3. 본 설계서의 범위', level=2)

add_table(doc,
    ['구분', '포함', '제외'],
    [
        ['**기능 영역**',   '기술 문의 · 케이스 · 지식베이스 · 선박 이력',          '제3자 수리 중개 · 금융·보험 서비스'],
        ['**사용자**',      '선주 PM · 선박 사관 · HMS 엔지니어 · 관리자',          '일반 공중 · 비고객 선주'],
        ['**시스템 연계**', 'Hi-4S · SmartCare · AS 계약 ERP · 메일 게이트웨이',     'HD현대 외 타 조선소 건조선 데이터'],
        ['**단계 범위**',   'Phase 1(무상 MVP) ~ Phase 4(통합 포털) 전체 설계',     '실제 구현 · 상세 UI 시안 · 세부 API 스펙'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════
#  본론 1 — 시스템 아키텍처
# ════════════════════════════════════════════

doc.add_heading('본론 1 — 시스템 아키텍처', level=1)

doc.add_heading('1-1. 전체 구조 (5계층)', level=2)

chart(doc, img_arch, '24/7 무상 기술 지원 웹사이트 — 시스템 아키텍처 5계층 구조')

body(doc,
    '시스템은 **5개 계층(User · Frontend · Application · Data · External)** 으로 구성. '
    '각 계층은 책임이 분리되어 있어 향후 유상 서비스 확장 시 **Application 계층에 모듈을 추가**하면 됨 — '
    '데이터·프론트·외부 연계는 재사용')

add_table(doc,
    ['계층', '구성 요소', '주요 책임'],
    [
        ['**User**',        '선주 PM · 선박 사관 · HMS 엔지니어 · 도메인 전문가 · 관리자',
         '역할 기반 권한(RBAC)으로 각자 다른 화면 제공'],
        ['**Frontend**',    '웹 포털(PC) · 모바일 앱(iOS/Android) · 이메일 게이트웨이',
         '선주가 편한 경로로 접근 — 모든 채널은 동일 백엔드 사용'],
        ['**Application**', '인증(SSO) · 케이스 관리 · 채팅/Tier 라우팅 · 지식베이스 · 알림·리포트',
         '5개 핵심 서비스 — 각 서비스가 독립 모듈로 운영'],
        ['**Data**',        'Core DB · Knowledge Store · Event Log',
         '선박·케이스·유저·메시지·감사 로그 저장'],
        ['**External**',    'Hi-4S · SmartCare · ERP · Mail Gateway',
         '기존 HMS 시스템과 API·이벤트 기반 연동'],
    ])

doc.add_heading('1-2. 외부 시스템 연계 상세', level=2)

add_table(doc,
    ['외부 시스템', '연계 방식', '데이터 흐름'],
    [
        ['**Hi-4S**',            'REST API + Webhook',      '선박 경보 발생 → Webhook → 케이스 자동 초안 생성'],
        ['**SmartCare 관제**',    '내부 이벤트 버스',         '관제 담당자의 수동 알림 → 포털 케이스 연동'],
        ['**AS 계약 ERP**',       'REST API (읽기 전용 초기)','선박별 계약 상태 조회 · 유상 서비스 정산 연동(Phase 3~)'],
        ['**Mail Gateway**',     'SMTP/IMAP 양방향',          '웹 케이스 → 메일 발송 · 선주 메일 회신 → 케이스 자동 업데이트'],
        ['**SSO (OAuth 2.0)**',   '선주사 내부 IdP 연동',     '선주사 직원 이메일로 원클릭 로그인'],
    ])

doc.add_heading('1-3. 비기능 요구사항 (NFR)', level=2)

add_table(doc,
    ['항목', '요구 수준'],
    [
        ['**가용성**',        '99.9% (월 43분 이내 다운타임) — 무중단 배포 지원'],
        ['**응답 시간**',      '화면 로딩 2초 이내 · 채팅 메시지 전송 500ms 이내'],
        ['**동시 사용자**',    '초기 500명 · Phase 4까지 10,000명 확장 가능하도록 설계'],
        ['**데이터 보관**',    '케이스 이력 10년 · 메시지 5년 · 감사 로그 7년'],
        ['**보안**',           'TLS 1.3 · RBAC · 감사 로그 · 개인정보 암호화 저장'],
        ['**접근성**',          'WCAG 2.1 AA 준수 · 한/영/일/중 4개 언어'],
        ['**모바일**',          'iOS 14+ / Android 10+ / 위성 저대역폭 환경 최적화'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════
#  본론 2 — 정보 아키텍처 (IA)
# ════════════════════════════════════════════

doc.add_heading('본론 2 — 정보 아키텍처 (IA) · 화면 구성', level=1)

doc.add_heading('2-1. 사이트 맵', level=2)

chart(doc, img_sitemap, '정보 아키텍처 — 5개 1차 메뉴 · 유상 서비스는 자연스러운 확장 영역')

body(doc,
    '**5개 1차 메뉴**(케이스·지식베이스·선박별 현황·리포트·내 계정)를 중심으로 구성. '
    '유상 서비스는 별도 섹션이 아닌 **기존 메뉴 안에서 "원격 진단 요청"·"부품 주문" 버튼**으로 자연스럽게 노출 — '
    '선주가 새 시스템으로 이동하는 느낌을 주지 않음')

doc.add_heading('2-2. 주요 화면 상세 (Screen Specifications)', level=2)

add_table(doc,
    ['#', '화면', '주요 구성 요소', '핵심 액션'],
    [
        ['1', '**홈 대시보드**',
         '내 선대 상태 요약 · 진행 중 케이스 · 최근 경보 · 공지사항',
         '케이스 바로가기 · 긴급 알림 확인'],
        ['2', '**케이스 목록**',
         '필터(선박·상태·Tier·기간) · 정렬 · SLA 시계 · 담당자',
         '새 케이스 등록 · 검색 · 엑셀 내보내기'],
        ['3', '**케이스 상세**',
         '타임라인 · 채팅창 · 첨부파일 · 상태 · 담당자 · 핸드오프 기록',
         '메시지 작성 · 파일 첨부 · 상태 변경 · 에스컬레이션'],
        ['4', '**지식베이스**',
         '도메인별 카테고리 · 검색 · FAQ · 매뉴얼 · 해결 사례',
         '기사 조회 · 관련 케이스 연결 · 북마크'],
        ['5', '**선박별 현황**',
         '선박 프로필 · Hi-4S 상태 요약 · 정비 일정 · 케이스 이력 · 계약 현황',
         '경보 이력 조회 · 유상 서비스 요청 (Phase 3~)'],
        ['6', '**리포트**',
         '월간 SLA 준수율 · 재발 이슈 Top 5 · 처리 시간 분포',
         'PDF 내보내기 · 이메일 정기 발송 설정'],
        ['7', '**내 계정**',
         '프로필 · 알림 설정(이메일/SMS/푸시) · 사용자 관리 · SSO 연결',
         '권한 관리 · 초대 · 비밀번호 변경'],
    ])

doc.add_heading('2-3. 홈 대시보드 레이아웃 상세', level=2)

body(doc,
    '선주 PM이 **로그인 직후 가장 먼저 보는 화면**. \"내가 지금 알아야 할 것\"을 한눈에 파악')

bullet(doc, '**상단 알림 배너** — 긴급·High 등급 케이스만 빨간 배너로 표시')
bullet(doc, '**좌측 내 선대 카드** — 선박별 상태 요약 (운항 중/경보 발생/정비 중)')
bullet(doc, '**중앙 진행 중 케이스 리스트** — SLA 시계와 함께 최근 3~5건')
bullet(doc, '**우측 지식 추천** — 최근 문의 기반으로 관련 FAQ·매뉴얼 자동 추천')
bullet(doc, '**하단 월간 요약** — 이번 달 케이스 수·해결 시간·재발 이슈 수')

doc.add_page_break()

# ════════════════════════════════════════════
#  본론 3 — 핵심 기능 상세
# ════════════════════════════════════════════

doc.add_heading('본론 3 — 핵심 기능 상세', level=1)

doc.add_heading('3-1. 24/7 실시간 채팅 · Tier 에스컬레이션', level=2)

chart(doc, img_tier_flow, 'Tier 에스컬레이션 플로우 — 단순은 자동, 복잡은 전문가에게')

body(doc,
    '모든 문의는 **Tier 1부터 자동 접수**되고, 해결이 되지 않으면 다음 Tier로 자동 에스컬레이션. '
    '선주 PM은 \"지금 누가 보고 있는가\"를 항상 실시간으로 확인 가능')

add_table(doc,
    ['Tier', '응답 주체', '응답 시간 (SLA)', '커버 범위'],
    [
        ['Tier 1', '자동 봇 + AI 지식베이스 매칭',              '즉시 (0초)',    'FAQ · 매뉴얼 조회성 질문 · 상태 문의'],
        ['Tier 2', '당번 엔지니어 (HMS 해외지사 Follow-the-Sun)', '15분 이내',     '일반 기술 문의 · 케이스 1차 진단'],
        ['Tier 3', '도메인 전문가 풀 (전기·추진·제어·배전반)',    '2시간 이내',    '도메인 특화 진단 · 원격 조치 제안'],
        ['Tier 4', 'R&D · 한국 본사 엔지니어',                   '24시간 이내',   '복잡 건 · 설계 수준 · 안전 관련'],
    ])

doc.add_heading('3-2. 케이스 관리 — 접수부터 종료까지', level=2)

body(doc,
    '케이스는 **6단계 상태(Status)를 자동 추적**. 각 단계는 SLA와 연결되어 기한 초과 시 자동 알림')

add_table(doc,
    ['단계', '상태', '진입 트리거', 'SLA (긴급 기준)'],
    [
        ['1', '**접수 (Submitted)**',       '선주 등록 · 경보 자동 생성 · 메일 수신',  '즉시 시스템 수신'],
        ['2', '**분류 (Triaged)**',          'AI 자동 분류 또는 담당자 수동 분류',     '5분 이내'],
        ['3', '**진단 중 (In Diagnosis)**',  '담당 엔지니어 배정 · 1차 분석 시작',    '15분 이내'],
        ['4', '**조치 중 (In Resolution)**', '원격 조치 · 추가 정보 확인',            '2시간 이내'],
        ['5', '**확인 대기 (Awaiting)**',     '해결 방안 전달 · 선주 확인 요청',       '-'],
        ['6', '**종료 (Closed)**',            '선주 확인 완료 · 이력 자동 저장',       '24시간 이내'],
    ])

body(doc,
    '각 상태 변경 시 **자동 타임라인 기록**. 담당자 변경(핸드오프) · 메시지 · 첨부파일 · SLA 체크가 모두 타임라인에 남음')

doc.add_heading('3-3. 지식베이스 · FAQ', level=2)

bullet(doc, '**도메인별 카테고리** — 전기/전자 · 추진 · 제어/자동화 · 배전반 · 공조 · 기타')
bullet(doc, '**AI 시맨틱 검색** — 자연어 질문에 가장 가까운 기사 자동 제시')
bullet(doc, '**케이스 → 기사 자동 환류** — 자주 반복되는 해결 사례를 월간 리뷰로 기사화')
bullet(doc, '**선박 타입·장비 기준 필터** — 컨테이너선·LNG선·벌크선 등 선종별')
bullet(doc, '**다국어 지원** — 영문 원문 + 한/일/중 자동 번역 병기')

doc.add_heading('3-4. 알림 · 통합 (Notifications)', level=2)

add_table(doc,
    ['채널', '발송 조건', '설정 주체'],
    [
        ['**이메일**',     '케이스 상태 변경 · SLA 임박 · 주간 요약',           '선주 PM 개별 설정'],
        ['**SMS**',        '긴급·High 등급 케이스만',                          'PM + 관리자 승인'],
        ['**모바일 푸시**', '모바일 앱 설치자에 한해 모든 업데이트',            '사용자 선택'],
        ['**웹 팝업**',     '로그인 중 실시간 업데이트',                        '기본 ON'],
        ['**메일 회신 연동**','선주가 이메일에 회신 → 케이스 메시지 자동 추가', '기본 ON'],
    ])

doc.add_heading('3-5. 리포트 · 대시보드', level=2)

bullet(doc, '**SLA 리포트** — 월간 응답/해결 시간 준수율 · 미준수 사유 분석')
bullet(doc, '**재발 이슈 리포트** — 같은 선박·같은 장비에서 반복된 문의 Top N')
bullet(doc, '**Top FAQ 리포트** — 선주가 가장 많이 찾은 기사 — 제품 개선 인풋')
bullet(doc, '**비용 분석** — 유상 서비스 이용 현황 (Phase 3~)')
bullet(doc, '**자동 배포** — 지정된 수신자에게 매월 1일 PDF 자동 발송')

doc.add_heading('3-6. 권한 · 역할 (RBAC)', level=2)

add_table(doc,
    ['역할', '주요 권한', '제약'],
    [
        ['**선주 PM (Master)**',  '선사 소유 전 선박 조회 · 사용자 초대 · 권한 부여',       '타 선사 데이터 접근 불가'],
        ['**선주 PM (Member)**',   '배정된 선박 조회 · 케이스 등록/조회',                   '사용자 관리 불가'],
        ['**선박 사관**',          '해당 선박 케이스 등록 · 모바일 중심 UI',                '리포트·재무 정보 접근 불가'],
        ['**HMS 엔지니어 (Tier 2)**','당번 시간대 모든 케이스 · 에스컬레이션 권한',         '유상 서비스 단가 변경 불가'],
        ['**HMS 전문가 (Tier 3)**', '도메인 케이스 전체 · 지식베이스 기사 작성',            '-'],
        ['**관리자**',              '시스템 설정 · 사용자 전체 · 감사 로그',                 '선주 데이터 열람은 감사 기록 남김'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════
#  본론 4 — 유저 플로우
# ════════════════════════════════════════════

doc.add_heading('본론 4 — 유저 플로우 (대표 시나리오 4건)', level=1)

doc.add_heading('시나리오 1. 새벽 경보 → 해결까지 (대표 플로우)', level=2)

chart(doc, img_user_flow, '유저 플로우 예시 — "새벽 경보" 시나리오 (발생부터 종료까지)')

body(doc,
    '**상황**: 대서양 항해 중 선박에서 새벽 03:00(KST)에 배전반 경보 발생. '
    '기존 방식이라면 PM은 아침까지 모르거나, 알아도 영업시간까지 기다려야 함. 본 시스템에서는 다음 흐름')

add_table(doc,
    ['단계', '시점', '시스템 행동', '선주 PM 행동'],
    [
        ['1', 'T+0초',    'Hi-4S가 경보 감지 → Webhook으로 포털 전달',        '(잠자는 중)'],
        ['2', 'T+5초',    '포털이 케이스 초안 자동 생성 — 경보 정보 자동 입력', '(잠자는 중)'],
        ['3', 'T+10초',   '선주 PM 모바일 앱에 푸시 · SMS(긴급 등급)',         '푸시 수신 · 인지'],
        ['4', 'T+30초',   'Tier 2 당번 엔지니어(휴스턴 센터)에게 자동 배정',    '-'],
        ['5', 'T+3분',    '당번 엔지니어가 1차 데이터 분석 시작',             'PM 확인 (선택)'],
        ['6', 'T+15분',   '최초 응답 메시지 + 초기 진단 요약 전송',             '(선택) 추가 정보 제공'],
        ['7', 'T+2시간',  '원격 조치 실행 or 추가 진단 · 해결안 전달',          '조치 확인'],
        ['8', 'T+해결',   '선주 확인 → 케이스 종료 · 이력 자동 저장 · FAQ 후보 태깅','종료 확인'],
    ])

doc.add_heading('시나리오 2. 사관이 모바일로 사진 첨부 문의', level=2)

body(doc,
    '**상황**: 항해 중 사관이 특정 장비 상태를 눈으로 보고 이상 판단. '
    '모바일 앱에서 장비 사진 + 짧은 설명을 올려 케이스 생성')

bullet(doc, '① 모바일 앱 로그인 (SSO · 생체 인증)')
bullet(doc, '② 해당 선박 선택 → \"새 케이스\" 버튼')
bullet(doc, '③ 카테고리 선택 (전기/제어/추진 등) → 사진 2~3장 + 음성 메모 첨부')
bullet(doc, '④ 위성 저대역폭 환경 대응 — 업로드 전 자동 압축 · 재전송 큐')
bullet(doc, '⑤ 접수 즉시 Tier 1 봇이 유사 이력 3건 자동 추천 (사관 즉시 참조 가능)')
bullet(doc, '⑥ Tier 2 당번 엔지니어가 15분 내 응답')

doc.add_heading('시나리오 3. HMS 엔지니어의 Tier 2 대응', level=2)

body(doc,
    '**상황**: 휴스턴 센터 당번 엔지니어가 본인 근무 시간 중 케이스를 인계받음')

bullet(doc, '① 당번 시작 시 **인수인계 노트** 자동 표시 — 직전 센터(유럽)가 남긴 진행 중 케이스 요약')
bullet(doc, '② 각 케이스의 SLA 시계·담당 이력 확인')
bullet(doc, '③ 새 케이스 배정 시 자동 알림 + 맥락(선박 최근 이력·경보)이 자동 첨부')
bullet(doc, '④ 해결 불가 시 \"Tier 3로 에스컬레이션\" 버튼 — 도메인 전문가 풀 자동 매칭')
bullet(doc, '⑤ 퇴근 직전 **핸드오프 노트** 작성 (체크리스트 기반)')

doc.add_heading('시나리오 4. 유상 서비스 요청 — PM 업무 개선 포인트', level=2)

body(doc,
    '**상황**: 무상 Tier 2~3에서 해결 방안을 제시했으나, **\"현장 파견\"** 또는 **\"원격 진단 깊이 있는 수행\"** 이 필요한 상황. '
    'PM이 기존에는 별도 메일로 영업팀에 문의하고, 견적 주고받고, 발주까지 1~2주 걸렸던 업무')

add_table(doc,
    ['단계', '기존 방식 (메일)', '웹 포털 (유상 서비스 버튼)'],
    [
        ['1', 'PM이 영업팀에 별도 메일',             '케이스 상세 화면의 \"원격 진단 요청\" 버튼 클릭'],
        ['2', '영업팀이 계약 조건 확인 · 견적 산정', '자동으로 계약 기반 견적 표시 · 승인 요청'],
        ['3', 'PM 사내 결재 → 영업팀 회신',           '모바일 전자 결재 즉시 가능'],
        ['4', '영업팀이 엔지니어 일정 확인 · 파견',    '엔지니어 풀에서 자동 매칭 · 선주 확인'],
        ['5', '1~2주 소요',                          '평균 24시간 이내 확정'],
    ])

insight(doc,
    '유상 서비스의 핵심 = "PM이 별도 프로세스를 밟지 않는 것"\n'
    '이미 케이스 안에서 진행되던 대화가 → 버튼 하나로 유상 전환 →\n'
    '기존 계약·결재 구조가 자동 연결되는 것이 "PM 업무 개선"의 본질.')

doc.add_page_break()

# ════════════════════════════════════════════
#  본론 5 — 데이터 모델 · 보안
# ════════════════════════════════════════════

doc.add_heading('본론 5 — 데이터 모델 · 보안', level=1)

doc.add_heading('5-1. 핵심 데이터 엔터티 (ERD)', level=2)

chart(doc, img_data, '데이터 모델 — 선박·케이스를 중심으로 한 관계 구조')

add_table(doc,
    ['엔터티', '주요 속성', '역할'],
    [
        ['**User**',             'id, name, role, email, company, ship_ids',      '사용자 · RBAC 권한'],
        ['**Ship**',             'imo, name, type, flag, delivery_date',          '선박 마스터'],
        ['**Case**',             'id, ship_id, title, severity, status, sla_tier','기술 문의 단위'],
        ['**Message**',          'case_id, author, text, attachments, ts',        '케이스 내 대화'],
        ['**KnowledgeArticle**',  'id, title, domain, body, linked_cases',         'FAQ · 매뉴얼 · 해결 사례'],
        ['**AlertEvent**',       'ship_id, source, severity, raw_data, case_id',  'Hi-4S 경보 자동 연동'],
        ['**SLA_Log**',           'case_id, stage, target_ts, actual_ts, breach',  'SLA 준수/미준수 기록'],
        ['**Handover**',          'case_id, from_user, to_user, note, ts',         '교대 인수인계 로그'],
    ])

doc.add_heading('5-2. 보안 · 인증', level=2)

bullet(doc, '**전송 암호화** — TLS 1.3 전 구간 · 위성 통신 구간도 포함')
bullet(doc, '**저장 암호화** — 첨부파일·개인정보·선박 데이터 AES-256')
bullet(doc, '**인증** — SSO(OAuth 2.0 / SAML) · MFA 옵션 · 모바일 생체 인증')
bullet(doc, '**권한** — RBAC 5단계 (Master/Member/사관/엔지니어/관리자)')
bullet(doc, '**감사 로그** — 모든 조회·변경·다운로드 7년 보존 · 관리자 접근도 기록')
bullet(doc, '**데이터 분리** — 선사별 테넌트 분리 (Logical Multi-tenancy)')
bullet(doc, '**취약점 관리** — 연 1회 외부 침투 테스트 · 월 1회 자동 스캔')

doc.add_heading('5-3. 개인정보 · 선박 데이터 취급 원칙', level=2)

insight(doc,
    '선박 운항 데이터·경보 데이터는 "선주의 자산".\n'
    'HMS는 기술 지원 목적으로 "열람·분석"하지만, 2차 활용(마케팅 등)은 원칙적으로 금지.\n'
    '유상 서비스 정산 외에는 제3자 공유하지 않음.')

doc.add_page_break()

# ════════════════════════════════════════════
#  본론 6 — 무상 → 유상 확장 로드맵
# ════════════════════════════════════════════

doc.add_heading('본론 6 — 무상 → 유상 확장 로드맵 (PM 업무 개선 중심)', level=1)

chart(doc, img_roadmap, '무상 → 유상 확장 로드맵 — 선주가 쓰던 창구 안에서 자연 전환')

doc.add_heading('6-1. Phase별 상세', level=2)

add_table(doc,
    ['Phase', '기간', '포함 기능', '대상'],
    [
        ['**Phase 1 — 무상 MVP**',  '0~6개월',
         'FAQ · 케이스 등록/조회 · 실시간 채팅(Tier 1~2) · 선박별 이력 · 기본 알림',
         '기존 AS 선주 10~20개사 Closed Beta'],
        ['**Phase 2 — 24/7 정식**', '6~12개월',
         '4-Tier 완전 운영 · Follow-the-Sun · SLA 문서화 · 모바일 앱 · 리포트',
         '전 기존 선주 확대'],
        ['**Phase 3 — 유상 서비스**','1~2년',
         '원격 진단 · 부품 주문 · 엔지니어 파견 요청 · ERP 정산 연동',
         '기존 무상 사용자 대상 자연 전환'],
        ['**Phase 4 — 통합 포털**', '2년+',
         'LCA 계약 관리 · 예측정비 대시보드 · KPI 리포트 · 파트너 연계',
         '장기 계약 선주 중심'],
    ])

doc.add_heading('6-2. 유상 서비스 설계 — PM 업무 개선 관점', level=2)

body(doc,
    '유상 서비스는 \"새 결제 시스템\"이 아니라 **\"PM이 기존에 따로 하던 업무를 웹 포털 안에서 버튼 하나로 끝낼 수 있게\"** 설계')

add_table(doc,
    ['유상 서비스', 'PM의 기존 업무 (제거되는 부분)', '웹 포털에서의 처리 방식'],
    [
        ['**원격 진단 심화**',      '영업팀 메일 → 견적 → 발주 → 엔지니어 배정 (평균 1주)',
         '케이스 상세에서 \"원격 진단 요청\" → 계약 기반 견적 자동 표시 → 전자 결재'],
        ['**부품 주문**',            '부품 카탈로그 PDF 검색 → 메일 견적 → 발주',
         '선박 프로필 기반 호환 부품 자동 추천 → 클릭 주문 · ERP 자동 연동'],
        ['**엔지니어 파견**',        '일정 조율 메일 수차례 · PM이 중간 조정',
         '엔지니어 풀 가용일 자동 표시 → 선주 선택 → 자동 확정'],
        ['**LCA 계약 갱신**',        '계약 만료 2~3개월 전 영업팀 방문 · 협의',
         '대시보드에서 현재 사용 패턴 기반 추천안 표시 → 전자 서명'],
        ['**정기 리포트 제작**',     'PM이 매월 메일·엑셀로 수작업 취합',
         '자동 생성된 월간 리포트 · 필요 시 커스텀 리포트 신청'],
    ])

insight(doc,
    '유상 서비스의 핵심 목표 = "매출 증대"가 아니라 "PM 업무 축소 → 자연 전환"\n'
    '편해서 쓰게 되는 유상 서비스가 진짜 PM 업무 개선이며,\n'
    '그 결과로 매출은 "따라옴".')

doc.add_heading('6-3. 무상 → 유상 전환 UX — 장벽 제거 설계', level=2)

bullet(doc, '**회원가입 없음** — 무상 사용자가 이미 계정 있음 → 유상도 즉시 접근')
bullet(doc, '**결제 팝업 없음** — 이미 체결된 LCA·AS 계약이 기본 결제 수단')
bullet(doc, '**승인 절차 내재화** — 선주사 사내 결재가 웹 안에서 완결')
bullet(doc, '**사용 패턴 기반 제안** — \"이 케이스는 원격 진단을 추가하면 3시간 → 30분\" 같은 인라인 추천')
bullet(doc, '**투명한 요금** — 서비스별 단가가 케이스 상세에 자동 표시 · 숨겨진 비용 없음')

doc.add_page_break()

# ════════════════════════════════════════════
#  본론 7 — 기대 효과
# ════════════════════════════════════════════

doc.add_heading('본론 7 — 기대 효과 (정량·정성)', level=1)

doc.add_heading('7-1. PM 업무 개선 정량 효과', level=2)

chart(doc, img_pm, 'PM 업무 개선 정량 효과 — 웹 포털 도입 후 소요시간 변화 (예상치)')

add_ba_table(doc,
    [
        ('**메일 처리 시간**',        '하루 평균 2~3시간',             '35% 수준 — 하루 1시간 이내'),
        ('**케이스 응답 시간**',       '평균 4~12시간',                 '15% 수준 — 평균 15분~1시간'),
        ('**이력 조회 시간**',         '재발 이슈 조회 30분~1시간',     '8% 수준 — 10초 이내 검색'),
        ('**주간 회의 준비**',          '엑셀 취합 · 자료 작성 3~4시간', '40% 수준 — 자동 리포트 활용'),
        ('**재발 이슈 대응**',          '매번 처음부터 진단',            '30% 수준 — 기존 해결사례 즉시 참조'),
        ('**유상 서비스 요청 처리**',   '1~2주 (메일·결재 왕복)',        'Phase 3부터 평균 24시간 이내'),
    ],
    before_label='Before (메일 기반)',
    after_label='After (웹 포털 도입)')

doc.add_heading('7-2. 정성 효과 — 3자 관점별', level=2)

add_table(doc,
    ['관점', '핵심 효과'],
    [
        ['**선주 PM**',
         '반복 업무 자동화 · 본업(선대 관리·의사결정)에 집중 가능 · 재발 이슈 즉시 참조 · 회사 내 지식 자산화'],
        ['**선박 사관**',
         '위성 환경에서도 모바일로 즉시 문의 · 15분 내 응답 · 과거 선박 이력 기반 맞춤 지원'],
        ['**HMS 엔지니어**',
         'Tier 분리로 단순 질문 자동 처리 · 전문가는 고부가 업무 집중 · 케이스 핸드오프 시스템화로 인수인계 부담 감소'],
        ['**HMS 경영**',
         '고객 관계 지속성 확보 · 무상→유상 전환 깔때기 확보 · OEM 포털 포지션(3사 대비) 구축 가속'],
        ['**HMS 영업/마케팅**',
         'Top FAQ · 재발 이슈 데이터가 제품 개선·영업 제안 근거 · LCA 전환 비중 확대 기반'],
    ])

doc.add_heading('7-3. 정량 KPI (Phase별)', level=2)

add_table(doc,
    ['Phase', 'KPI', '목표'],
    [
        ['Phase 1', '활성 선주사 수 (Closed Beta)', '10~20개사 도달'],
        ['Phase 1', '월간 케이스 등록 건수',        '선박당 월 1건 이상'],
        ['Phase 2', 'SLA 응답 준수율',              '긴급 95% / 일반 90% 이상'],
        ['Phase 2', '전체 기존 선주사 중 활성 비율',  '70% 이상'],
        ['Phase 3', '무상 → 유상 전환율',            '활성 선주사의 30% 이상'],
        ['Phase 3', 'PM 하루 평균 메일 처리 시간 절감', '60% 이상 감소'],
        ['Phase 4', 'LCA 체결 선주 비중',             'AS 선주의 40% 이상'],
        ['Phase 4', 'NPS (선주 대상 만족도)',         '70 이상'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════
#  결론
# ════════════════════════════════════════════

doc.add_heading('결론', level=1)

doc.add_heading('1. 요약', level=2)

body(doc,
    '본 설계서는 **24/7 무상 기술 지원 웹사이트**를 5계층 아키텍처 · 5개 주 메뉴 · 4-Tier 응답 · 4-Phase 로드맵으로 정의. '
    '무상으로 접근성을 확보하고, 신뢰가 쌓이면 **PM 업무 개선을 목적으로 설계된 유상 서비스**가 기존 계약 구조 안에서 자연스럽게 확장')

bullet(doc, '**기술 구조** — 5계층 분리로 Application 계층에 유상 모듈 추가만으로 확장 가능')
bullet(doc, '**IA/기능** — 선주가 \"새로 배우지 않아도 되는\" 5개 메뉴 + Tier 에스컬레이션')
bullet(doc, '**PM 중심 효과** — 메일 처리·이력 조회·유상 요청을 모두 \"버튼 하나\"로')
bullet(doc, '**로드맵** — 기존 선주 Closed Beta → 24/7 확장 → 유상 → LCA 통합 포털')

doc.add_heading('2. 성공 조건', level=2)

add_table(doc,
    ['영역', '성공 조건'],
    [
        ['**고객 확보**',   'Phase 1에서 \"열성 선주 10개사\" 확보 — 체감 피드백으로 설계 보정'],
        ['**조직 정렬**',    'HMS 해외지사 4-Tier 체계 · SLA 준수 문화 정착'],
        ['**시스템 연계**', 'Hi-4S·SmartCare·ERP의 API 개방 · 내부 리소스 확보'],
        ['**계약 개편**',    '신규 AS/LCA 계약에 \"포털 편입\" 조항 기본 포함'],
        ['**데이터 거버넌스**', '선주 데이터 취급 원칙 공표 · 감사 체계 마련'],
    ])

doc.add_heading('3. 리스크 · 완화', level=2)

add_table(doc,
    ['리스크', '완화 방안'],
    [
        ['**오아시스 실패 이미지 재연**',
         'B2B 중개 플랫폼이 아닌 \"OEM 직통 기술 지원\"임을 명확히 커뮤니케이션'],
        ['**선주 정착률 저조**',
         '기존 메일·경보 접점에 웹 진입로 심기 (QR · 자동 링크) · 메일 공존 유지'],
        ['**HMS 내부 리소스 부족**',
         'MVP는 SaaS형 플랫폼(Salesforce Service Cloud 등) 활용 · 점진적 내재화'],
        ['**데이터 보안 우려**',
         'SSO · RBAC · 감사 로그 · 외부 침투 테스트로 신뢰 확보'],
        ['**유상 전환 저조**',
         '전환 자체가 아니라 \"PM 업무 개선\"이 목적 — 장기 관점 KPI로 평가'],
    ])

doc.add_heading('4. 다음 단계 (Immediate Next Steps)', level=2)

bullet(doc, '**Step 1** — Phase 1 MVP 상세 요구사항 명세서(PRD) 작성 · 기술 검토 회의 소집')
bullet(doc, '**Step 2** — Closed Beta 대상 선주 10곳 후보 선정 · 사전 인터뷰')
bullet(doc, '**Step 3** — 기술 스택 결정 (SaaS vs 자체 구축) · 연간 예산 추정')
bullet(doc, '**Step 4** — HMS 해외지사 담당 부서 합의 · 4-Tier 당번 체계 초안 수립')
bullet(doc, '**Step 5** — 6개월 후 MVP 런칭 · 3개월 운영 피드백 후 Phase 2 확정')

spacer(doc, 10)

insight(doc,
    '"메일함에 갇혀 있던 지식을 선박 단위 자산으로 바꾼다."\n'
    '\n'
    'PM의 하루를 바꾸는 도구가 곧 HMS의 장기 관계를 만드는 플랫폼입니다.\n'
    '무상으로 시작해 신뢰를 쌓고, 유상으로 자연스럽게 확장하는 단 하나의 창구 —\n'
    '\n'
    '이것이 본 설계서가 정의하는 "24/7 무상 기술 지원 웹사이트"입니다.')

doc.add_page_break()

# ════════════════════════════════════════════
#  참고문헌
# ════════════════════════════════════════════

doc.add_heading('참고문헌 (References)', level=1)

body(doc,
    '본 설계서의 아키텍처·IA·데이터 모델·보안 설계는 다음 공식 자료와 선도 OEM 사례를 참조')

spacer(doc, 6)

ref_categories = [
    ('1. OEM 선도기업 포털 · 원격 지원 시스템', [
        ('Wärtsilä Smart Support Centre — 24/7 글로벌 기술 지원',
         'https://www.wartsila.com/marine/products/data-service-downloads/smart-support-centre'),
        ('My Wärtsilä Portal — 선주·운영자용 통합 창구',
         'https://www.wartsila.com/marine/services/my-wartsila'),
        ('ABB Ability™ Remote Diagnostics and Predictive Maintenance',
         'https://new.abb.com/abb-ability/transport/marine/remote-diagnostics-and-predictive-maintenance'),
        ('ABB Integrated Operations Centers (IOC)',
         'https://new.abb.com/marine/systems-and-solutions/digital-solutions'),
        ('Kongsberg Maritime — Remote Support (24/7)',
         'https://www.kongsberg.com/maritime/services/kongsberg-remote-services/remote-support/'),
        ('Kongsberg K-IMS (Information Management System)',
         'https://www.kongsberg.com/maritime/products/digital/k_ims_applications/kongsberg-remote-services/'),
    ]),
    ('2. 정보 아키텍처 · UX 설계 레퍼런스', [
        ('Nielsen Norman Group — Information Architecture Principles',
         'https://www.nngroup.com/articles/ia-vs-navigation/'),
        ('WCAG 2.1 AA 접근성 가이드라인',
         'https://www.w3.org/TR/WCAG21/'),
        ('Material Design — Dashboard Layout Patterns',
         'https://m3.material.io/foundations/layout/applying-layout/dashboard'),
    ]),
    ('3. 보안 · 데이터 거버넌스', [
        ('OWASP ASVS 4.0 — Application Security Verification Standard',
         'https://owasp.org/www-project-application-security-verification-standard/'),
        ('ISO/IEC 27001:2022 — 정보보안 경영시스템',
         'https://www.iso.org/standard/27001'),
        ('OAuth 2.0 Framework (RFC 6749)',
         'https://datatracker.ietf.org/doc/html/rfc6749'),
    ]),
    ('4. HD현대마린솔루션 · 사내 자료', [
        ('HD현대마린솔루션 — 디지털 서비스 공식 페이지',
         'https://www.hd-marinesolution.com/kor/CMS/Contents/Contents.do?mCode=MN094'),
        ('HD현대마린솔루션 — 통합 디지털 솔루션 Hi-4S',
         'https://www.hd-marinesolution.com/kor/CMS/Contents/Contents.do?mCode=MN095'),
        ('HMS 테크 — 24/7 무상 기술지원 웹사이트 구축 제안 (2026.04)',
         '사내 자료 · 본 설계서의 기획 근거'),
        ('HMS 테크 — 원격 기술 지원 24/7 전문가 대기 체계 제안 (2026.04)',
         '사내 자료 · Tier 구조 참조'),
        ('HMS 테크 — 배전반 서비스 리포트 고도화 제안 (2026.04)',
         '사내 자료 · 도메인별 리포트 설계 참조'),
    ]),
    ('5. 시장 · 산업 동향', [
        ('Fortune Business Insights — Maritime Satellite Communication Market',
         'https://www.fortunebusinessinsights.com/maritime-satellite-communication-market-113315'),
        ('DNV — Maritime Forecast to 2050',
         'https://www.dnv.com/maritime/publications/maritime-forecast-2023/index.html'),
    ]),
]

for cat_title, refs in ref_categories:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(cat_title)
    run.font.size = Pt(11); run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    for i, (title, url) in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.left_indent = Cm(0.4)

        run_num = p.add_run(f'[{i}]  ')
        run_num.font.size = Pt(9); run_num.font.bold = True
        run_num.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        run_num.font.name = '맑은 고딕'
        run_num._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        run_title = p.add_run(title)
        run_title.font.size = Pt(9)
        run_title.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run_title.font.name = '맑은 고딕'
        run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        run_url = p.add_run(f'\n      {url}')
        run_url.font.size = Pt(8); run_url.font.italic = True
        run_url.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
        run_url.font.name = '맑은 고딕'
        run_url._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

spacer(doc, 10)

body(doc,
    '※ 모든 해외 자료는 영문 원문 기준. 사내 자료는 발표 자료집 또는 팀 공유 드라이브 참조. '
    '공개 URL은 접속 시점에 따라 변경될 수 있음')

doc.save(OUTPUT)
print(f"\n문서 생성 완료: {OUTPUT}")
