#!/usr/bin/env python3
"""배전반 서비스 리포트 고도화 제안 — 워드 생성"""
import os, re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE_DIR, "배전반 서비스 리포트 고도화 제안.docx")

PRIMARY = '1B3A5C'
SECONDARY = '2B579A'
LIGHT = 'F0F4F8'
LIGHT_BG = 'E8F0FE'


def init_doc(doc):
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)
    for level, (size, color) in enumerate([(20, PRIMARY), (16, SECONDARY), (13, '333333')], 1):
        h = doc.styles[f'Heading {level}']
        h.font.name = '맑은 고딕'
        h.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        h.font.size = Pt(size)
        h.font.bold = True
        r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
        h.font.color.rgb = RGBColor(r, g, b)
        h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def set_shading(cell, color):
    tcPr = cell._element.get_or_add_tcPr()
    s = tcPr.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    tcPr.append(s)


def set_text(cell, text, bold=False, color_rgb=None, size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(str(text))
    run.font.size = size
    run.font.bold = bold
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color_rgb:
        run.font.color.rgb = color_rgb


def keep_together(table):
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trPr.append(trPr.makeelement(qn('w:cantSplit'), {}))


def spacer(doc, pt=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(pt)
    p.add_run('').font.size = Pt(1)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_shading(cell, PRIMARY)
        set_text(cell, h, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                set_shading(cell, LIGHT)
            if isinstance(val, str) and val.startswith('**') and val.endswith('**'):
                set_text(cell, val[2:-2], bold=True, color_rgb=RGBColor(0x1B, 0x3A, 0x5C))
            else:
                set_text(cell, val)
    keep_together(table)
    spacer(doc, 12)
    return table


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if not part:
            continue
        run = p.add_run(part[2:-2] if part.startswith('**') else part)
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        if part.startswith('**') and part.endswith('**'):
            run.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        else:
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    p.clear()
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if not part:
            continue
        run = p.add_run(part[2:-2] if part.startswith('**') else part)
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        if part.startswith('**') and part.endswith('**'):
            run.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        else:
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def insight(doc, text, bg='E8F0FE', border='1B3A5C'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_shading(cell, bg)
    tcPr = cell._element.get_or_add_tcPr()
    borders = tcPr.makeelement(qn('w:tcBorders'), {})
    for side, w in [('left', '24'), ('top', '4'), ('bottom', '4'), ('right', '4')]:
        b = borders.makeelement(qn(f'w:{side}'), {
            qn('w:val'): 'single', qn('w:sz'): w,
            qn('w:color'): border if side == 'left' else 'D5D5D5',
            qn('w:space'): '0',
        })
        borders.append(b)
    tcPr.append(borders)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run('\n')
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.bold = True
        r, g, b = int(border[:2], 16), int(border[2:4], 16), int(border[4:], 16)
        run.font.color.rgb = RGBColor(r, g, b)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    keep_together(table)
    spacer(doc, 12)


def code_block(doc, text, bg='F5F5F5'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_shading(cell, bg)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    lines = text.rstrip().split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run('\n')
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.font.name = 'D2Coding'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    keep_together(table)
    spacer(doc, 12)


def divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)


# ═══════════════════════════════════════════════════════════════
doc = Document()
init_doc(doc)

# 표지
divider(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('배전반 서비스 리포트 고도화 제안')
run.font.size = Pt(24); run.font.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('리포트 시각화 + AI 자동 분석을 통한 고객 신뢰 강화')
run.font.size = Pt(13); run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
spacer(doc, 20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('2026. 04\nHD현대마린솔루션테크')
run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
divider(doc)

doc.add_page_break()

# 1. 현재 문제
doc.add_heading('1. 현재 문제', level=1)
doc.add_heading('지금 리포트의 한계', level=2)
add_table(doc, ['문제', '상세'], [
    ['**텍스트/숫자 나열**', '측정값이 표로만 나열되어 직관적으로 뭐가 문제인지 안 보임'],
    ['**시각화 없음**', '그래프, 차트, 다이어그램 없이 숫자만 나열'],
    ['**추이 분석 없음**', '"지금 이 수치다"만 있고, "과거 대비 어떻게 변했는지" 없음'],
    ['**향후 전망 없음**', '"현재 추세면 언제 위험해지는지" 예측이 없음'],
    ['**리포트 품질 편차**', '작성하는 엔지니어에 따라 품질이 크게 다름'],
    ['**작성 시간 오래 걸림**', '수동으로 데이터 정리 + 문서 작성에 반나절~하루 소요'],
])

doc.add_heading('왜 문제인가', level=2)
code_block(doc,
'''리포트 품질이 낮으면:
  → 선주가 "뭐가 문제인지 모르겠다"
  → 수리 결정이 느려짐
  → 수주가 늦어지거나 안 됨

리포트 품질이 높으면:
  → 선주가 "이건 빨리 고쳐야겠다"
  → 수리 결정이 빨라짐
  → 수주가 빨라짐''')
insight(doc, '즉, 리포트 품질 = 매출 속도')

doc.add_heading('선도기업과의 차이', level=2)
add_table(doc, ['항목', '선도기업 (바르질라/콩스버그)', '우리 (현재)'], [
    ['데이터 시각화', '그래프, 차트, 히트맵, 게이지', '표/텍스트'],
    ['추이 분석', '시간별 트렌드 그래프 자동 생성', '없음'],
    ['향후 전망', 'AI 기반 예측 트렌드라인', '없음'],
    ['리포트 생성', '자동/반자동', '100% 수동'],
    ['품질 일관성', '템플릿 기반 표준화', '작성자마다 다름'],
])

doc.add_page_break()

# 2. 해결 방안
doc.add_heading('2. 해결 방안 — 3단계', level=1)

doc.add_heading('Step 1. 리포트 템플릿 + 자동 시각화 (즉시 가능)', level=2)
body(doc, '**방법:** Excel/Google Sheets 기반 자동 시각화 템플릿')
body(doc, '**구조**')
code_block(doc,
'''시트 1: 데이터 입력 (엔지니어가 측정값 채움)
         ↓ 자동 연결
시트 2: 시각화 자동 생성 (차트, 그래프)
         ↓ PDF 출력
시트 3: 리포트 출력용 (선주 발송)''')

body(doc, '**자동 생성되는 시각화 항목**')
add_table(doc, ['측정 데이터', '시각화 형태', '보여주는 것'], [
    ['종합 건강도 점수', '**게이지 차트**', '100점 만점 중 현재 점수, 전월 대비 변화'],
    ['절연저항', '**꺾은선 그래프**', '최근 12개월 추이 + 정상/주의/위험 기준선'],
    ['차단기 동작시간', '**가로 막대 차트**', '정상 범위 vs 현재값 비교 (색상 구분)'],
    ['온도', '**히트맵 / 막대**', '부위별 온도, 과열 포인트 빨간색 강조'],
    ['부품 잔여 수명', '**수평 프로그레스 바**', '남은 수명 시각적 표시 + 교체 시기'],
    ['전체 평가', '**레이더 차트**', '절연/온도/동작시간/부식 등 항목별 점수'],
])
add_table(doc, ['항목', '내용'], [
    ['**난이도**', '쉬움 (Excel 스킬이면 됨)'],
    ['**소요 기간**', '1~2주'],
    ['**비용**', '0원'],
])

doc.add_heading('Step 2. AI 자동 분석 코멘트 (즉시 가능)', level=2)
body(doc, '**방법:** 측정값을 AI(ChatGPT/Claude)에 입력 → 분석 문구 자동 생성')
body(doc, '**입력 예시**')
code_block(doc,
'''선박: OO호
배전반: MSB #1
절연저항: 150MΩ (6개월 전 320MΩ, 12개월 전 480MΩ)
ACB-01 동작시간: 85ms (정상 60~90ms)
ACB-02 동작시간: 105ms (정상 60~90ms)
ACB-03 동작시간: 88ms (정상 60~90ms)
Bus-bar R상 온도: 65°C (정상 50°C 이하)''')

body(doc, '**AI 자동 생성 코멘트**')
code_block(doc,
'''■ 절연저항 분석
12개월간 480 → 320 → 150MΩ으로 지속 하락 중입니다.
월평균 약 27MΩ씩 감소하고 있으며, 현재 추세가 유지될 경우
약 5~6개월 후 위험 기준(50MΩ) 진입이 예상됩니다.
다음 입거 시 절연 정밀 점검을 권장합니다.

■ 차단기 동작시간 분석
ACB-02의 동작시간이 105ms로 정상 범위(60~90ms) 대비
17% 지연되어 있습니다. 기구부 열화가 추정되며,
미조치 시 트립 실패 위험이 있습니다.
교체를 권장합니다. (긴급도: 높음)

ACB-01(85ms), ACB-03(88ms)은 정상 범위 내입니다.

■ 온도 분석
Bus-bar R상 온도가 65°C로 정상(50°C 이하) 대비
15°C 높습니다. 접촉부 불량이 의심됩니다.
재조임 및 접촉 저항 측정을 권장합니다.

■ 종합 소견
배전반 종합 건강도 72/100 (전월 대비 -8점).
ACB-02 교체와 Bus-bar 접촉부 점검이 시급합니다.
절연저항은 현재 양호하나 하락 추세 모니터링이 필요합니다.''', bg='E8F0FE')

insight(doc, '엔지니어가 숫자만 입력 → AI가 전문가 수준의 분석 문구 자동 작성')
add_table(doc, ['항목', '내용'], [
    ['**난이도**', '매우 쉬움 (ChatGPT에 복붙이면 됨)'],
    ['**소요 기간**', '즉시'],
    ['**비용**', '0원 (또는 ChatGPT 구독료 월 $20)'],
])

doc.add_heading('Step 3. 통합 리포트 자동 생성 시스템 (중기)', level=2)
body(doc, '**방법:** 데이터 입력 → 시각화 + AI 분석 + PDF 리포트 자동 완성')
body(doc, '**구조**')
code_block(doc,
'''엔지니어가 앱/웹에서 측정값 입력
         ↓
시스템이 자동으로:
  ① 과거 데이터 DB에서 해당 선박 이력 조회
  ② 그래프/차트 자동 생성
  ③ AI가 분석 코멘트 작성
  ④ 향후 전망 (트렌드 기반 예측 그래프)
  ⑤ 권고 사항 + 견적 자동 첨부
  ⑥ PDF 리포트 완성
         ↓
PM이 검토 (수정 사항 있으면 수정)
         ↓
선주에게 발송''')

body(doc, '**구현 방법**')
add_table(doc, ['기술', '용도'], [
    ['**Power BI / Tableau**', '데이터 시각화 + 대시보드'],
    ['**Python + matplotlib**', '커스텀 그래프 자동 생성'],
    ['**GPT API / Claude API**', '분석 코멘트 자동 생성'],
    ['**자동화 도구 (Zapier 등)**', '입력 → 리포트 생성 → 발송 자동 연결'],
])
add_table(doc, ['항목', '내용'], [
    ['**난이도**', '중간 (개발 필요)'],
    ['**소요 기간**', '3~6개월'],
    ['**비용**', '중간 (Power BI 라이선스 + 개발비)'],
])

doc.add_page_break()

# 3. Before/After
doc.add_heading('3. 리포트 고도화 Before / After', level=1)

doc.add_heading('Before (현재)', level=2)
code_block(doc,
'''OO호 배전반 점검 결과

1. 절연저항
   MSB #1: 150MΩ
   ESB #1: 320MΩ

2. 차단기 동작시간
   ACB-01: 85ms
   ACB-02: 105ms
   ACB-03: 88ms

3. 온도 측정
   Bus-bar R상: 65°C
   Bus-bar S상: 48°C
   Bus-bar T상: 47°C

4. 소견
   ACB-02 교체 권장.
   절연저항 점검 필요.

→ 숫자만 나열, 그래프 없음
→ 과거 대비 비교 없음
→ 향후 전망 없음
→ 선주가 봐도 심각성 모름''', bg='FDEDEC')

doc.add_heading('After (고도화 후)', level=2)
code_block(doc,
'''OO호 배전반 정기점검 리포트
HD현대마린솔루션테크 | 2026-04
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

■ 종합 건강도
        ┌──────────────┐
        │    72 / 100  │  ← 게이지 차트
        │   (주의 필요)│
        └──────────────┘
   전월 80 → 이번달 72 (▼8점)

■ 절연저항 추이 (최근 12개월)
   MΩ
   500 ┤ ●
   400 ┤   ●  ●
   300 ┤        ●  ●
   200 ┤-----------●--●---- 주의 기준선
   100 ┤                 ●  ●
    50 ┤======================== 위험 기준선
       └─┬──┬──┬──┬──┬──┬──┬──→
        5월 7월 9월 11월 1월 3월

   ⚠ AI 분석: 월평균 27MΩ 하락 중.
   현재 추세 유지 시 5~6개월 후 위험 진입 예상.
   다음 입거 시 정밀 점검 권장.

■ 차단기 동작시간
   ACB-01  ████████████████░░░░  85ms  ✅ 정상
   ACB-02  ████████████████████████ 105ms ⚠
   ACB-03  █████████████████░░░░  88ms  ✅ 정상
           ├────────┤ 정상 범위 (60~90ms)

   ⚠ ACB-02 정상 대비 17% 지연.
   기구부 열화 추정. 교체 권장.

■ 부품 수명 예측
   ACB-01 ████████████████░░░░ 잔여 14개월
   ACB-02 ████░░░░░░░░░░░░░░░ 잔여 3개월 🔴
   ACB-03 █████████████████░░░ 잔여 16개월
   계전기  ████████████░░░░░░░ 잔여 10개월

■ 온도 측정
   R상  ████████████████████████ 65°C  🔴 과열
   S상  ██████████████████░░░░░░ 48°C  ✅ 정상
   T상  █████████████████░░░░░░░ 47°C  ✅ 정상
        기준: 50°C 이하

   ⚠ R상 접촉부 과열. 재조임 + 접촉저항 측정 권장.

■ 종합 권고 (AI 자동 생성)
   🔴 긴급: ACB-02 교체 (₩4.5M)
   🔴 긴급: Bus-bar R상 접촉부 점검
   🟡 권장: 절연저항 정밀 점검 (다음 입거)
   🟢 양호: ACB-01, 03 (모니터링 유지)
   📎 견적서 별첨

■ 향후 전망
   현재 추세 유지 시:
   • 3개월 후: ACB-02 트립 실패 가능성 높음
   • 6개월 후: 절연저항 위험 구간 진입 예상
   • 권장: 다음 입거(6개월 이내) 시 종합 정비''', bg='E8F0FE')

doc.add_page_break()

# 4. 기대 효과
doc.add_heading('4. 기대 효과', level=1)

doc.add_heading('리포트 품질 향상 → 매출 속도 향상', level=2)
add_table(doc, ['항목', 'Before', 'After'], [
    ['리포트 작성 시간', '반나절~하루', '**30분 (자동 생성)**'],
    ['시각화', '없음 (숫자 나열)', '**그래프, 차트, 게이지 자동**'],
    ['추이 분석', '없음', '**12개월 트렌드 자동**'],
    ['향후 전망', '없음', '**AI 예측 트렌드라인**'],
    ['권고 사항', '엔지니어 주관적 소견', '**데이터 기반 + AI 분석**'],
    ['리포트 품질', '작성자마다 다름', '**템플릿 기반 표준화**'],
    ['선주 의사결정', '느림 (뭐가 문제인지 모름)', '**빠름 (한눈에 보임)**'],
])

doc.add_heading('경쟁력 강화', level=2)
add_table(doc, ['효과', '설명'], [
    ['**선주 신뢰도 향상**', '"이 회사는 리포트가 프로페셔널하다" → 장기 계약 유도'],
    ['**수리 결정 속도 향상**', '시각화 → 심각성 직관적 이해 → 빠른 의사결정 → 수주 단축'],
    ['**차별화**', '경쟁 AS 업체 대비 리포트 품질로 차별화'],
    ['**데이터 축적**', '표준화된 데이터 → 향후 AI 분석/예지정비 기반'],
    ['**브랜드 가치**', 'HD현대마린솔루션테크 = "리포트도 다르다"'],
])

doc.add_heading('5. 구현 로드맵', level=1)
add_table(doc, ['Phase', '기간', '내용', '비용'], [
    ['**Step 1**', '1~2주', 'Excel 리포트 템플릿 (데이터 입력 → 차트 자동)', '0원'],
    ['**Step 2**', '즉시', 'AI 분석 코멘트 자동 생성 (ChatGPT/Claude 활용)', '0원 (또는 월 $20)'],
    ['**Step 3**', '1~3개월', 'Power BI 대시보드 (선박별 건강도 대시보드)', '낮음'],
    ['**Step 4**', '3~6개월', '통합 자동 리포트 시스템 (입력→시각화→분석→PDF)', '중간'],
])
insight(doc, 'Step 1, 2는 이번 주에도 가능합니다.')

doc.add_heading('6. 발표에서 활용하는 방법', level=1)

doc.add_heading('방법 1. 기존 발표에 추가', level=2)
body(doc, 'AM 벤치마킹 발표의 **제안 2 (부품/견적 플랫폼)**에 "리포트 고도화" 항목 추가')

doc.add_heading('방법 2. 별도 제안으로 분리', level=2)
bullet(doc, '제안 1. 고객 서비스 앱 (고객 접점)')
bullet(doc, '제안 2. 부품/견적 온라인 플랫폼 (매출 채널)')
bullet(doc, '제안 3. PM 업무 AI 자동화 (내부 효율화)')
bullet(doc, '**제안 4. 리포트 시각화 고도화 (고객 신뢰) ← 추가**')

doc.add_heading('방법 3. 실제 샘플 시연', level=2)
body(doc, '발표 때 **Before/After 리포트 샘플을 직접 보여주면** 임팩트 극대화.')
body(doc, '"이게 지금 리포트입니다" → "이게 고도화된 리포트입니다"')
body(doc, '→ 차이가 시각적으로 드러나니까 설득력 높음.')

doc.add_heading('발표 멘트 예시', level=2)
insight(doc,
        '"같은 데이터입니다. 같은 측정값입니다.\n'
        '하지만 보여주는 방식이 다르면, 선주의 반응이 다릅니다.\n'
        '숫자만 보면 \'나중에 하지\' 하던 선주가,\n'
        '그래프로 하락 추세를 보면 \'빨리 해야겠다\'로 바뀝니다.\n'
        '리포트 품질은 곧 수주 속도입니다."')

doc.add_heading('참고 — 활용 가능한 도구', level=1)
add_table(doc, ['도구', '용도', '비용'], [
    ['**Excel / Google Sheets**', '기본 템플릿 + 차트', '무료'],
    ['**ChatGPT / Claude**', 'AI 분석 코멘트 자동 생성', '무료~$20/월'],
    ['**Power BI**', '대시보드 + 자동 리포트', 'Microsoft 365 포함 or $10/월'],
    ['**Tableau**', '고급 시각화', '유료'],
    ['**Canva**', '리포트 디자인 템플릿', '무료~유료'],
    ['**Python (matplotlib/plotly)**', '커스텀 그래프 자동 생성', '무료'],
    ['**Figma**', '리포트 디자인 프로토타입', '무료'],
])

doc.save(OUTPUT)
print(f'✅ 저장 완료: {OUTPUT}')
