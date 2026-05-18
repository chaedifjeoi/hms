#!/usr/bin/env python3
"""
원격 기술 지원 — 선도기업 벤치마킹 발표 보고서
서론-본론-결론 구조 / 10분 발표용
"""

import os, re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

# ── 경로 ──
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, "report_images")
os.makedirs(IMG_DIR, exist_ok=True)
OUTPUT = os.path.join(BASE_DIR, "원격 기술 지원 — 선도기업 벤치마킹 보고서.docx")

# ── 컬러 (네이비/블루 기조) ──
PRIMARY = '1B3A5C'
SECONDARY = '2B579A'
LIGHT = 'F0F4F8'
LIGHT_BG = 'E8F0FE'
RED = 'E74C3C'
YELLOW = 'F39C12'
GREEN = '27AE60'

plt.rcParams['font.family'] = 'AppleGothic'
plt.rcParams['axes.unicode_minus'] = False
C_PRIMARY = '#1B3A5C'
C_SECONDARY = '#2B579A'
C_GRAY = '#BDC3C7'


# ══════════════════════════════════════
#  헬퍼 함수
# ══════════════════════════════════════

def init_doc(doc):
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)

    for level, (size, color) in enumerate([
        (20, PRIMARY), (16, SECONDARY), (13, '333333'),
    ], 1):
        h = doc.styles[f'Heading {level}']
        h.font.name = '맑은 고딕'
        h.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        h.font.size = Pt(size)
        h.font.bold = True
        r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
        h.font.color.rgb = RGBColor(r, g, b)
        h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def set_shading(cell, color):
    tcPr = cell._element.get_or_add_tcPr()
    s = tcPr.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    tcPr.append(s)


def set_text(cell, text, bold=False, color_rgb=None, size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(str(text))
    run.font.size = size
    run.font.bold = bold
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color_rgb:
        run.font.color.rgb = color_rgb


def keep_together(table):
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trPr.append(trPr.makeelement(qn('w:cantSplit'), {}))
    for row in table.rows[:-1]:
        for cell in row.cells:
            for p in cell.paragraphs:
                pPr = p._p.get_or_add_pPr()
                pPr.append(pPr.makeelement(qn('w:keepNext'), {}))
                pPr.append(pPr.makeelement(qn('w:keepLines'), {}))
    for cell in table.rows[-1].cells:
        for p in cell.paragraphs:
            pPr = p._p.get_or_add_pPr()
            pPr.append(pPr.makeelement(qn('w:keepLines'), {}))


def spacer(doc, pt=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(pt)
    run = p.add_run('')
    run.font.size = Pt(1)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_shading(cell, PRIMARY)
        set_text(cell, h, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                set_shading(cell, LIGHT)
            if isinstance(val, str) and val.startswith('**') and val.endswith('**'):
                set_text(cell, val[2:-2], bold=True)
            else:
                set_text(cell, val)
    keep_together(table)
    spacer(doc, 12)
    return table


def add_gap_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_shading(cell, PRIMARY)
        set_text(cell, h, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                set_shading(cell, LIGHT)
            if ci == len(headers) - 1:
                if '큼' in str(val):
                    set_shading(cell, 'FDEDEC')
                    set_text(cell, val, bold=True, color_rgb=RGBColor(0xE7, 0x4C, 0x3C))
                elif '중간' in str(val):
                    set_shading(cell, 'FEF9E7')
                    set_text(cell, val, bold=True, color_rgb=RGBColor(0xF3, 0x9C, 0x12))
                elif '양호' in str(val) or '우위' in str(val):
                    set_shading(cell, 'EAFAF1')
                    set_text(cell, val, bold=True, color_rgb=RGBColor(0x27, 0xAE, 0x60))
                else:
                    set_text(cell, val)
            else:
                set_text(cell, val)
    keep_together(table)
    spacer(doc, 12)
    return table


def add_ba_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['항목', 'Before (현재)', 'After (원격 기술 지원)']):
        cell = table.rows[0].cells[i]
        set_shading(cell, PRIMARY)
        set_text(cell, h, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, (item, before, after) in enumerate(rows):
        set_text(table.rows[ri + 1].cells[0], item, bold=True)
        c_b = table.rows[ri + 1].cells[1]
        set_shading(c_b, 'F5F5F5')
        set_text(c_b, before, color_rgb=RGBColor(0x99, 0x99, 0x99))
        c_a = table.rows[ri + 1].cells[2]
        set_shading(c_a, LIGHT_BG)
        set_text(c_a, after, bold=True, color_rgb=RGBColor(0x1B, 0x3A, 0x5C))
    keep_together(table)
    spacer(doc, 12)
    return table


def insight(doc, text, bg='E8F0FE', border='1B3A5C'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_shading(cell, bg)
    tcPr = cell._element.get_or_add_tcPr()
    borders = tcPr.makeelement(qn('w:tcBorders'), {})
    for side, w in [('left', '24'), ('top', '4'), ('bottom', '4'), ('right', '4')]:
        b = borders.makeelement(qn(f'w:{side}'), {
            qn('w:val'): 'single', qn('w:sz'): w,
            qn('w:color'): border if side == 'left' else 'D5D5D5',
            qn('w:space'): '0',
        })
        borders.append(b)
    tcPr.append(borders)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    r, g, b = int(border[:2], 16), int(border[2:4], 16), int(border[4:], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    keep_together(table)
    spacer(doc, 12)


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        else:
            run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    p.clear()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        else:
            run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def chart(doc, path, caption, width=Inches(5.8)):
    spacer(doc, 18)
    doc.add_picture(path, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(f'[그림] {caption}')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    spacer(doc, 18)


# ══════════════════════════════════════
#  차트 생성
# ══════════════════════════════════════

def make_effect_chart():
    fig, ax = plt.subplots(figsize=(7, 3.5))
    categories = ['엔지니어\n방문 횟수', '정비 비용', '다운타임', '대응 시간']
    before = [100, 100, 100, 100]
    after = [30, 50, 40, 15]
    x = np.arange(len(categories))
    w = 0.3
    ax.bar(x - w/2, before, w, label='Before (기존)', color=C_GRAY, zorder=3)
    ax.bar(x + w/2, after, w, label='After (원격 지원 도입)', color=C_PRIMARY, zorder=3)
    for i in range(len(categories)):
        red = before[i] - after[i]
        ax.text(x[i], 108, f'▼{red}%', ha='center', fontsize=9, fontweight='bold', color=C_PRIMARY)
    ax.set_ylabel('상대값 (%)', fontsize=9, color='#666666')
    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontsize=10, fontweight='bold')
    ax.set_ylim(0, 125)
    ax.legend(fontsize=9)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_axisbelow(True)
    plt.title('원격 기술 지원 도입 효과 (글로벌 선도기업 사례 기반)', fontsize=13,
              fontweight='bold', color=C_PRIMARY, pad=15)
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "remote_effect.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def make_gap_radar():
    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    cats = ['실시간\n원격 진단', '예측 정비\n(AI)', 'AR 원격\n가이드', '24/7\n관제센터',
            '원격 접속\n(제어)', '디지털\n트윈', '통합\n플랫폼']
    comp = [9, 9, 8, 9, 9, 8, 9]
    us = [4, 5, 1, 7, 3, 2, 4]
    N = len(cats)
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]
    comp += comp[:1]
    us += us[:1]
    ax.plot(angles, comp, 'o-', linewidth=2, color=C_GRAY, label='바르질라/콩스버그/ABB', markersize=5)
    ax.fill(angles, comp, alpha=0.05, color=C_GRAY)
    ax.plot(angles, us, 's-', linewidth=2.5, color=C_PRIMARY, label='HD현대마린솔루션테크', markersize=5)
    ax.fill(angles, us, alpha=0.1, color=C_PRIMARY)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(cats, fontsize=8, fontweight='bold')
    ax.set_ylim(0, 10)
    ax.set_yticks([2, 4, 6, 8, 10])
    ax.set_yticklabels(['2', '4', '6', '8', '10'], fontsize=7, color='gray')
    ax.grid(color='gray', linestyle='--', linewidth=0.5, alpha=0.5)
    ax.legend(loc='lower right', bbox_to_anchor=(1.15, -0.05), fontsize=9)
    plt.title('원격 기술 지원 Gap 분석', fontsize=12, fontweight='bold', color=C_PRIMARY, pad=20)
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "remote_gap.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


print("차트 생성 중...")
img_effect = make_effect_chart()
img_gap = make_gap_radar()
print("차트 생성 완료")

# ══════════════════════════════════════
#  문서 생성
# ══════════════════════════════════════

print("Word 문서 생성 중...")
doc = Document()
init_doc(doc)

# ────── 표지 ──────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(160)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
run = p.add_run('━' * 40)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run('선도기업 벤치마킹을 통한\n원격 기술 지원 도입 방향 제안')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
run = p.add_run('Remote Technical Support Benchmarking Report')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 40)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.font.size = Pt(10)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(140)

for t in ['2026. 04', 'HD현대마린솔루션테크']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(t)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_page_break()

# ════════════════════════════════════════════
#  서론 — 왜 원격 기술 지원인가 (약 2분)
# ════════════════════════════════════════════

doc.add_heading('서론 — 왜 원격 기술 지원인가', level=1)

doc.add_heading('배경', level=2)

body(doc,
    '글로벌 해양 산업은 **디지털 전환의 가속기**에 진입. '
    '선박의 기술적 문제 발생 시 엔지니어를 현장에 파견하는 기존 방식은 '
    '**시간·비용·인력** 세 가지 측면에서 한계에 도달')

bullet(doc, '엔지니어 파견 시 평균 **2~5일** 소요 (이동 + 대기 + 작업)')
bullet(doc, '긴급 상황에서도 **가용 인력 부족** → 대응 지연 → 선주 불만')
bullet(doc, '파견 1건당 **항공료 + 체재비 + 인건비** = 상당한 비용 발생')
bullet(doc, '친환경 선박 증가 → **신기술 문의 급증** → 기존 대응 방식으로 감당 불가')

spacer(doc, 6)

doc.add_heading('시장 동향', level=2)

body(doc,
    '관련 시장이 빠르게 성장 중. '
    '해양 위성통신 시장 **45억$(2025) → 118억$(2034)**, CAGR **11.4%**. '
    '해양 소프트웨어 시장 **18.2억$(2026) → 28.7억$(2035)**, CAGR **7.9%**')

body(doc,
    '선도 OEM(바르질라, 콩스버그, ABB)은 이미 원격 기술 지원을 **핵심 서비스 상품**으로 운영 중. '
    '단순 전화 상담이 아닌, **실시간 데이터 기반 원격 진단 + AR 가이드 + 예측 정비**를 '
    '통합한 디지털 서비스 체계를 구축')

insight(doc,
    '원격 기술 지원 = 단순 비용 절감 수단이 아닌,\n'
    '고객 경험을 근본적으로 바꾸는 "서비스 모델의 전환"')

doc.add_page_break()

# ════════════════════════════════════════════
#  본론 1 — 선도기업 벤치마킹 (약 4분)
# ════════════════════════════════════════════

doc.add_heading('본론 1 — 선도기업 벤치마킹', level=1)

# ── 바르질라 ──
doc.add_heading('바르질라 (Wärtsilä) — Expert Insight', level=2)

body(doc,
    '바르질라는 **Expert Insight** 서비스를 통해 원격 기술 지원의 최전선에 있음. '
    'AI + 룰 기반 진단을 결합하여 실시간 선박 데이터를 분석하고, '
    '**이상 신호 발생 전에 문제를 조기 감지**하는 예측 정비 체계를 운영')

add_table(doc,
    ['기능', '상세'],
    [
        ['**실시간 데이터 모니터링**', '선박 엔진/시스템의 실시간 데이터를 육상 센터에서 지속 분석'],
        ['**AI 기반 이상 감지**', '정상 운전 프로파일 대비 이탈 패턴을 AI가 자동 감지'],
        ['**온라인 협업**', '바르질라 전문가 ↔ 선박 승무원 간 웹 앱으로 실시간 채팅/케이스 관리'],
        ['**예측 정비 권고**', '데이터 분석 결과 기반 사전 정비 권고 → 계획 정비로 전환'],
    ])

body(doc,
    '2025년 일본 페리 운항사, LNG 운반선 14척 등 **Lifecycle Agreement**에 '
    'Expert Insight를 포함하는 사례가 지속 증가 중')

# ── ABB ──
doc.add_heading('ABB Marine — Remote Diagnostics', level=2)

body(doc,
    'ABB는 **원격 진단 서비스의 ROI를 수치로 증명**한 대표 사례')

add_table(doc,
    ['지표', '효과'],
    [
        ['**엔지니어 방문**', '**70% 감소** — 원격으로 대부분의 문제 진단·해결'],
        ['**정비 비용**', '**50% 절감** — 예측 모니터링으로 불필요한 정비 제거'],
        ['**모니터링 규모**', '**600척+** 선박을 24/7 원격 모니터링 중'],
        ['**서비스 구조**', '3단계 — Troubleshoot(진단) → Proactive(선제) → Predictive(예측)'],
    ])

body(doc,
    '**핵심:** 선상 소프트웨어 + 육상 분석 엔진을 결합한 구조. '
    '선박에서 데이터 수집 → 육상에서 분석 → 결과를 선박에 전달하는 **양방향 루프**')

# ── 콩스버그 ──
doc.add_heading('콩스버그 (Kongsberg) — Remote Virtual Service', level=2)

body(doc,
    '콩스버그는 **AR(증강현실) 기반 원격 가이드**에서 가장 앞서 있음. '
    '선박 승무원이 카메라를 켜면, 육상 엔지니어가 **실시간 화면에 마킹·주석·파일 공유**를 하며 '
    '작업을 지시하는 방식')

add_table(doc,
    ['기능', '상세'],
    [
        ['**라이브 비디오 스트리밍**', '선박 현장 ↔ 육상 엔지니어 간 실시간 영상 공유'],
        ['**AR 마킹/주석**', '화면 위에 화살표, 텍스트, 도형으로 작업 위치·순서 지시'],
        ['**파일 공유**', '매뉴얼, 도면, 체크리스트를 실시간으로 선박에 전송'],
        ['**위성통신 지원**', '대양 항해 중에도 위성 네트워크로 원격 지원 가능'],
        ['**원격 시스템 접속**', '육상 엔지니어가 선박 제어 시스템에 직접 원격 접속하여 진단'],
    ])

doc.add_page_break()

# ── 벤치마킹 핵심 요약 ──
doc.add_heading('벤치마킹 핵심 요약', level=2)

body(doc,
    '3사의 원격 기술 지원 전략을 관통하는 키워드 = **"현장에 가지 않고 해결한다"**')

body(doc,
    '바르질라는 **AI 예측 정비**, ABB는 **원격 진단 + 수치 증명**, '
    '콩스버그는 **AR 원격 가이드**로 각기 다른 강점을 보유. '
    '그러나 공통적으로 **실시간 데이터 분석 → 원격 진단 → 선제적 조치**라는 '
    '동일한 서비스 흐름을 구축하고 있음')

chart(doc, img_effect, '원격 기술 지원 도입 효과 — 글로벌 선도기업 실적 기반')

doc.add_page_break()

# ════════════════════════════════════════════
#  본론 2 — 우리 현황과 Gap (약 2분)
# ════════════════════════════════════════════

doc.add_heading('본론 2 — HD현대마린솔루션테크 현황과 Gap', level=1)

doc.add_heading('보유 역량', level=2)

body(doc,
    '원격 기술 지원의 기반이 될 수 있는 인프라는 **이미 상당 부분 보유**')

bullet(doc, '**Hi-4S** — 24/7 스마트케어 시스템, Fleet/Vessel 실시간 모니터링 가능')
bullet(doc, '**디지털 관제센터** — 350척+ 모니터링, 150명+ 전문가 협업 체제')
bullet(doc, '**SmartCare 부산 센터** — 2024.05 출범, 동남권 선박 관리 총괄, 24시간 운영')
bullet(doc, '**9,890척** AS 네트워크 — HD현대 건조 선박의 유일한 AS 사업자')

spacer(doc, 6)

doc.add_heading('Gap 분석', level=2)

add_gap_table(doc,
    ['영역', '선도기업 (바르질라/ABB/콩스버그)', 'HD현대마린솔루션테크', 'Gap'],
    [
        ['실시간 원격 진단', '데이터 기반 원격 진단 + 조치 권고', '전화/메일로 상황 파악', '큼'],
        ['예측 정비 (AI)', 'AI가 이상 징후 사전 감지', 'OceanWise 부분 적용', '중간'],
        ['AR 원격 가이드', 'AR 마킹 + 실시간 영상 지시', '미도입', '큼'],
        ['원격 시스템 접속', '육상에서 선박 시스템 직접 접속', '미도입', '큼'],
        ['24/7 관제센터', '전용 센터 + 글로벌 네트워크', 'SmartCare 센터 운영 중', '양호'],
        ['디지털 트윈', '선박 설비의 가상 복제 → 시뮬레이션', '미도입', '큼'],
        ['통합 서비스 플랫폼', '하나의 웹에서 원격 + 현장 + 이력 통합', '시스템 분산', '큼'],
    ])

chart(doc, img_gap, '원격 기술 지원 역량 Gap 분석', width=Inches(4.5))

insight(doc,
    '관제 인프라와 AS 네트워크는 이미 보유\n'
    '부족한 건 "현장에 가지 않고 문제를 해결하는" 원격 진단·가이드 기술',
    bg='FDEDEC', border=RED)

doc.add_page_break()

# ════════════════════════════════════════════
#  본론 3 — 도입 제안 (약 2분)
# ════════════════════════════════════════════

doc.add_heading('본론 3 — 원격 기술 지원 도입 제안', level=1)

doc.add_heading('도입 방향', level=2)

body(doc,
    '기존 Hi-4S/SmartCare 인프라를 활용하여 **3단계 원격 기술 지원 체계** 구축')

add_table(doc,
    ['단계', '내용', '기술 요소'],
    [
        ['**1단계: 원격 진단**', '배전반 상태 데이터를 실시간 수집 → 육상 전문가가 원격 분석', '센서 데이터 + 대시보드'],
        ['**2단계: AR 원격 가이드**', '현장 승무원 카메라 ↔ 육상 엔지니어 간 실시간 영상 지시', 'AR 마킹 + 화상 통화'],
        ['**3단계: 예측 정비**', '축적 데이터 기반 이상 징후 사전 감지 → 선제적 정비 권고', 'AI 분석 + 디지털 트윈'],
    ])

doc.add_heading('배전반 AS에 적용한 시나리오', level=2)

body(doc, '**시나리오: 선박 항해 중 배전반 차단기 트립 발생**')

spacer(doc, 4)

add_table(doc,
    ['단계', '기존 방식', '원격 기술 지원 도입 후'],
    [
        ['1. 문제 인지', '선박 → 전화/메일로 상황 설명', '센서가 트립 감지 → 관제센터 자동 알림'],
        ['2. 상황 파악', '구두 설명으로 파악 (오류 가능)', 'AR 영상으로 현장 실시간 확인'],
        ['3. 원인 분석', '엔지니어 파견 후 현장 확인', '원격 데이터 분석 + 이력 조회로 즉시 진단'],
        ['4. 조치 지시', '현장 도착 후 작업', 'AR 화면에 마킹하며 승무원에게 실시간 지시'],
        ['5. 소요 시간', '2~5일 (파견 포함)', '30분~2시간 (원격 완료)'],
    ])

doc.add_heading('기대 효과', level=2)

add_ba_table(doc, [
    ('문제 대응 시간', '2~5일 (엔지니어 파견)', '30분~2시간 (원격 즉시)'),
    ('엔지니어 파견', '매 건 현장 방문', '70% 원격 해결, 30%만 방문'),
    ('정비 비용', '긴급 파견 + 체재비', '최대 50% 절감'),
    ('선박 다운타임', '파견 대기 중 운항 제한', '원격 진단으로 대기 시간 최소화'),
    ('고객 만족도', '대응 지연 → 불만', '즉시 대응 → 신뢰 강화'),
    ('기술 이전', '현장에서만 가능', 'AR로 원격 교육/가이드 상시 가능'),
])

doc.add_page_break()

# ════════════════════════════════════════════
#  결론 (약 2분)
# ════════════════════════════════════════════

doc.add_heading('결론', level=1)

spacer(doc, 24)

insight(doc,
    '바르질라는 AI로 고장을 예측하고,\n'
    'ABB는 원격 진단으로 엔지니어 방문을 70% 줄였고,\n'
    '콩스버그는 AR로 현장에 가지 않고 수리를 지시합니다.\n'
    '\n'
    '우리는 아직 전화와 메일로 상황을 파악합니다.')

spacer(doc, 12)

body(doc,
    '그러나 **기반 인프라는 이미 갖추고 있음**')

bullet(doc, 'Hi-4S + 디지털 관제센터 + SmartCare 부산 = **원격 진단의 기반**')
bullet(doc, '**9,890척** 독점 AS 네트워크 = **데이터 수집의 기반**')
bullet(doc, '헬스체크 모델의 축적 데이터 = **예측 정비의 기반**')

spacer(doc, 12)

body(doc,
    '필요한 건 이 기반 위에 **원격 진단 + AR 가이드 + 예측 정비**를 올리는 것')

spacer(doc, 18)

insight(doc,
    '원격 기술 지원은 비용 절감 도구가 아닙니다.\n'
    '"전화 한 통이면 30분 안에 해결된다"는 경험이\n'
    '고객을 우리 서비스에 묶어두는 가장 강력한 무기가 됩니다.\n'
    '\n'
    '그 무기를 만드는 것이 이 제안의 핵심입니다.')

doc.add_page_break()

# ════════════════════════════════════════════
#  참고 자료
# ════════════════════════════════════════════

doc.add_heading('참고 자료', level=1)

refs = [
    ('Wärtsilä Expert Insight Service',
     'https://www.wartsila.com/marine/services/lifecycle-agreements/expert-insight-service'),
    ('Wärtsilä Smart Support Centre',
     'https://www.wartsila.com/marine/products/data-service-downloads/smart-support-centre'),
    ('Wärtsilä Lifecycle Agreement — Japanese Ferry Operator (2025)',
     'https://www.wartsila.com/are/media/news/04-11-2025-wartsila-lifecyle-agreement-selected-by-japanese-ferry-operator-to-support-service-reliability-3678494'),
    ('Riviera — Wärtsilä Expert Insight Takes Predictive Maintenance to the Next Level',
     'https://www.rivieramm.com/expert-views/expert-views/waumlrtsilauml-expert-insight-takes-predictive-maintenance-to-the-next-level'),
    ('ABB Ability™ Remote Diagnostics and Predictive Maintenance',
     'https://new.abb.com/abb-ability/transport/marine/remote-diagnostics-and-predictive-maintenance'),
    ('ABB — Remote Diagnostic Services Brochure',
     'https://library.e.abb.com/public/97e02350b7e6330bc1257c47004b1622/RDS%20Marine_Brochure%202014.pdf'),
    ('Maritime Executive — ABB Latest Update of Remote Diagnostic Services',
     'https://maritime-executive.com/corporate/abbs-latest-update-of-remote-diagnostic-services'),
    ('Kongsberg Maritime — Remote Virtual Service',
     'https://www.kongsberg.com/maritime/services/kongsberg-remote-services/remote-virtual-service/'),
    ('Kongsberg Maritime — Remote Support',
     'https://www.kongsberg.com/maritime/services/kongsberg-remote-services/remote-support/'),
    ('Kongsberg Maritime — KONGSBERG Remote Services (K-IMS)',
     'https://www.kongsberg.com/maritime/products/digital/k_ims_applications/kongsberg-remote-services/'),
    ('Fortune Business Insights — Maritime Satellite Communication Market (2025-2034)',
     'https://www.fortunebusinessinsights.com/maritime-satellite-communication-market-113315'),
    ('Market Reports World — Maritime Software Market Size & Share (2026-2035)',
     'https://www.marketreportsworld.com/market-reports/maritime-software-market-14725342'),
    ('HD현대마린솔루션 — 디지털 서비스',
     'https://www.hd-marinesolution.com/kor/CMS/Contents/Contents.do?mCode=MN094'),
    ('서울경제 — HD현대마린솔루션 부산 스마트케어 센터 설치',
     'https://www.sedaily.com/NewsView/2D9ACT14IL'),
]

for i, (title, url) in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2

    run_num = p.add_run(f'[{i}]  ')
    run_num.font.size = Pt(9)
    run_num.font.bold = True
    run_num.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run_num.font.name = '맑은 고딕'
    run_num._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    run_title = p.add_run(title)
    run_title.font.size = Pt(9)
    run_title.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run_title.font.name = '맑은 고딕'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    run_url = p.add_run(f'\n      {url}')
    run_url.font.size = Pt(8)
    run_url.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run_url.font.name = '맑은 고딕'
    run_url._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ── 저장 ──
doc.save(OUTPUT)
print(f"\n문서 생성 완료: {OUTPUT}")
