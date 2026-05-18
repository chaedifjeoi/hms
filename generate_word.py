#!/usr/bin/env python3
"""
선도기업 AM 전략 벤치마킹 보고서 — Word 문서 생성
가독성 가이드 + python-docx 레시피 기반
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

# ── 경로 설정 ──
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, "report_images")
os.makedirs(IMG_DIR, exist_ok=True)
OUTPUT_PATH = os.path.join(BASE_DIR, "선도기업 AM 전략 벤치마킹 보고서.docx")

# ══════════════════════════════════════════════
#  컬러 팔레트 — 4색 (네이비/블루 기조)
# ══════════════════════════════════════════════
PRIMARY = '1B3A5C'     # 네이비 — 제목, 표 헤더, 핵심 강조
SECONDARY = '2B579A'   # 블루 — 소제목, 그래프 보조
LIGHT_BG = 'E8F0FE'    # 연한 블루 배경
RED = 'E74C3C'
YELLOW = 'F39C12'
LIGHT = 'F0F4F8'
GRAY_BAR = 'BDC3C7'

# matplotlib
plt.rcParams['font.family'] = 'AppleGothic'
plt.rcParams['axes.unicode_minus'] = False

C_PRIMARY = '#1B3A5C'
C_SECONDARY = '#2B579A'
C_GRAY = '#BDC3C7'
C_LIGHT = '#F0F4F8'


# ══════════════════════════════════════════════
#  헬퍼 함수들 (레시피 기반)
# ══════════════════════════════════════════════

def init_document(doc):
    """레시피 1. 문서 초기 설정 — 폰트 + 여백 + 줄간격"""
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)

    for level, (size, color) in enumerate([
        (20, PRIMARY),
        (16, SECONDARY),
        (13, '333333'),
    ], 1):
        h = doc.styles[f'Heading {level}']
        h.font.name = '맑은 고딕'
        h.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        h.font.size = Pt(size)
        h.font.bold = True
        r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
        h.font.color.rgb = RGBColor(r, g, b)
        h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True  # 제목이 다음 요소와 같은 페이지에

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def set_cell_shading(cell, color_hex):
    """셀 배경색"""
    tcPr = cell._element.get_or_add_tcPr()
    shading = tcPr.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    tcPr.append(shading)


def set_cell_text(cell, text, bold=False, color_rgb=None, size=Pt(10),
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    """셀 텍스트 스타일"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(str(text))
    run.font.size = size
    run.font.bold = bold
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color_rgb:
        run.font.color.rgb = color_rgb


def add_spacer(doc, pt=12):
    """여백용 빈 단락 (높이 제어)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(pt)
    run = p.add_run('')
    run.font.size = Pt(1)


def keep_table_on_one_page(table):
    """표가 페이지 경계에서 잘리지 않도록 설정"""
    # 각 행이 페이지를 넘어가며 분할되지 않도록
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = trPr.makeelement(qn('w:cantSplit'), {})
        trPr.append(cantSplit)
    # 표 내 모든 단락에 keepNext + keepLines 설정
    for row in table.rows[:-1]:  # 마지막 행 제외
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                pPr = paragraph._p.get_or_add_pPr()
                keepNext = pPr.makeelement(qn('w:keepNext'), {})
                keepLines = pPr.makeelement(qn('w:keepLines'), {})
                pPr.append(keepNext)
                pPr.append(keepLines)
    # 마지막 행도 keepLines
    for cell in table.rows[-1].cells:
        for paragraph in cell.paragraphs:
            pPr = paragraph._p.get_or_add_pPr()
            keepLines = pPr.makeelement(qn('w:keepLines'), {})
            pPr.append(keepLines)


def add_styled_table(doc, headers, rows, col_widths=None):
    """레시피 3. 스타일링된 표"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, PRIMARY)
        set_cell_text(cell, header, bold=True,
                      color_rgb=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(10))

    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if row_idx % 2 == 1:
                set_cell_shading(cell, LIGHT)
            if isinstance(value, str) and value.startswith('**') and value.endswith('**'):
                set_cell_text(cell, value[2:-2], bold=True, size=Pt(10))
            else:
                set_cell_text(cell, value, size=Pt(10))

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width

    keep_table_on_one_page(table)
    add_spacer(doc, 12)
    return table


def add_gap_table(doc, headers, rows):
    """레시피 4. Gap 분석 표"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, PRIMARY)
        set_cell_text(cell, header, bold=True,
                      color_rgb=RGBColor(0xFF, 0xFF, 0xFF))

    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if row_idx % 2 == 1:
                set_cell_shading(cell, LIGHT)

            if col_idx == len(headers) - 1:
                if '큼' in str(value):
                    set_cell_shading(cell, 'FDEDEC')
                    set_cell_text(cell, value, bold=True,
                                  color_rgb=RGBColor(0xE7, 0x4C, 0x3C))
                elif '중간' in str(value):
                    set_cell_shading(cell, 'FEF9E7')
                    set_cell_text(cell, value, bold=True,
                                  color_rgb=RGBColor(0xF3, 0x9C, 0x12))
                elif '양호' in str(value) or '우위' in str(value):
                    set_cell_shading(cell, 'EAFAF1')
                    set_cell_text(cell, value, bold=True,
                                  color_rgb=RGBColor(0x27, 0xAE, 0x60))
                else:
                    set_cell_text(cell, value)
            else:
                set_cell_text(cell, value)

    keep_table_on_one_page(table)
    add_spacer(doc, 12)
    return table


def add_before_after_table(doc, rows):
    """레시피 6. Before/After 비교표"""
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['항목', 'Before (현재)', 'After (웹 포털)']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, PRIMARY)
        set_cell_text(cell, header, bold=True,
                      color_rgb=RGBColor(0xFF, 0xFF, 0xFF))

    for row_idx, (item, before, after) in enumerate(rows):
        set_cell_text(table.rows[row_idx + 1].cells[0], item, bold=True)

        cell_before = table.rows[row_idx + 1].cells[1]
        set_cell_shading(cell_before, 'F5F5F5')
        set_cell_text(cell_before, before, color_rgb=RGBColor(0x99, 0x99, 0x99))

        cell_after = table.rows[row_idx + 1].cells[2]
        set_cell_shading(cell_after, LIGHT_BG)
        set_cell_text(cell_after, after, bold=True,
                      color_rgb=RGBColor(0x1B, 0x3A, 0x5C))

    keep_table_on_one_page(table)
    add_spacer(doc, 12)
    return table


def add_insight_box(doc, text, bg_color='E8F0FE', border_color='1B3A5C'):
    """레시피 5. 왼쪽 보더 인사이트 박스"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)

    set_cell_shading(cell, bg_color)

    tcPr = cell._element.get_or_add_tcPr()
    borders = tcPr.makeelement(qn('w:tcBorders'), {})
    for side, width in [('left', '24'), ('top', '4'), ('bottom', '4'), ('right', '4')]:
        border = borders.makeelement(qn(f'w:{side}'), {
            qn('w:val'): 'single',
            qn('w:sz'): width,
            qn('w:color'): border_color if side == 'left' else 'D5D5D5',
            qn('w:space'): '0',
        })
        borders.append(border)
    tcPr.append(borders)

    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    r, g, b = int(border_color[:2], 16), int(border_color[2:4], 16), int(border_color[4:], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 인사이트 박스가 잘리지 않도록
    keep_table_on_one_page(table)
    add_spacer(doc, 12)


def add_cover(doc, title, subtitle, date, department):
    """레시피 2. 표지"""
    # 상단 여백
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(160)

    # 상단 장식선
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run('━' * 40)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.font.size = Pt(10)

    # 메인 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 부제
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(subtitle)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 하단 장식선
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.font.size = Pt(10)

    # 하단 여백
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(140)

    # 발표 정보
    for text in [date, department]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.add_page_break()


def add_body_text(doc, text):
    """본문 서술 텍스트"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3

    # **Bold** 패턴 처리
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        else:
            run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_bullet_text(doc, text):
    """불릿 포인트 텍스트"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2

    p.clear()
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        else:
            run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_chart(doc, img_path, caption_text, width=Inches(5.8)):
    """그래프 삽입 + 캡션 (전후 18pt 여백)"""
    add_spacer(doc, 18)  # 그래프 전 18pt
    doc.add_picture(img_path, width=width)
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 캡션
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f'[그림] {caption_text}')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    add_spacer(doc, 18)  # 그래프 후 18pt


# ══════════════════════════════════════════════
#  차트 생성 — 네이비/블루 2색 통일
# ══════════════════════════════════════════════

def create_market_chart():
    """시장 규모 성장 추이"""
    fig, ax = plt.subplots(figsize=(7, 3.5))

    years = [2024, 2026, 2028, 2030, 2032]
    market = [371, 406, 444, 486, 532]

    ax.plot(years, market, color=C_PRIMARY, linewidth=2.5,
            marker='o', markerfacecolor=C_SECONDARY, markersize=8, zorder=3)
    ax.fill_between(years, market, alpha=0.08, color=C_SECONDARY)

    # 시작/끝 주석
    for x, y, label in [(2024, 371, '371억$\n(55조원)'), (2032, 532, '532억$\n(78조원)')]:
        ax.annotate(label, xy=(x, y), xytext=(0, 25),
                    textcoords='offset points', fontsize=9, fontweight='bold',
                    color=C_PRIMARY, ha='center',
                    bbox=dict(boxstyle='round,pad=0.4', facecolor=C_LIGHT,
                              edgecolor=C_SECONDARY, alpha=0.9),
                    arrowprops=dict(arrowstyle='->', color=C_SECONDARY, lw=1.5))

    ax.set_title('글로벌 선박 수리·보수 시장 규모', fontsize=13,
                 fontweight='bold', color=C_PRIMARY, pad=15)
    ax.set_ylabel('억 달러', fontsize=9, color='#666666')
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    plt.tight_layout()
    path = os.path.join(IMG_DIR, "market.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def create_gap_radar():
    """Gap 분석 레이더 차트"""
    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))

    categories = ['고객용\n웹 포털', '서비스 요청\n온라인화', '호선정보\n고객 조회',
                   '서비스\n이력 조회', '기술 문서\n접근', '리포트\n시각화',
                   '통합\n플랫폼', 'AI/데이터\n분석', 'AS 네트워크\n/기술력']

    competitor = [9, 9, 9, 9, 9, 9, 8, 8, 8]
    hdms = [2, 2, 2, 3, 3, 2, 5, 7, 9]

    N = len(categories)
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]
    competitor += competitor[:1]
    hdms += hdms[:1]

    ax.plot(angles, competitor, 'o-', linewidth=2, color='#BDC3C7',
            label='바르질라/콩스버그', markersize=5)
    ax.fill(angles, competitor, alpha=0.05, color='#BDC3C7')
    ax.plot(angles, hdms, 's-', linewidth=2.5, color=C_PRIMARY,
            label='HD현대마린솔루션테크', markersize=5)
    ax.fill(angles, hdms, alpha=0.1, color=C_PRIMARY)

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=8, fontweight='bold')
    ax.set_ylim(0, 10)
    ax.set_yticks([2, 4, 6, 8, 10])
    ax.set_yticklabels(['2', '4', '6', '8', '10'], fontsize=7, color='gray')
    ax.grid(color='gray', linestyle='--', linewidth=0.5, alpha=0.5)
    ax.legend(loc='lower right', bbox_to_anchor=(1.15, -0.05), fontsize=9)

    plt.title('Gap 분석 — 선도기업 vs HD현대마린솔루션테크',
              fontsize=12, fontweight='bold', color=C_PRIMARY, pad=20)
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "gap_radar.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def create_roadmap():
    """로드맵 간트차트"""
    fig, ax = plt.subplots(figsize=(8, 3))

    phases = ['Phase 4: 확장', 'Phase 3: 중기', 'Phase 2: 단기', 'Phase 1: 즉시']
    starts = [12, 6, 3, 0]
    durations = [6, 6, 3, 3]
    # 진한 → 연한 네이비 그라데이션
    colors = ['#B0BEC5', '#5B9BD5', C_SECONDARY, C_PRIMARY]

    for i, (phase, start, dur, color) in enumerate(zip(phases, starts, durations, colors)):
        ax.barh(i, dur, left=start, height=0.5, color=color,
                edgecolor='white', linewidth=1.5, zorder=3)

    ax.set_yticks(range(len(phases)))
    ax.set_yticklabels(phases, fontsize=10, fontweight='bold')
    ax.set_xlabel('개월', fontsize=9, color='#666666')
    ax.set_xlim(-0.5, 20)
    ax.set_xticks([0, 3, 6, 9, 12, 15, 18])
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='x', alpha=0.3, linestyle='--')
    ax.set_axisbelow(True)

    plt.title('구현 로드맵', fontsize=13, fontweight='bold', color=C_PRIMARY, pad=15)
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "roadmap.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


# ══════════════════════════════════════════════
#  차트 생성
# ══════════════════════════════════════════════

print("차트 생성 중...")
img_market = create_market_chart()
img_radar = create_gap_radar()
img_roadmap = create_roadmap()
print("차트 생성 완료")

# ══════════════════════════════════════════════
#  문서 생성
# ══════════════════════════════════════════════

print("Word 문서 생성 중...")
doc = Document()
init_document(doc)

# ── 1. 표지 ──
add_cover(doc,
    title="선도기업 AM 전략 벤치마킹을 통한\nHD현대마린솔루션테크의 미래 방향 제안",
    subtitle="AM Strategy Benchmarking Report",
    date="2026. 04",
    department="HD현대마린솔루션테크")

# ── 2. Executive Summary (가이드: 1페이지로 핵심) ──
doc.add_heading('Executive Summary', level=1)

add_insight_box(doc,
    '선도기업(바르질라, 콩스버그)은 웹 포털에서 모든 AM 서비스를 제공\n'
    '→ HD현대마린솔루션테크는 기술력·네트워크 우위에도 디지털 고객 접점 부재\n'
    '→ 웹 포털 구축으로 Gap 해소 + 수주 속도 향상 + 고객 충성도 강화 가능')

add_styled_table(doc,
    ['구분', '현황', '제안'],
    [
        ['글로벌 AM 시장', '371억$(2024) → 532억$(2032), CAGR 4.6%', '성장 시장 선점 필요'],
        ['선도기업 수준', '웹 포털 기반 24/7 셀프서비스 완비', '바르질라/콩스버그 벤치마킹'],
        ['우리 현황', '전화/이메일 중심, 디지털 창구 부재', '**웹 포털 구축**'],
        ['AS 네트워크', '9,890척 독점 AS + 최고 기술력 보유', '강점을 디지털로 연결'],
    ])

doc.add_page_break()

# ── 3. 시장 현황 ──
doc.add_heading('1. AM 시장 현황', level=1)

doc.add_heading('시장 규모', level=2)

add_chart(doc, img_market, '글로벌 선박 수리·보수 시장 규모 전망 (CAGR 4.6%)')

add_styled_table(doc,
    ['시장', '2024년', '2032~33년', 'CAGR'],
    [
        ['선박 수리·보수', '371억$(55조원)', '532억$(78조원)', '4.6%'],
        ['해양 소프트웨어', '29.9억$', '78.4억$', '11.3%'],
    ])

doc.add_heading('핵심 트렌드', level=2)

add_body_text(doc,
    'AM 시장의 핵심 변화 = **디지털 전환**')

add_bullet_text(doc, '**디지털 전환** — 단순 AS → 웹 기반 플랫폼으로 고객 접점 전환')
add_bullet_text(doc, '**고객 셀프서비스** — 전화/메일 → 웹에서 서비스 요청·이력 조회·문서 열람')
add_bullet_text(doc, '**데이터 기반 서비스** — 수집 데이터 시각화 → 분석 리포트 제공')
add_bullet_text(doc, '**플랫폼 통합** — 흩어진 레거시 → 하나의 웹 플랫폼으로 통합')

add_spacer(doc, 6)

doc.add_page_break()

# ── 4. 벤치마킹 ──
doc.add_heading('2. 선도기업 벤치마킹', level=1)

doc.add_heading('바르질라 (Wärtsilä) — 웹 포털에서 모든 서비스를 처리', level=2)

add_styled_table(doc,
    ['기능', '상세'],
    [
        ['**서비스 요청 (TechRequest)**', '웹에서 기술 지원 요청 → 과거 요청 이력 + 답변 조회'],
        ['**기술 문서 열람**', '매뉴얼, 기술 회보 24/7 웹에서 열람'],
        ['**견적 요청/추적**', '온라인 견적 → 실시간 진행 상태 확인 → 문서 조회'],
        ['**클레임/워런티**', '요청 등록 → 처리 추적 → 전체 이력 조회'],
        ['**호선 정보 관리**', '고객이 자사 설비 정보를 웹에서 직접 관리'],
    ])

doc.add_heading('콩스버그 (Kongsberg) — 선박과 육상을 웹으로 연결', level=2)

add_styled_table(doc,
    ['기능', '상세'],
    [
        ['**K-IMS (웹 포털)**', '선박-육상 간 정보 관리 플랫폼, 브라우저 접속'],
        ['**정보 동기화**', '문서, 점검 이력, 기술 데이터를 선박↔육상 실시간 동기화'],
        ['**KM Performance (2026)**', '레거시 시스템을 하나의 웹 생태계로 통합'],
        ['**데이터 대시보드**', '항해 계획, 운항 모니터링, 트림 최적화 시각화'],
    ])

doc.add_heading('MAN + Rolls-Royce — 경쟁사끼리도 연결', level=2)

add_styled_table(doc,
    ['기능', '상세'],
    [
        ['**mya 플랫폼**', 'OEM 데이터 통합 오픈 플랫폼'],
        ['**전략**', '경쟁 OEM 데이터까지 단일 인터페이스로 통합 → 선주에게 하나의 창구'],
        ['**의미**', 'AM 시장에서 개방형 웹 생태계 자체가 경쟁력이 되는 시대'],
    ])

doc.add_heading('벤치마킹 요약 — 3사 공통점', level=2)

add_body_text(doc,
    '바르질라, 콩스버그, MAN+RR 3사의 AM 전략을 관통하는 키워드는 하나 — **"웹"**')

add_body_text(doc,
    '3사 모두 별도 앱 설치 없이 **브라우저만으로 모든 서비스에 접근**할 수 있는 환경을 구축했음. '
    '서비스 요청, 이력 조회, 기술 문서 열람까지 **고객이 직접 수행**하는 셀프서비스 체계로 전환 완료')

add_body_text(doc,
    '단순 데이터 나열이 아닌 **대시보드와 그래프로 시각화**하여 고객의 의사결정을 돕고, '
    '흩어져 있던 레거시 시스템들을 **하나의 통합 웹 플랫폼**으로 재편하는 것이 공통된 방향')

doc.add_page_break()

# ── 5. 현황 분석 ──
doc.add_heading('3. HD현대마린솔루션(테크) 현황 분석', level=1)

doc.add_heading('보유 역량', level=2)

add_body_text(doc,
    '선도기업 대비 기술력·네트워크는 이미 우위에 있음')

add_bullet_text(doc, '**9,890척** AS 네트워크 (2024 3Q 누적) — HD현대 건조 선박의 **유일한 AS 사업자**')
add_bullet_text(doc, '**Hi-4S** 24/7 스마트케어 + **디지털 관제센터** 350척+ 모니터링, 150명+ 전문가')
add_bullet_text(doc, '**ISS 2.0** 통합 스마트십 + **OceanWise** AI 항로 최적화 (연료 **5.3%** 절감)')
add_bullet_text(doc, '**헬스체크 모델** — 무상 점검 → 리포트 → 수리 권고 → 수주 (검증된 비즈니스 모델)')

add_spacer(doc, 6)

doc.add_heading('Gap 분석', level=2)

add_gap_table(doc,
    ['영역', '바르질라/콩스버그', 'HD현대마린솔루션(테크)', 'Gap'],
    [
        ['고객용 웹 포털', 'Wärtsilä Online, K-IMS', '없음 (전화/이메일)', '큼'],
        ['서비스 요청 온라인화', '웹에서 접수→추적→이력', '전화/메일 접수', '큼'],
        ['호선정보 고객 조회', '웹 포털에서 즉시 조회', '내부 시스템 (고객 접근 불가)', '큼'],
        ['서비스 이력 조회', '고객이 직접 조회', '요청 시 제공', '큼'],
        ['기술 문서 접근', '24/7 웹 셀프서비스', '요청 시 제공', '큼'],
        ['리포트 시각화', '대시보드 + 그래프', '텍스트/숫자 나열', '큼'],
        ['통합 플랫폼', '레거시 통합 완료/진행', 'Hi-4S + ISS 2.0 (분리)', '중간'],
        ['AI/데이터 분석', '예지정비, 성능 최적화', 'OceanWise, ISS 2.0', '양호'],
        ['AS 네트워크/기술력', '글로벌 네트워크', '9,890척 독점 AS', '우위'],
    ])

add_chart(doc, img_radar, 'Gap 분석 — 선도기업 vs HD현대마린솔루션테크', width=Inches(4.5))

add_insight_box(doc,
    '기술력과 네트워크는 선도기업 이상\n'
    '부족한 건 고객이 우리 서비스에 접근하는 "디지털 창구"',
    bg_color='FDEDEC', border_color=RED)

doc.add_page_break()

# ── 6. 제안 ──
doc.add_heading('4. 제안 — 웹 포털 구축', level=1)

doc.add_heading('웹 포털 핵심 기능', level=2)

add_body_text(doc,
    'Gap 분석 기반, 웹 포털 핵심 기능을 **6개 영역**으로 설계. '
    '핵심 목표 = 고객이 **전화/메일 없이 스스로 서비스에 접근**할 수 있는 환경 구축')

features = [
    ('A. 호선 정보 조회', [
        ['선박 검색', '선박명/IMO 번호로 검색 → 선박 정보 페이지'],
        ['배전반 스펙 조회', '설치된 배전반 종류, 모델, 설치 연도, 회로도'],
        ['기자재 목록', '해당 선박에 설치된 담당 기자재 전체 리스트'],
    ]),
    ('B. 서비스 요청 + 진행 추적', [
        ['온라인 서비스 요청', '웹에서 작업 요청 접수 (선박 선택 → 내용 입력 → 첨부)'],
        ['자동 PM 배정', '요청 접수 → 담당 PM에게 자동 배정 + 알림'],
        ['진행 현황 추적', '접수→견적→승인→공사→완료 단계별 실시간 표시'],
    ]),
    ('C. 점검/수리 이력 관리', [
        ['헬스체크 이력', '과거 방문 일자, 점검 항목, 결과 요약'],
        ['수리 이력', '건별 작업 내용, 교체 부품, 비용'],
        ['타임라인 뷰', '선박별 전체 서비스 히스토리를 시간순 표시'],
    ]),
    ('D. 리포트 시각화 대시보드', [
        ['건강도 점수', '종합 점수 게이지 차트 (100점 만점)'],
        ['측정값 추이', '절연저항, 차단기 동작시간, 온도 등 트렌드 그래프'],
        ['부품 수명 예측', '잔여 수명 프로그레스 바 + 교체 권장 시기'],
    ]),
    ('E. 기술 문서 라이브러리', [
        ['매뉴얼', '배전반 운용/정비 매뉴얼 호선별 제공'],
        ['단선도/복선도', '해당 선박 배전반 도면 열람/다운로드'],
        ['기술 회보', 'Technical Bulletin (안전 공지, 업데이트 정보)'],
    ]),
    ('F. 알림/리마인더', [
        ['정기 점검 알림', '마지막 점검 1년 경과 → 헬스체크 권장'],
        ['부품 교체 알림', 'ACB 설치 후 5년 경과 → 점검/교체 검토 권장'],
        ['서비스 진행 알림', '견적 발송, PO 수신, 공사 완료 등 단계 변경 시 알림'],
    ]),
]

for title, rows in features:
    doc.add_heading(title, level=3)
    add_styled_table(doc, ['기능', '설명'], rows)

doc.add_page_break()

# ── 7. 기대 효과 ──
doc.add_heading('5. 기대 효과 (고객/회사)', level=1)

doc.add_heading('고객 관점 — Before / After', level=2)

add_before_after_table(doc, [
    ('배전반 스펙 확인', '전화 문의', '웹에서 즉시 검색'),
    ('서비스 요청', '전화/메일', '웹 24/7 접수'),
    ('진행 현황', '전화 문의', '웹에서 실시간 추적'),
    ('과거 이력', '요청 후 대기', '웹에서 직접 조회'),
    ('기술 문서', '요청 후 수신', '웹 24/7 열람'),
    ('리포트', '텍스트 나열', '그래프/대시보드'),
    ('점검 시기', '고객이 관리', '자동 알림'),
])

doc.add_heading('회사 관점', level=2)

add_bullet_text(doc, '**수주 속도 향상** — 시각화 리포트 → 선주 빠른 의사결정 → 수주 단축')
add_bullet_text(doc, '**업무 부하 감소** — "스펙 알려주세요" 류 반복 전화 대폭 감소')
add_bullet_text(doc, '**선제적 영업** — 알림/리마인더로 우리가 먼저 접근 → 수주 기회 확대')
add_bullet_text(doc, '**데이터 축적** — 고객 행동 데이터 → 맞춤형 서비스/마케팅 기반 확보')
add_bullet_text(doc, '**인수인계 용이** — 모든 이력 시스템화 → PM 교체 시 맥락 유실 없음')

doc.add_page_break()

# ── 7. 로드맵 ──
doc.add_heading('6. 구현 로드맵', level=1)

add_chart(doc, img_roadmap, '구현 로드맵 타임라인')

add_styled_table(doc,
    ['Phase', '기간', '내용', '투자 규모'],
    [
        ['**Phase 1: 즉시**', '1~3개월', '리포트 시각화 템플릿 + 내부 프로세스 정비', '낮음'],
        ['**Phase 2: 단기**', '3~6개월', '웹 포털 MVP (호선정보, 서비스 요청, 이력 관리)', '중간'],
        ['**Phase 3: 중기**', '6~12개월', '대시보드 + 기술 문서 라이브러리 + 알림 시스템', '중간'],
        ['**Phase 4: 확장**', '12개월~', '모회사 Hi-4S/ISS 2.0 연동 + 전사 확대', '그룹 협의'],
    ])

add_body_text(doc,
    '**Phase 1부터 시작하는 이유:**')

add_bullet_text(doc, '추가 투자 거의 없이 **내일부터 착수 가능**')
add_bullet_text(doc, '내부 효율화 효과 **먼저 체감**')
add_bullet_text(doc, 'Phase 2 웹 포털 구축을 위한 **데이터/프로세스 기반** 마련')

doc.add_page_break()

# ══════════════════════════════════════════════════
#  7. 기대 효과 종합 (기대효과_상세.md 기반)
# ══════════════════════════════════════════════════

doc.add_heading('7. 선도기업은 이미 성과를 내고 있다', level=1)

# ── 바르질라 실적 ──
doc.add_heading('바르질라 — 디지털 전환 후 실적', level=2)

add_styled_table(doc,
    ['지표', '2023년', '2024년', '성장률'],
    [
        ['**총 매출**', '€60.2억', '€64.5억', '+7.2%'],
        ['**서비스 매출**', '~€30억', '~€34억', '+12%'],
        ['**서비스 매출 비중**', '50%', '53%', '↑'],
        ['**영업이익**', '€5.0억', '€6.9억', '+39%'],
        ['**서비스 수주**', '—', '—', '+15%'],
    ])

add_body_text(doc,
    '바르질라 전체 매출의 **53%가 서비스(AM)**에서 발생. '
    '그 서비스의 핵심 도구 = **Wärtsilä Online 웹 포털 + 라이프사이클 계약**')

add_body_text(doc,
    '2024~2026년 **2년 만에 80척+ 장기 계약 체결** (5~16년). '
    '라이프사이클 계약 갱신율 **90% 이상** — 한 번 들어오면 나가지 않는 구조')

add_body_text(doc,
    '원격 지원 성과: Smart Support Centre 기준 **원격 문제 해결률 95%**, '
    '데이터 기반 정비 비용 **14% 절감**')

# ── 콩스버그 실적 ──
doc.add_heading('콩스버그 — 디지털 통합으로 급성장', level=2)

add_styled_table(doc,
    ['지표', '2023년', '2024년', '성장률'],
    [
        ['**해양사업 매출**', '~NOK 200억', '~NOK 240억', '+20~30%'],
        ['**디지털 반복 매출**', 'NOK 2.1억(Q4)', 'NOK 2.4억(Q4)', '+16.5%'],
        ['**그룹 영업이익**', 'NOK 46억', 'NOK 65억', '+41%'],
        ['**수주잔고**', 'NOK 886억', 'NOK 1,279억', '+44%'],
    ])

add_body_text(doc,
    '**디지털 반복 매출이 분기 NOK 2.4억(약 240억 원)**에 달함. '
    '소프트웨어는 한 번 만들면 추가 비용 없이 반복 과금되는 구조')

doc.add_page_break()

# ── 우리 회사 기대효과 ──
doc.add_heading('8. 우리 회사에 돌아오는 기대효과', level=1)

# 매출 구조 전환
doc.add_heading('매출 구조 전환 — 건별 매출 → 안정적 반복 매출', level=2)

add_body_text(doc,
    '**현재:** 고장 발생 → 연락 → 수리 → 청구 → 끝 → (다음 고장까지 매출 0). '
    '매출 예측 불가, 건마다 가격 경쟁에 노출')

add_body_text(doc,
    '**도입 후:** 웹 포털에서 고객이 매일 접속 → 데이터 축적 → 선제적 서비스 제안 → '
    '**라이프사이클 계약 체결** → 연간 고정 요금 → 안정적 반복 매출')

add_styled_table(doc,
    ['Tier', '선박당 연간', '적용 척수', '연간 매출'],
    [
        ['**Tier 1 (Basic)**', '₩10M', '50척', '₩5억'],
        ['**Tier 2 (Standard)**', '₩20M', '50척', '₩10억'],
        ['**Tier 3 (Premium)**', '₩50M', '20척', '₩10억'],
        ['**합계 (3년차 목표)**', '—', '120척', '**₩25억 + 추가 수주 ₩10억 = ₩35억**'],
    ])

add_body_text(doc,
    '**120척 = 9,890척 네트워크의 1.2%에 불과.** '
    '5%(~500척)만 계약해도 → **연 ₩100억+ 반복 매출**')

# 고객 Lock-in
doc.add_heading('고객 Lock-in — 한 번 들어오면 나가지 않는 구조', level=2)

add_body_text(doc,
    '웹 포털에 **3~5년치 서비스 이력, 측정 데이터, 리포트**가 축적되면 '
    '이 데이터는 경쟁사에 없음 → 고객이 전환하면 모든 이력을 잃음 → **전환 비용(Switching Cost) 상승** → 이탈 방지')

add_styled_table(doc,
    ['Lock-in 요소', '전환 시 고객이 잃는 것'],
    [
        ['**서비스 이력**', '3~5년치 점검/수리 이력 전부'],
        ['**측정 데이터**', '절연저항, 차단기 동작시간 등 트렌드 분석 기반'],
        ['**리포트 아카이브**', '시각화 리포트 이력, 과거 건강도 비교 불가'],
        ['**부품 수명 추적**', '각 부품 잔여 수명 데이터, 교체 시기 예측 불가'],
        ['**알림 설정**', '맞춤형 알림 체계 → 수동 관리로 복귀'],
    ])

# 수주 속도 향상
doc.add_heading('수주 속도 향상 — 리포트가 영업을 한다', level=2)

add_body_text(doc,
    '**현재:** "절연저항 5.2MΩ" 숫자 나열 → 선주가 심각성 미체감 → 결정까지 3~5주')

add_body_text(doc,
    '**시각화 후:** 게이지 차트에 **빨간 구간 바늘** → 한눈에 "위험" 인지 → '
    '꺾은선 그래프로 **3개월 연속 하락 추이** 확인 → 빠른 결재 → **수주 1~2주로 단축 (50~70%↓)**')

# PM 업무 부하
doc.add_heading('PM 업무 부하 감소', level=2)

add_styled_table(doc,
    ['문의 유형', '현재', '웹 포털 후', '절감'],
    [
        ['배전반 모델 문의', '주 5~10건, 건당 10~20분', '고객이 웹에서 직접 검색', '주 2~3시간'],
        ['과거 수리 이력', '주 3~5건, 건당 15~30분', '고객이 웹에서 직접 조회', '주 1~2시간'],
        ['진행 상황 문의', '주 10~15건, 건당 5~10분', '웹에서 실시간 추적', '주 2~3시간'],
        ['도면/문서 요청', '주 3~5건, 건당 10~15분', '웹 라이브러리 24/7', '주 1~2시간'],
    ])

add_body_text(doc,
    'PM 1인당 **주 10~15시간 절감**, 5명 기준 연간 **2,400~3,600시간** → '
    '이 시간을 **영업·고객 관계·기술 자문** 등 고부가가치 업무에 투입')

# 선제적 영업
doc.add_heading('선제적 영업 — 고객이 부르기 전에 우리가 먼저', level=2)

add_body_text(doc,
    '**현재:** 고객이 고장 나서 전화해야 AS 시작. 고장 없으면 접점 없음')

add_body_text(doc,
    '**도입 후:** 웹 포털에 축적된 데이터를 자동 분석 → '
    '"이 선박 MSB #1 절연저항 3개월 연속 하락" 감지 → '
    '**자동 알림** → PM이 먼저 접근 → 경쟁사보다 먼저 수주 선점')

# 그룹 전략 정합
doc.add_heading('그룹 AX 전략 정합', level=2)

add_body_text(doc,
    '**Hi-4S**는 엔진/주요기기를, **ISS 2.0**은 통합 스마트십 데이터를 커버. '
    '**배전반은 빈 영역** — 우리가 웹 포털로 이 영역을 채우면 '
    '그룹 차원의 **"선박 토탈 케어"** 완성')

doc.add_page_break()

# ── 정량적 기대효과 종합표 ──
doc.add_heading('9. 기대 효과 종합 정리', level=1)

add_before_after_table(doc, [
    ('매출 구조', '건별 불규칙', '라이프사이클 계약 — 연 ₩25~35억 안정 매출'),
    ('서비스 요청', '전화/이메일 (업무시간)', '웹 포털 24/7 즉시 접수'),
    ('호선정보 확인', '전화→조회→회신 (10~20분)', '웹 즉시 검색 (3초)'),
    ('수주 기간', '리포트→결정 3~5주', '대시보드→결정 1~2주 (50~70%↓)'),
    ('PM 반복 업무', '전화 대응에 주 20시간+', '웹 셀프서비스 → 주 10~15시간 절감'),
    ('고객 이탈률', '건별 거래 → 이탈 가능', '데이터 Lock-in → 갱신율 90% 목표'),
    ('리포트 품질', '숫자 나열', '그래프/대시보드 시각화'),
])

doc.add_heading('전략적 효과', level=2)

add_styled_table(doc,
    ['전략적 효과', '설명', '벤치마크'],
    [
        ['**선도기업 Gap 해소**', '바르질라/콩스버그 수준 디지털 고객 경험', '바르질라: 서비스 매출 €34억, 53%'],
        ['**고객 Lock-in**', '데이터 축적 + 장기 계약 → 전환 비용 상승', '바르질라: 갱신율 90%+'],
        ['**선제적 영업 전환**', '수동 대기 → 데이터 기반 자동 영업', 'Expert Insight 벤치마킹'],
        ['**서비스 브랜드화**', '점검 이력 인증 → 중고선 가치 향상', '업계 최초 차별화 요소'],
        ['**그룹 AX 정합**', 'Hi-4S/ISS 2.0 보완 → 선박 토탈 케어', 'HD현대 그룹 전략 방향'],
    ])

doc.add_page_break()

# ── 11. 마무리 ──
doc.add_heading('마무리', level=1)

add_spacer(doc, 36)

add_insight_box(doc,
    '바르질라는 웹 포털에서 서비스를 관리하고,\n'
    '콩스버그는 웹 플랫폼으로 선박과 육상을 연결합니다.\n'
    '\n'
    '우리 고객은 아직 전화와 이메일로 연락합니다.')

add_spacer(doc, 18)

add_insight_box(doc,
    '선도기업과의 차이는 기술력이 아닙니다.\n'
    '우리는 9,890척의 AS 네트워크와 최고의 기술력을 이미 갖고 있습니다.\n'
    '부족한 건 고객이 그 서비스에 접근하는 디지털 창구입니다.\n'
    '\n'
    '웹 포털 하나가 그 창구가 됩니다.',
    bg_color=LIGHT_BG, border_color=PRIMARY)

doc.add_page_break()

# ── 참고 자료 ──
doc.add_heading('참고 자료', level=1)

refs = [
    ('Wärtsilä Online — 고객 웹 포털',
     'https://www.wartsila.com/wartsila-online'),
    ('Wärtsilä Online — TechRequest / 서비스 요청',
     'https://www.wartsila.com/marine/services/technical-support'),
    ('Wärtsilä Online — 기술 문서 / Digital Commerce Solutions',
     'https://www.wartsila.com/wartsila-online/digital-commerce-solutions'),
    ('Wärtsilä — 2026 해양 트렌드',
     'https://www.wartsila.com/insights/article/from-big-data-to-lifecycle-optimisation-4-trends-that-will-affect-shipping-in-2026'),
    ('Kongsberg K-IMS — 웹 기반 정보 관리 시스템',
     'https://www.kongsberg.com/maritime/products/digital/k_ims_applications/k-ims/'),
    ('Kongsberg — KM Performance 플랫폼 (2026)',
     'https://worldoil.com/news/2026/2/13/kongsberg-maritime-launches-unified-digital-portfolio-introduces-km-performance-platform/'),
    ('Kongsberg — 통합 디지털 솔루션 포트폴리오 발표',
     'https://www.kongsberg.com/maritime/news-and-events/news-archive/2026/harmonised-digital-solutions-portfolio-and-km-performance-launched/'),
    ('MAN Energy Solutions — PrimeServ / Marine Services',
     'https://www.man-es.com/services/industries/marine'),
    ('HD현대마린솔루션 — 디지털 서비스',
     'https://www.hd-marinesolution.com/kor/CMS/Contents/Contents.do?mCode=MN094'),
    ('서울경제 — HD현대마린솔루션 부산 스마트케어 센터 설치',
     'https://www.sedaily.com/NewsView/2D9ACT14IL'),
    ('Wärtsilä Financial Statements 2024 — 매출 €64.5억, 서비스 53%, 영업이익 +39%',
     'https://www.wartsila.com/media/news/05-02-2025-wartsila-s-financial-statements-bulletin-january-december-2024-3545975'),
    ('Wärtsilä Smart Support Centre — 원격 해결률 95%',
     'https://www.wartsila.com/marine/products/data-service-downloads/smart-support-centre'),
    ('Wärtsilä Data-driven Maintenance Planning — 유지보수 비용 14% 절감',
     'https://www.wartsila.com/marine/services/lifecycle-agreements/data-driven-maintenance-planning-service'),
    ('Wärtsilä Seapeak 10척 16년 계약 (2024)',
     'https://www.wartsila.com/media/news/30-07-2024-renewed-wartsila-lifecycle-agreement-will-enhance-performance-and-maintenance-of-ten-seapeak-lng-carriers-3478612'),
    ('Wärtsilä OPearl LNG 14척 10년 계약 (2025)',
     'https://www.wartsila.com/media/news/30-09-2025-wartsila-lifecycle-agreement-will-provide-support-to-14-lng-carriers-for-maintaining-operational-reliability-3664000'),
    ('Kongsberg Q4 2024 Results — 디지털 매출 +17%, 반복 매출 +16.5%',
     'https://www.kongsberg.com/maritime/news-and-events/news-archive/2025/financial-results-q4-2024/'),
    ('Kongsberg 2024 Annual Report — 수주잔고 +44%, EBIT +41%',
     'https://www.kongsberg.com/newsroom/news-archive/2025/annual-report-2024-outlook/'),
    ('Cognitive Market Research — Maritime Digitization Market $1,650억→$4,900억',
     'https://www.cognitivemarketresearch.com/maritime-digitization-market-report'),
    ('IMARC Group — Marine Telematics Market Size & Trends (2025-2033)',
     'https://www.imarcgroup.com/marine-telematics-market'),
    ('Fortune Business Insights — Maritime Satellite Communication Market (2025-2034)',
     'https://www.fortunebusinessinsights.com/maritime-satellite-communication-market-113315'),
]

for i, (title, url) in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2

    run_num = p.add_run(f'[{i}]  ')
    run_num.font.size = Pt(9)
    run_num.font.bold = True
    run_num.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run_num.font.name = '맑은 고딕'
    run_num._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    run_title = p.add_run(title)
    run_title.font.size = Pt(9)
    run_title.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run_title.font.name = '맑은 고딕'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    run_url = p.add_run(f'\n      {url}')
    run_url.font.size = Pt(8)
    run_url.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run_url.font.name = '맑은 고딕'
    run_url._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ── 저장 ──
doc.save(OUTPUT_PATH)
print(f"\n문서 생성 완료: {OUTPUT_PATH}")
